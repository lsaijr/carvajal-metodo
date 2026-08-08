import os, json, uuid, threading, calendar, secrets, hashlib, time as _time
import cloudinary
import cloudinary.uploader
import bcrypt
WEASYPRINT_OK = False
WeasyprintHTML = None
from datetime import date, datetime
from flask import Flask, request, jsonify, render_template_string, send_from_directory
import requests as req
import zipfile, re, html as htmllib

app = Flask(__name__)

# ── Config desde variables de entorno ────────────────────────
CLAUDE_KEY  = os.environ.get('CLAUDE_KEY', '')
GEMINI_KEY  = os.environ.get('GEMINI_KEY', '')
GROQ_KEY    = os.environ.get('GROQ_KEY', '')
RESEND_KEY  = os.environ.get('RESEND_KEY', '')
MAIL_CC         = [m.strip() for m in os.environ.get('MAIL_CC','').split(',') if m.strip()]
MAIL_TO     = os.environ.get('MAIL_TO', 'isai.josue@gmail.com').strip()
MAIL_FROM   = os.environ.get('MAIL_FROM', 'envios@centrocarvajal.com')
DEMO_MAIL   = os.environ.get('DEMO_MAIL', 'isai.josue@gmail.com')
PLANES_DIR  = os.path.join(os.path.dirname(__file__), 'planes_generados')
os.makedirs(PLANES_DIR, exist_ok=True)

# ── Inicializar usuarios al primer request ────────────────────
_usuarios_inicializados = False

@app.before_request
def _init_on_first_request():
    global _usuarios_inicializados
    if not _usuarios_inicializados:
        _usuarios_inicializados = True
        try:
            _inicializar_usuarios()
        except Exception as e:
            print(f'[startup] Error inicializando usuarios: {e}')

# ── Auth — sesiones en memoria con expiración 8h ─────────────────────────
# {token: {'email': str, 'expires': timestamp}}
auth_sessions = {}
SESSION_DURATION = 8 * 3600  # 8 horas en segundos

# ── Reset tokens temporales — {token: {'email': str, 'expires': timestamp}}
reset_tokens = {}
RESET_DURATION = 3600  # 1 hora

# ── Reportes — acceso protegido solo con contraseña ───────────────────────
reporte_sessions = {}
REPORTE_SESSION_DURATION = 8 * 3600  # 8 horas
REPORTE_PASSWORD = os.environ.get('ADMIN_PASSWORD', 'reportes2026')

# ── Cloudinary ────────────────────────────────────────────────
CLOUDINARY_CLOUD_NAME = os.environ.get('CLOUDINARY_CLOUD_NAME', '')
CLOUDINARY_API_KEY    = os.environ.get('CLOUDINARY_API_KEY', '')
CLOUDINARY_API_SECRET = os.environ.get('CLOUDINARY_API_SECRET', '')

if CLOUDINARY_CLOUD_NAME:
    cloudinary.config(
        cloud_name = CLOUDINARY_CLOUD_NAME,
        api_key    = CLOUDINARY_API_KEY,
        api_secret = CLOUDINARY_API_SECRET,
        secure     = True
    )

# ── Jobs en memoria ───────────────────────────────────────────
jobs = {}  # jobId -> {'status': ..., 'msg': ..., 'html_url': ...}

# ── Fallback de créditos Claude ────────────────────────────────
# Lista thread-safe: se llena cuando Claude devuelve 400 sin créditos.
# El worker la lee, envía alerta y la limpia.
_claude_sin_creditos = []

# ════════════════════════════════════════════════════════════
# RUTAS
# ════════════════════════════════════════════════════════════

@app.route('/', methods=['GET'])
def index():
    # Landing page — Método Carvajal
    with open(os.path.join(os.path.dirname(__file__), 'index-carvajal.html'), encoding='utf-8') as f:
        return f.read()


@app.route('/plan-diario', methods=['GET'])
def plan_diario():
    # Plan diario personal con seguimiento (localStorage)
    with open(os.path.join(os.path.dirname(__file__), 'plan-diario.html'), encoding='utf-8') as f:
        return f.read()


@app.route('/reportes/', methods=['GET', 'POST'])
def reportes():
    """Reporte protegido con solo contraseña (sin usuario)."""
    from flask import make_response
    tok = request.cookies.get('cv_reporte_session', '')
    now = _time.time()
    # Limpiar sesiones expiradas
    expiradas = [t for t, s in reporte_sessions.items() if s['expires'] < now]
    for t in expiradas:
        del reporte_sessions[t]
    sesion = reporte_sessions.get(tok)

    if request.method == 'POST':
        pwd = request.form.get('password', '').strip()
        if pwd == REPORTE_PASSWORD:
            tok = secrets.token_hex(32)
            reporte_sessions[tok] = {'expires': now + REPORTE_SESSION_DURATION}
            with open(os.path.join(os.path.dirname(__file__), 'reportes.html'), encoding='utf-8') as f:
                html = f.read()
            resp = make_response(html)
            resp.set_cookie('cv_reporte_session', tok, max_age=REPORTE_SESSION_DURATION, httponly=True, samesite='Lax')
            return resp
        return _reportes_login_form(error='Contraseña incorrecta.'), 401

    if sesion:
        with open(os.path.join(os.path.dirname(__file__), 'reportes.html'), encoding='utf-8') as f:
            return f.read()

    return _reportes_login_form()


def _reportes_login_form(error=''):
    error_html = f'<div class="error">{error}</div>' if error else ''
    return f'''<!DOCTYPE html>
<html lang="es">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>Acceso a reportes · Centro Carvajal</title>
<style>
*{{margin:0;padding:0;box-sizing:border-box}}
body{{font-family:'Segoe UI',sans-serif;background:#f4f5ef;min-height:100vh;display:flex;align-items:center;justify-content:center;padding:20px}}
.box{{background:#fff;border-radius:12px;padding:44px 40px;max-width:400px;width:100%;box-shadow:0 8px 40px rgba(0,0,0,.1);text-align:center}}
h1{{font-size:22px;color:#1a1410;margin-bottom:8px}}
p{{font-size:13px;color:#6b7280;margin-bottom:28px}}
input{{width:100%;padding:13px 16px;border:1px solid #d4dcc0;border-radius:8px;font-size:14px;margin-bottom:14px;font-family:inherit;outline:none}}
input:focus{{border-color:#8fa832}}
button{{width:100%;padding:14px;background:#1a1410;color:#b8935a;border:none;border-radius:8px;font-size:14px;font-weight:500;cursor:pointer;font-family:inherit}}
.error{{background:#fef2f2;color:#c0392b;border:1px solid #fecaca;padding:10px;border-radius:6px;margin-bottom:14px;font-size:13px}}
</style>
</head>
<body>
<div class="box">
  <h1>Reportes Carvajal</h1>
  <p>Ingresa la contraseña para acceder al reporte.</p>
  <form method="POST" action="/reportes/">
    {error_html}
    <input type="password" name="password" placeholder="Contraseña" autofocus>
    <button type="submit">Ingresar</button>
  </form>
</div>
</body>
</html>'''


@app.route('/panel', methods=['GET'])
def panel():
    from flask import redirect
    # Si ya tiene sesión válida y viene con ?next=, redirigir directo
    tok = request.cookies.get('cv_session', '')
    sesion = _verificar_sesion(tok)
    next_url = request.args.get('next', '').strip()
    if sesion and next_url and next_url.startswith('/'):
        return redirect(next_url)
    # Pasar el ?next al HTML para que el JS lo capture post-login
    with open(os.path.join(os.path.dirname(__file__), 'panel-carvajal.html'), encoding='utf-8') as f:
        html = f.read()
    if next_url:
        # Inyectar el next_url en el JS para capturarlo tras login exitoso
        html = html.replace('</head>', f'<script>window._cv_next="{next_url}";</script></head>', 1)
    return html

# ── Demo — formulario estético simplificado ──────────────────
@app.route('/demo')
def demo():
    with open(os.path.join(os.path.dirname(__file__), 'formulario-estetica-v2.html'), encoding='utf-8') as f:
        return f.read()

@app.route('/demo/recomendar', methods=['POST'])
def demo_recomendar():
    """Proxy multi-modelo para el formulario demo — las keys nunca salen del servidor."""
    try:
        data   = request.get_json(force=True)
        perfil = data.get('perfil', '')
        modelo = data.get('modelo', 'claude')
        if not perfil:
            return jsonify({'error': 'Sin datos de perfil'}), 400

        base_prompt = (
            'Eres una especialista en estética con años de experiencia. Tu forma de comunicarte es cálida, respetuosa y genuinamente personalizada. '
            'Le hablas al cliente por su nombre y usas "tú". '
            'TONO: profesional pero humano — como una especialista que realmente leyó su caso y le habla con consideración, no como un robot generando texto genérico ni como una amiga informal. '
            'EVITA absolutamente: frases genéricas como "¡Hola!", exclamaciones excesivas, emojis en el texto, frases vacías como "es un placer atenderte", lenguaje clínico frío, o recomendaciones que podrían ser para cualquier persona. '
            'BUSCA: oraciones que demuestren que leíste el perfil específico del cliente — menciona su situación real, sus áreas de interés, su nivel de estrés o actividad física cuando sea relevante. Que cada frase aporte algo concreto. '
            'Responde SIEMPRE en HTML simple usando solo estas etiquetas: <p>, <strong>, <ul>, <li>. '
            'NO uses markdown, NO uses encabezados, NO uses tablas. '
            'Estructura: '
            '1. Abre mencionando el nombre del cliente y una observación específica sobre su perfil que demuestre que lo leíste — algo como reconocer su preocupación principal o su situación particular. '
            '2. Explica brevemente por qué los tratamientos que vas a recomendar tienen sentido para SU caso específico, en lenguaje claro sin tecnicismos. '
            '3. Lista los tratamientos recomendados — nombre del tratamiento en negrita y una frase concreta sobre qué resultado puede esperar esta persona en particular. '
            '4. Cierra con una invitación a agendar su consulta, transmitiendo que en ese espacio podrán profundizar y resolver todas sus dudas.'
        )
        if modelo == 'groq':
            sys_prompt = base_prompt + (
                ' IMPORTANTE: sé muy detallada y personalizada. '
                'Integra datos específicos del perfil: edad, nivel de estrés, actividad física, historial de tratamientos previos. '
                'Cada tratamiento merece 2-3 oraciones — qué hace, por qué encaja con este perfil y qué resultado concreto y realista puede esperar. '
                'Máximo 450 palabras. Que al leerlo sienta que fue escrito para ella o él, no para cualquier persona.'
            )
        else:
            sys_prompt = base_prompt + ' Máximo 300 palabras. Cada frase debe aportar valor concreto — nada de relleno.'
        user_msg = 'Perfil del paciente:\n' + perfil

        # ── Claude ──────────────────────────────────────────
        if modelo == 'claude':
            r = req.post('https://api.anthropic.com/v1/messages',
                headers={'Content-Type':'application/json','x-api-key':CLAUDE_KEY,'anthropic-version':'2023-06-01'},
                json={'model':'claude-haiku-4-5-20251001','max_tokens':800,
                      'system':sys_prompt,'messages':[{'role':'user','content':user_msg}]},
                timeout=30)
            j = r.json()
            if 'content' in j and j['content']:
                return jsonify({'html': j['content'][0]['text']})
            return jsonify({'error': 'Sin respuesta de Claude', 'detail': j}), 500

        # ── Gemini ──────────────────────────────────────────
        elif modelo == 'gemini':
            url = f'https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash-lite:generateContent?key={GEMINI_KEY}'
            r = req.post(url,
                headers={'Content-Type':'application/json'},
                json={'system_instruction':{'parts':[{'text':sys_prompt}]},
                      'contents':[{'parts':[{'text':user_msg}]}],
                      'generationConfig':{'maxOutputTokens':800}},
                timeout=30)
            j = r.json()
            print(f'[demo/gemini] status={r.status_code} response={json.dumps(j)[:300]}')
            if 'candidates' in j and j['candidates']:
                return jsonify({'html': j['candidates'][0]['content']['parts'][0]['text']})
            # Devolver el error completo al frontend para debug
            error_msg = j.get('error', {}).get('message', 'Sin respuesta de Gemini')
            return jsonify({'error': f'Gemini: {error_msg}', 'detail': j}), 500

        # ── Groq ─────────────────────────────────────────────
        elif modelo == 'groq':
            r = req.post('https://api.groq.com/openai/v1/chat/completions',
                headers={'Content-Type':'application/json','Authorization':f'Bearer {GROQ_KEY}'},
                json={'model':'llama-3.3-70b-versatile','max_tokens':800,
                      'messages':[{'role':'system','content':sys_prompt},{'role':'user','content':user_msg}]},
                timeout=30)
            j = r.json()
            if 'choices' in j and j['choices']:
                return jsonify({'html': j['choices'][0]['message']['content']})
            return jsonify({'error': 'Sin respuesta de Groq', 'detail': j}), 500

        return jsonify({'error': f'Modelo desconocido: {modelo}'}), 400

    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/demo/cita', methods=['POST'])
def demo_cita():
    """Recibe solicitud de cita del demo y envía correo a DEMO_MAIL."""
    try:
        d = request.get_json(force=True)

        # Construir cuerpo del correo en HTML
        nombre   = d.get('nombre','')
        cuerpo = f"""
<div style="font-family:Arial,sans-serif;max-width:600px;margin:0 auto;color:#1a1a18">

<div style="background:#f0f7ff;border-bottom:3px solid #4299e1;padding:24px 28px;border-radius:8px 8px 0 0">
  <h2 style="margin:0;font-size:20px;color:#2d3748">📋 Nueva solicitud de cita</h2>
  <p style="margin:6px 0 0;font-size:13px;color:#718096">Formulario demo · metodo.centrocarvajal.com/demo</p>
</div>

<div style="background:#ffffff;padding:24px 28px;border:1px solid #bee3f8;border-top:none">

  <h3 style="font-size:12px;color:#4299e1;text-transform:uppercase;letter-spacing:.1em;margin:0 0 12px">👤 Datos del paciente</h3>
  <table style="font-size:14px;border-collapse:collapse;width:100%;margin-bottom:24px">
    <tr><td style="padding:5px 0;color:#718096;width:160px">Nombre</td><td style="padding:5px 0"><strong>{nombre}</strong></td></tr>
    <tr><td style="padding:5px 0;color:#718096">Correo</td><td style="padding:5px 0">{d.get('email','')}</td></tr>
    <tr><td style="padding:5px 0;color:#718096">WhatsApp</td><td style="padding:5px 0">{d.get('tel','') or '—'}</td></tr>
    <tr><td style="padding:5px 0;color:#718096">Edad</td><td style="padding:5px 0">{d.get('edad','')}</td></tr>
    <tr><td style="padding:5px 0;color:#718096">Género</td><td style="padding:5px 0">{d.get('genero','')}</td></tr>
    {f'<tr><td style="padding:5px 0;color:#718096">Embarazo/lactancia</td><td style="padding:5px 0">{d.get("embarazo","")}</td></tr>' if d.get('genero','') == 'mujer' else ''}
    <tr><td style="padding:5px 0;color:#718096">Cómo nos conoció</td><td style="padding:5px 0">{d.get('referido','') or '—'}</td></tr>
  </table>

  <h3 style="font-size:12px;color:#4299e1;text-transform:uppercase;letter-spacing:.1em;margin:0 0 12px">🎯 Perfil estético</h3>
  <table style="font-size:14px;border-collapse:collapse;width:100%;margin-bottom:24px">
    <tr><td style="padding:5px 0;color:#718096;width:160px">Áreas de interés</td><td style="padding:5px 0">{d.get('areas','') or '—'}</td></tr>
    <tr><td style="padding:5px 0;color:#718096">Problemas faciales</td><td style="padding:5px 0">{d.get('piel_checks','') or '—'}</td></tr>
    <tr><td style="padding:5px 0;color:#718096">Tono de piel</td><td style="padding:5px 0">{d.get('tono','') or '—'}</td></tr>
    <tr><td style="padding:5px 0;color:#718096">Zonas corporales</td><td style="padding:5px 0">{d.get('cuerpo_zona','') or '—'}</td></tr>
    <tr><td style="padding:5px 0;color:#718096">Objetivos corporales</td><td style="padding:5px 0">{d.get('cuerpo_obj','') or '—'}</td></tr>
    <tr><td style="padding:5px 0;color:#718096">Zonas de depilación</td><td style="padding:5px 0">{d.get('vello_zonas','') or '—'}</td></tr>
    <tr><td style="padding:5px 0;color:#718096">Situación capilar</td><td style="padding:5px 0">{d.get('capilar','') or '—'}</td></tr>
    <tr><td style="padding:5px 0;color:#718096">Tratamientos previos</td><td style="padding:5px 0">{d.get('trat_prev','') or 'Ninguno'}</td></tr>
    <tr><td style="padding:5px 0;color:#718096">Mayor preocupación</td><td style="padding:5px 0">{d.get('prioridad','') or '—'}</td></tr>
  </table>

  <h3 style="font-size:12px;color:#4299e1;text-transform:uppercase;letter-spacing:.1em;margin:0 0 12px">💡 Hábitos y preferencias</h3>
  <table style="font-size:14px;border-collapse:collapse;width:100%;margin-bottom:24px">
    <tr><td style="padding:5px 0;color:#718096;width:160px">Actividad física</td><td style="padding:5px 0">{d.get('ejercicio','') or '—'}</td></tr>
    <tr><td style="padding:5px 0;color:#718096">Nivel de estrés</td><td style="padding:5px 0">{d.get('estres','')}/10</td></tr>
    <tr><td style="padding:5px 0;color:#718096">Urgencia</td><td style="padding:5px 0">{d.get('urgencia','') or '—'}</td></tr>
    <tr><td style="padding:5px 0;color:#718096">Presupuesto mensual</td><td style="padding:5px 0">{d.get('presupuesto','') or '—'}</td></tr>
  </table>

  <h3 style="font-size:12px;color:#4299e1;text-transform:uppercase;letter-spacing:.1em;margin:0 0 12px">📅 Cita solicitada</h3>
  <table style="font-size:14px;border-collapse:collapse;width:100%;margin-bottom:8px">
    <tr><td style="padding:5px 0;color:#718096;width:160px">Fecha</td><td style="padding:5px 0"><strong>{d.get('fecha','')}</strong></td></tr>
    <tr><td style="padding:5px 0;color:#718096">Horario</td><td style="padding:5px 0"><strong>{d.get('horario','')}</strong></td></tr>
    <tr><td style="padding:5px 0;color:#718096">Nota adicional</td><td style="padding:5px 0">{d.get('nota','') or '—'}</td></tr>
  </table>

</div>

<div style="background:#f0f7ff;padding:14px 28px;border:1px solid #bee3f8;border-top:none;border-radius:0 0 8px 8px;text-align:center">
  <p style="font-size:11px;color:#718096;margin:0">Evaluación IA · metodo.centrocarvajal.com/demo</p>
</div>

</div>
"""

        resend_payload = {
            'from':    f'Evaluación IA <{MAIL_FROM}>',
            'to':      [DEMO_MAIL],
            'subject': f'Nueva solicitud de cita para — {nombre}',
            'html':    cuerpo
        }

        r = req.post('https://api.resend.com/emails',
            headers={'Authorization': f'Bearer {RESEND_KEY}', 'Content-Type': 'application/json'},
            json=resend_payload, timeout=15)

        print(f'[demo/cita] resend status={r.status_code} body={r.text[:200]}')

        if r.status_code == 200:
            return jsonify({'ok': True})
        return jsonify({'ok': False, 'detail': r.text}), 500

    except Exception as e:
        print(f'[demo/cita] error: {e}')
        return jsonify({'ok': False, 'error': str(e)}), 500

@app.route('/formulario', methods=['GET'])
def formulario():
    # Formulario de producción — Claude hardcodeado, sin selector de modelo
    with open(os.path.join(os.path.dirname(__file__), 'formulario-produccion.html'), encoding='utf-8') as f:
        return f.read()


@app.route('/formulario-demo', methods=['GET'])
def formulario_demo():
    # Formulario de pruebas — con selector de modelo (Claude / Gemini / Llama)
    with open(os.path.join(os.path.dirname(__file__), 'formulario-demo.html'), encoding='utf-8') as f:
        return f.read()


@app.route('/formulario-seleccion', methods=['GET'])
def formulario_seleccion():
    # Herramienta interna: revisión de campos del formulario para decidir qué se mantiene/quita
    return send_from_directory(os.path.dirname(os.path.abspath(__file__)), 'formulario-seleccion.html')


@app.route('/api/formulario-seleccion/enviar', methods=['POST'])
def api_formulario_seleccion_enviar():
    data = request.get_json(silent=True) or {}
    campos = data.get('campos', [])
    if not isinstance(campos, list) or not campos:
        return jsonify({'error': 'Sin campos para enviar'}), 400

    def esc(s): return str(s).replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

    label_estado = {'mantener': ('✓ Mantener', '#3f8f4f', '#eaf6ec'), 'quitar': ('✕ Quitar', '#c0392b', '#fbeceb'), 'sin_marcar': ('○ Sin marcar', '#8a7a68', '#f2ede4')}
    pasos = {}
    for c in campos:
        pasos.setdefault(c.get('paso', ''), []).append(c)

    secciones_html = ''
    for paso, items in pasos.items():
        filas = ''
        for c in items:
            txt, color, bg = label_estado.get(c.get('decision', 'sin_marcar'), label_estado['sin_marcar'])
            filas += f'<tr><td style="padding:7px 12px;font-size:13px;border-bottom:1px solid #eee">{esc(c.get("campo",""))}</td><td style="padding:7px 12px;font-size:12px;font-weight:600;color:{color};background:{bg};border-bottom:1px solid #eee;white-space:nowrap">{txt}</td></tr>'
        secciones_html += f'<div style="margin:18px 24px 0"><div style="font-size:12px;font-weight:700;color:#b8935a;text-transform:uppercase;letter-spacing:1px;margin-bottom:6px">{esc(paso)}</div><table style="width:100%;border-collapse:collapse;background:#fff;border:1px solid #eee;border-radius:4px;overflow:hidden">{filas}</table></div>'

    total_mantener = sum(1 for c in campos if c.get('decision') == 'mantener')
    total_quitar = sum(1 for c in campos if c.get('decision') == 'quitar')
    total_pendiente = sum(1 for c in campos if c.get('decision') not in ('mantener', 'quitar'))

    cuerpo = f'''<!DOCTYPE html><html><head><meta charset="UTF-8"></head><body style="background:#f0e8de;padding:20px;font-family:sans-serif">
<div style="max-width:640px;margin:0 auto;background:#fff;border:1px solid #ddd">
<div style="background:#1a1410;padding:20px 24px">
<div style="color:#b8935a;font-size:11px;letter-spacing:3px;text-transform:uppercase">Centro Carvajal · Revisión de Formulario</div>
<div style="color:#fff;font-size:18px;margin-top:4px">Resultado de selección de campos</div>
<div style="color:rgba(255,255,255,.5);font-size:12px;margin-top:6px">{total_mantener} mantener · {total_quitar} quitar · {total_pendiente} sin marcar</div>
</div>
{secciones_html}
<div style="background:#1a1410;padding:12px 24px;text-align:center;font-size:10px;color:rgba(255,255,255,.3);margin-top:20px">Centro Carvajal · centrocarvajal.com</div>
</div></body></html>'''

    try:
        enviar_resend('Revisión de campos — Formulario Método Carvajal', cuerpo, 'isai.josue@gmail.com')
        return jsonify({'ok': True})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/planes_generados/<path:filename>')
def serve_plan(filename):
    return send_from_directory(PLANES_DIR, filename)

@app.route('/status')
def status():
    job_id = request.args.get('job', '')
    job = jobs.get(job_id, {'status': 'working', 'msg': 'Procesando...', 'pct': 5})
    return jsonify(job)


# ── Endpoint formulario web (/enviar) ─────────────────────────
@app.route('/enviar', methods=['POST'])
def enviar():
    raw = request.form.get('data', '')
    if not raw:
        return jsonify({'error': 'No se recibieron datos del formulario'}), 400

    try:
        form = json.loads(raw)
    except Exception:
        return jsonify({'error': 'JSON invalido en los datos del formulario'}), 400

    # Sanitizar entradas antes de mapear — previene inyecciones y caracteres de control
    form = _sanitizar_form(form)
    data = _mapear_formulario(form)

    if not data.get('nombre'):
        return jsonify({'error': 'El campo Nombre completo es obligatorio'}), 400

    # Guardar fotos opcionales (se envían al correo del staff)
    fotos = []
    for i in range(1, 5):
        foto = request.files.get(f'foto_{i}')
        if foto:
            tmp = f'/tmp/foto_{uuid.uuid4().hex}_{foto.filename}'
            foto.save(tmp)
            fotos.append(tmp)

    job_id = uuid.uuid4().hex[:16]
    jobs[job_id] = {'status': 'working', 'msg': 'Iniciando generacion del plan...'}

    modelo = request.form.get('modelo', 'claude')  # viene del FormData, no del JSON

    # session_id: usar el que manda el frontend (guardado progresivo) o generar uno nuevo
    session_id = request.form.get('session_id', '').strip() or None
    if not session_id:
        session_id = f'carvajal_{uuid.uuid4().hex[:16]}'

    # ── Guardar sesión completa AHORA — antes de cualquier otra cosa ──
    # Si la IA falla, Railway se reinicia, o cualquier otra cosa, los datos quedan en Cloudinary
    try:
        guardar_sesion_cloudinary(session_id, data)
        print(f'[enviar] Sesión guardada: {session_id}')
    except Exception as e:
        print(f'[enviar] Error guardando sesión: {e}')

    # ── Generar .docx del cuestionario para adjuntar al Correo 1 ──
    # Así el staff siempre tiene el respaldo completo independientemente de la IA
    docx_respaldo = None
    try:
        docx_respaldo = generar_docx_cuestionario(data)
        print(f'[enviar] .docx respaldo generado: {docx_respaldo}')
    except Exception as e:
        print(f'[enviar] Error generando .docx respaldo: {e}')

    # ── CORREO 1: formulario inmediato + .docx adjunto ──
    # Sale antes de lanzar la IA. Si la IA falla, el staff tiene todo aquí.
    adjuntos_correo1 = list(fotos or [])
    if docx_respaldo and os.path.exists(docx_respaldo):
        adjuntos_correo1.append(docx_respaldo)
    try:
        fecha_recibido = datetime.now().strftime('%d/%m/%Y a las %H:%M')
        enviar_resend(
            f'📋 Formulario recibido — {data["nombre"]} ({fecha_recibido})',
            email_formulario_inmediato(data, fotos),
            MAIL_TO,
            adjuntos_extra=adjuntos_correo1,
            cc=MAIL_CC or None
        )
        print(f'[enviar] Correo 1 (formulario) enviado: {data["nombre"]}')
    except Exception as e:
        print(f'[enviar] Error correo 1: {e}')

    # Limpiar .docx respaldo después de enviarlo
    if docx_respaldo:
        try: os.unlink(docx_respaldo)
        except: pass

    t = threading.Thread(target=worker, args=(job_id, data, [], fotos, modelo, session_id), daemon=True)
    t.start()

    return jsonify({'jobId': job_id, 'nombre': data['nombre'], 'modelo': modelo, 'session_id': session_id})


# ── Endpoint carga .docx (/upload) ────────────────────────────
@app.route('/upload', methods=['POST'])
def upload():
    if 'docx' not in request.files:
        return jsonify({'error': 'No se recibio archivo'}), 400

    f = request.files['docx']
    tmp_path = f'/tmp/carvajal_{uuid.uuid4().hex}.docx'
    f.save(tmp_path)

    texto = leer_docx(tmp_path)
    os.unlink(tmp_path)

    if not texto:
        return jsonify({'error': 'No se pudo leer el archivo .docx'}), 400

    data, faltantes = parsear_cuestionario(texto)
    if not data.get('nombre'):
        return jsonify({'error': 'No se encontro nombre_completo en el documento'}), 400

    job_id = uuid.uuid4().hex[:16]
    jobs[job_id] = {'status': 'working', 'msg': 'Iniciando generacion del plan...'}

    modelo = request.form.get('modelo', 'claude')
    t = threading.Thread(target=worker, args=(job_id, data, faltantes, [], modelo), daemon=True)
    t.start()

    return jsonify({'jobId': job_id, 'nombre': data['nombre'], 'modelo': modelo})


# ════════════════════════════════════════════════════════════
# RUTA BORRADOR EDITABLE
# ════════════════════════════════════════════════════════════

@app.route('/borrador/<job_id>', methods=['GET'])
def ver_borrador(job_id):
    """Sirve el borrador editable. Descarga desde Cloudinary si no está en memoria."""
    from flask import Response
    html = descargar_borrador_cloudinary(job_id)
    if html:
        return Response(html, content_type='text/html; charset=utf-8')
    job = jobs.get(job_id)
    if not job or job.get('status') != 'done':
        return Response('<h2 style="font-family:sans-serif;padding:40px;color:#666">Borrador no encontrado o aún procesando.</h2>', status=404, content_type='text/html; charset=utf-8')
    return Response('<h2 style="font-family:sans-serif;padding:40px;color:#666">Borrador no disponible. Verifica Cloudinary.</h2>', status=404, content_type='text/html; charset=utf-8')


@app.route('/guardar/<job_id>', methods=['POST'])
def guardar_borrador(job_id):
    """Recibe el HTML completo editado y lo guarda en Cloudinary."""
    from flask import Response as R
    html = request.get_data(as_text=True)
    if not html:
        return jsonify({'ok': False, 'error': 'HTML vacío'}), 400
    url = subir_borrador_cloudinary(html, job_id)
    # Actualizar nombre en jobs si existe
    if job_id in jobs:
        jobs[job_id]['borrador_actualizado'] = True
    return jsonify({'ok': True, 'url': url or ''})


# ── Sesiones — guardado progresivo del formulario ────────────

@app.route('/sesion/guardar', methods=['POST'])
def sesion_guardar():
    """Guarda o actualiza los datos parciales del formulario en Cloudinary."""
    try:
        d = request.get_json(force=True)
        session_id = d.get('session_id', '').strip()
        data_parcial = d.get('data', {})
        if not session_id:
            return jsonify({'ok': False, 'error': 'session_id requerido'}), 400
        # Solo permitir caracteres seguros en el session_id
        if not re.match(r'^[a-zA-Z0-9_\-]{5,80}$', session_id):
            return jsonify({'ok': False, 'error': 'session_id inválido'}), 400
        url = guardar_sesion_cloudinary(session_id, data_parcial)
        return jsonify({'ok': True, 'url': url or ''})
    except Exception as e:
        return jsonify({'ok': False, 'error': str(e)}), 500


@app.route('/sesion/<session_id>', methods=['GET'])
def sesion_recuperar(session_id):
    """Recupera los datos guardados de una sesión."""
    if not re.match(r'^[a-zA-Z0-9_\-]{5,80}$', session_id):
        return jsonify({'ok': False, 'error': 'session_id inválido'}), 400
    data = recuperar_sesion_cloudinary(session_id)
    if data is None:
        return jsonify({'ok': False, 'error': 'Sesión no encontrada o expirada'}), 404
    return jsonify({'ok': True, 'data': data})





# ════════════════════════════════════════════════════════════
# WORKER
# ════════════════════════════════════════════════════════════

def subir_plan_cloudinary(html_path, html_name, job_id=''):
    """Sube el HTML del plan a Cloudinary como raw file.
    Devuelve la URL pública o None si falla."""
    if not CLOUDINARY_CLOUD_NAME:
        print('Cloudinary no configurado, usando URL local')
        return None
    try:
        resultado = cloudinary.uploader.upload(
            html_path,
            folder        = 'carvajal/planes',
            public_id     = re.sub(r'\.html$', '', html_name),
            resource_type = 'raw',
            overwrite     = True,
            context       = f'job_id={job_id}' if job_id else None,
        )
        url = resultado.get('secure_url', '')
        print(f'Cloudinary OK: {url[:80]}')
        return url
    except Exception as e:
        print(f'Cloudinary error: {e}')
        return None


def subir_borrador_cloudinary(html_content, job_id):
    """Sube el HTML del borrador a Cloudinary como raw file."""
    if not CLOUDINARY_CLOUD_NAME:
        return None
    try:
        import tempfile
        with tempfile.NamedTemporaryFile(mode='wb', suffix='.html', delete=False) as tmp:
            tmp.write(html_content.encode('utf-8'))  # escribir bytes UTF-8 explícito
            tmp_path = tmp.name
        resultado = cloudinary.uploader.upload(
            tmp_path,
            folder='carvajal/borradores',
            public_id=f'borrador_{job_id}',
            resource_type='raw',
            overwrite=True,
            invalidate=True,  # invalida cache CDN para que _cargar_usuarios lea version fresca
        )
        os.unlink(tmp_path)
        return resultado.get('secure_url', '')
    except Exception as e:
        print(f'Error subiendo borrador: {e}')
        return None


def descargar_borrador_cloudinary(job_id):
    """Descarga el HTML del borrador desde Cloudinary."""
    if not CLOUDINARY_CLOUD_NAME:
        return None
    try:
        url = f'https://res.cloudinary.com/{CLOUDINARY_CLOUD_NAME}/raw/upload/carvajal/borradores/borrador_{job_id}.html'
        r = req.get(url, timeout=30)
        if r.status_code == 200:
            return r.content.decode('utf-8')  # forzar UTF-8, Cloudinary no declara charset
        return None
    except Exception as e:
        print(f'Error descargando borrador: {e}')
        return None


def render_borrador(plan_json, data, job_id):
    """Genera el HTML del borrador — ahora usa la misma plantilla que render_plan."""
    return render_plan(plan_json, data, job_id=job_id)


# ════════════════════════════════════════════════════════════
# SESIONES — persistencia de formulario en Cloudinary
# ════════════════════════════════════════════════════════════

SESION_TTL_HORAS = 72  # sesiones expiran a las 72 horas (cubre fines de semana)

def _session_public_id(session_id):
    return f'carvajal/sesiones/{session_id}'

def guardar_sesion_cloudinary(session_id, data_parcial):
    """Guarda o actualiza los datos parciales del formulario como JSON en Cloudinary."""
    if not CLOUDINARY_CLOUD_NAME:
        return None
    try:
        import tempfile, json as _json
        payload = {
            'session_id': session_id,
            'guardado_en': datetime.utcnow().isoformat() + 'Z',
            'data': data_parcial
        }
        with tempfile.NamedTemporaryFile(mode='w', suffix='.json', delete=False, encoding='utf-8') as tmp:
            _json.dump(payload, tmp, ensure_ascii=False)
            tmp_path = tmp.name
        resultado = cloudinary.uploader.upload(
            tmp_path,
            folder='carvajal/sesiones',
            public_id=session_id,
            resource_type='raw',
            overwrite=True,
        )
        os.unlink(tmp_path)
        print(f'[sesion] Guardada: {session_id}')
        return resultado.get('secure_url', '')
    except Exception as e:
        print(f'[sesion] Error guardando: {e}')
        return None


def recuperar_sesion_cloudinary(session_id):
    """Descarga y devuelve los datos de sesión guardados. Retorna None si no existe o expiró."""
    if not CLOUDINARY_CLOUD_NAME:
        return None
    try:
        import json as _json
        url = f'https://res.cloudinary.com/{CLOUDINARY_CLOUD_NAME}/raw/upload/carvajal/sesiones/{session_id}.json'
        r = req.get(url, timeout=10)
        if r.status_code != 200:
            return None
        payload = r.json()
        # Verificar expiración (24 horas)
        guardado = payload.get('guardado_en', '')
        if guardado:
            from datetime import timezone
            ts = datetime.fromisoformat(guardado.replace('Z', '+00:00'))
            delta = datetime.now(timezone.utc) - ts
            if delta.total_seconds() > SESION_TTL_HORAS * 3600:
                print(f'[sesion] Expirada: {session_id} ({round(delta.total_seconds()/3600, 1)}h)')
                return None
        return payload.get('data')
    except Exception as e:
        print(f'[sesion] Error recuperando: {e}')
        return None


def eliminar_sesion_cloudinary(session_id):
    """Elimina la sesión de Cloudinary una vez que el plan fue generado exitosamente."""
    if not CLOUDINARY_CLOUD_NAME:
        return
    try:
        cloudinary.uploader.destroy(
            f'carvajal/sesiones/{session_id}',
            resource_type='raw',
            invalidate=True
        )
        print(f'[sesion] Eliminada: {session_id}')
    except Exception as e:
        print(f'[sesion] Error eliminando: {e}')


def _render_borrador_legacy(plan_json, data, job_id):
    """LEGACY — ya no se usa. render_borrador ahora llama render_plan."""
    tpl_path = os.path.join(os.path.dirname(__file__), 'plantilla_borrador.html')
    with open(tpl_path, encoding='utf-8') as f:
        tpl = f.read()

    def esc(s): return str(s).replace('&','&amp;').replace('<','&lt;').replace('>','&gt;').replace('"','&quot;')

    nombre = data.get('nombre', '')

    # Diagnóstico rows
    badge_map = {'warning': 'badge-warning', 'critical': 'badge-critical', 'normal': 'badge-normal'}
    badge_label = {'warning': '⚠ Atención', 'critical': '✕ Crítico', 'normal': 'Normal'}
    diag_html = ''.join(
        f'<tr><td style="font-weight:500">{esc(f.get("area",""))}</td><td contenteditable="true">{esc(f.get("estado",""))}</td><td contenteditable="true">{esc(f.get("hallazgos",""))}</td><td><span class="{badge_map.get(f.get("alerta","normal"),"badge-normal")}">{badge_label.get(f.get("alerta","normal"),"Normal")}</span></td></tr>'
        for f in plan_json.get('diagnostico', {}).get('filas', [])
    )

    # Rutina
    tag_map = {'Nutricion':'tag-n','Sueno':'tag-s','Actividad':'tag-a','Mental':'tag-m','Estetico':'tag-e','Salud':'tag-h'}
    rutina_html = ''.join(
        f'<div class="rutina-row"><span class="rutina-hora">{esc(r["hora"])}</span><span contenteditable="true" style="flex:1">{esc(r["actividad"])}</span><span class="rutina-tag {tag_map.get(r["pilar"],"tag-n")}">{esc(r["pilar"])}</span></div>'
        for r in plan_json.get('rutina', {}).get('items', [])
    )

    # Pilares helpers
    p1 = plan_json.get('pilar1', {})
    p2 = plan_json.get('pilar2', {})
    p3 = plan_json.get('pilar3', {})
    p4 = plan_json.get('pilar4', {})
    p5 = plan_json.get('pilar5', {})
    comp = plan_json.get('compromiso', {})

    p1_perm = ''.join(f'<li contenteditable="true">{esc(i)}</li>' for i in p1.get('permitidos',[]))
    p1_evit = ''.join(f'<li contenteditable="true">{esc(i)}</li>' for i in p1.get('evitar',[]))
    p1_menu = ''.join(
        f'<tr><td class="dia">{esc(m.get("dia",""))}</td><td contenteditable="true">{esc(m.get("desayuno",""))}</td><td contenteditable="true">{esc(m.get("almuerzo",""))}</td><td contenteditable="true">{esc(m.get("cena",""))}</td><td contenteditable="true">{esc(m.get("snack",""))}</td></tr>'
        for m in p1.get('menu',[])
    )
    p1_supl = ''
    if p1.get('suplementacion'):
        items = ''.join(f'<li contenteditable="true" style="font-size:13px;padding:4px 0;border-bottom:1px solid #e5e7eb">{esc(s)}</li>' for s in p1['suplementacion'])
        p1_supl = f'<h4 style="font-size:12px;font-weight:500;color:#6b7280;text-transform:uppercase;letter-spacing:.06em;margin-bottom:8px">Suplementación</h4><ul style="list-style:none">{items}</ul>'

    p3_tec = ''.join(f'<li contenteditable="true">{esc(t)}</li>' for t in p3.get('tecnicas',[]))

    p4_proto = ''.join(f'<li contenteditable="true" style="padding:4px 0;font-size:13px">{esc(s)}</li>' for s in p4.get('protocolo',[]))
    p4_reg = ''.join(f'<li contenteditable="true">{esc(r)}</li>' for r in p4.get('reglas',[]))

    # Pilar 5 bimestres
    p5_bim = ''
    for bim in p5.get('bimestres',[]):
        rows = ''.join(f'<tr><td contenteditable="true" style="font-weight:500">{esc(t.get("nombre",""))}</td><td contenteditable="true">{esc(t.get("sesiones",""))}</td><td contenteditable="true" style="font-weight:500">{esc(t.get("inversion",""))}</td><td contenteditable="true">{esc(t.get("beneficio",""))}</td></tr>' for t in bim.get('tratamientos',[]))
        p5_bim += f'<div style="margin-bottom:16px"><div style="font-size:12px;font-weight:600;color:#8fa832;text-transform:uppercase;letter-spacing:.06em;margin-bottom:6px" contenteditable="true">{esc(bim.get("periodo",""))} · {esc(bim.get("titulo",""))}</div><table><thead><tr><th>Tratamiento</th><th>Sesiones</th><th>Inversión</th><th>Beneficio</th></tr></thead><tbody>{rows}</tbody></table><div style="font-size:12px;color:#6b7280;margin-top:6px;text-align:right" contenteditable="true">Total bimestre: ${bim.get("total",0):,}</div></div>'

    notas = p5.get('notas_criticas',[])
    p5_notas = ''
    if notas:
        items = ''.join(f'<p style="font-size:13px;padding:3px 0;border-bottom:1px solid #e5e7eb" contenteditable="true">{esc(n)}</p>' for n in notas)
        p5_notas = f'<div style="background:#fff7ed;border:1px solid #fed7aa;border-radius:4px;padding:12px 16px;margin-bottom:12px"><strong style="font-size:11px;color:#9a3412;text-transform:uppercase;letter-spacing:.06em;display:block;margin-bottom:6px">Notas críticas</strong>{items}</div>'

    p5_am = ''.join(f'<div style="display:flex;gap:8px;align-items:flex-start;padding:6px 0;border-bottom:1px solid #e5e7eb"><span style="font-size:11px;font-weight:600;background:#8fa832;color:#fff;border-radius:50%;width:18px;height:18px;display:flex;align-items:center;justify-content:center;flex-shrink:0">{s.get("paso","")}</span><div><div style="font-size:12px;font-weight:500" contenteditable="true">{esc(s.get("producto",""))}</div><div style="font-size:11px;color:#6b7280" contenteditable="true">{esc(s.get("descripcion",""))}</div></div></div>' for s in p5.get('rutina_am',[]))
    p5_pm = ''.join(f'<div style="display:flex;gap:8px;align-items:flex-start;padding:6px 0;border-bottom:1px solid #e5e7eb"><span style="font-size:11px;font-weight:600;background:#2d3a2e;color:#fff;border-radius:50%;width:18px;height:18px;display:flex;align-items:center;justify-content:center;flex-shrink:0">{s.get("paso","")}</span><div><div style="font-size:12px;font-weight:500" contenteditable="true">{esc(s.get("producto",""))}</div><div style="font-size:11px;color:#6b7280" contenteditable="true">{esc(s.get("descripcion",""))}</div></div></div>' for s in p5.get('rutina_pm',[]))

    comp_res = ''.join(f'<li contenteditable="true">{esc(r["texto"])}</li>' for r in comp.get('resultados',[]))
    comp_pasos = ''.join(f'<li style="padding:4px 0;font-size:13px" contenteditable="true">{esc(p)}</li>' for p in comp.get('proximos_pasos',[]))

    replacements = {
        '{{JOB_ID}}': job_id,
        '{{NOMBRE}}': esc(nombre),
        '{{EDAD}}': esc(data.get('edad','')),
        '{{OCUPACION}}': esc(data.get('ocupacion','')),
        '{{FECHA}}': esc(data.get('fecha','')),
        '{{DIAGNOSTICO_FILAS}}': diag_html,
        '{{RUTINA_NOTA}}': esc(plan_json.get('rutina',{}).get('nota','')),
        '{{RUTINA_FILAS}}': rutina_html,
        '{{P1_TITULO}}': esc(p1.get('titulo','')),
        '{{P1_OBJETIVO}}': esc(p1.get('objetivo','')),
        '{{P1_PERMITIDOS}}': p1_perm,
        '{{P1_EVITAR}}': p1_evit,
        '{{P1_MENU}}': p1_menu,
        '{{P1_SUPLEMENTACION}}': p1_supl,
        '{{P2_TITULO}}': esc(p2.get('titulo','')),
        '{{P2_OBJETIVO}}': esc(p2.get('objetivo','')),
        '{{P2_PLAN}}': esc(p2.get('plan_semanal','')),
        '{{P2_ADAPTACIONES}}': esc(p2.get('adaptaciones','')),
        '{{P3_TITULO}}': esc(p3.get('titulo','')),
        '{{P3_OBJETIVO}}': esc(p3.get('objetivo','')),
        '{{P3_TECNICAS}}': p3_tec,
        '{{P4_TITULO}}': esc(p4.get('titulo','')),
        '{{P4_OBJETIVO}}': esc(p4.get('objetivo','')),
        '{{P4_PROTOCOLO}}': p4_proto,
        '{{P4_REGLAS}}': p4_reg,
        '{{P5_TITULO}}': esc(p5.get('titulo','')),
        '{{P5_OBJETIVO}}': esc(p5.get('objetivo','')),
        '{{P5_BIMESTRES}}': p5_bim,
        '{{P5_NOTAS_CRITICAS}}': p5_notas,
        '{{P5_RUTINA_AM}}': p5_am,
        '{{P5_RUTINA_PM}}': p5_pm,
        '{{COMP_PARRAFO}}': esc(comp.get('parrafo','')),
        '{{COMP_RESULTADOS}}': comp_res,
        '{{COMP_PASOS}}': comp_pasos,
    }
    for k, v in replacements.items():
        tpl = tpl.replace(k, v)
    return tpl



# ════════════════════════════════════════════════════════════
# ADMIN — Página /planes con login
# ════════════════════════════════════════════════════════════

@app.route('/planes', methods=['GET'])
def admin_planes():
    return PLANES_HTML

PLANES_HTML = """<!DOCTYPE html>
<html lang="es">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>Centro Carvajal · Planes Generados</title>
<style>
:root{
  --dark:#1a1410;--olive:#8fa832;--olive-light:rgba(143,168,50,.1);
  --gold:#b8935a;--cream:#f4f5ef;--white:#fff;--gray:#6b7280;
  --border:rgba(143,168,50,.2);--red:#dc2626;
}
*{margin:0;padding:0;box-sizing:border-box}
body{font-family:'Segoe UI',system-ui,sans-serif;background:var(--cream);color:var(--dark);min-height:100vh}

/* ── LOGIN ── */
#login-screen{
  display:flex;align-items:center;justify-content:center;
  min-height:100vh;padding:20px;
}
.login-box{
  background:var(--white);border-radius:12px;padding:44px 40px;
  max-width:400px;width:100%;box-shadow:0 8px 40px rgba(0,0,0,.1);
  text-align:center;
}
.login-logo{margin:0 auto 20px}
.login-box h2{font-size:22px;font-weight:600;color:var(--dark);margin-bottom:6px}
.login-box p{font-size:13px;color:var(--gray);margin-bottom:28px}
.login-field{
  width:100%;padding:13px 16px;border:1px solid #d4dcc0;
  border-radius:8px;font-size:14px;margin-bottom:14px;
  font-family:inherit;outline:none;transition:border-color .2s;
}
.login-field:focus{border-color:var(--olive)}
.login-btn{
  width:100%;padding:14px;background:var(--dark);color:var(--cream);
  border:none;border-radius:8px;font-size:14px;font-weight:500;
  cursor:pointer;font-family:inherit;transition:background .2s;
}
.login-btn:hover{background:#2d3a2e}
.login-error{
  display:none;background:#fef2f2;border:1px solid #fecaca;
  border-radius:6px;padding:10px 14px;font-size:13px;color:var(--red);
  margin-bottom:14px;text-align:left;
}

/* ── APP SHELL ── */
#app{display:none;min-height:100vh;flex-direction:column}

/* Topbar */
.topbar{
  background:var(--dark);padding:14px 28px;
  display:flex;align-items:center;justify-content:space-between;
  position:sticky;top:0;z-index:100;
}
.topbar-left{display:flex;align-items:center;gap:14px}
.topbar-logo{
  width:36px;height:36px;border-radius:50%;
  border:1.5px solid var(--gold);
  display:flex;align-items:center;justify-content:center;
  font-family:Georgia,serif;font-size:12px;color:var(--gold);font-weight:600;
}
.topbar-title{
  font-size:15px;font-weight:600;color:var(--white);
}
.topbar-sub{font-size:11px;color:rgba(255,255,255,.35);margin-top:1px}
.topbar-right{display:flex;align-items:center;gap:12px}
.topbar-count{
  background:var(--olive-light);border:1px solid var(--border);
  border-radius:20px;padding:4px 12px;
  font-size:12px;font-weight:600;color:var(--olive);
}
.logout-btn{
  background:transparent;border:1px solid rgba(255,255,255,.15);
  color:rgba(255,255,255,.5);border-radius:6px;padding:6px 12px;
  font-size:12px;cursor:pointer;font-family:inherit;transition:all .2s;
}
.logout-btn:hover{border-color:rgba(255,255,255,.4);color:var(--white)}
.nav-link{
  background:transparent;border:1px solid rgba(255,255,255,.15);
  color:rgba(255,255,255,.5);border-radius:6px;padding:6px 12px;
  font-size:12px;cursor:pointer;font-family:inherit;text-decoration:none;
  transition:all .2s;display:inline-flex;align-items:center;gap:6px;
}
.nav-link:hover{border-color:rgba(255,255,255,.4);color:var(--white)}

/* Main content */
.main{flex:1;padding:28px;max-width:1200px;margin:0 auto;width:100%}

/* Toolbar */
.toolbar{
  display:flex;align-items:center;gap:12px;margin-bottom:24px;flex-wrap:wrap;
}
.search-wrap{position:relative;flex:1;min-width:220px}
.search-input{
  width:100%;padding:10px 16px 10px 38px;
  border:1px solid #d4dcc0;border-radius:8px;font-size:13px;
  font-family:inherit;outline:none;background:var(--white);
  transition:border-color .2s;
}
.search-input:focus{border-color:var(--olive)}
.search-icon{
  position:absolute;left:12px;top:50%;transform:translateY(-50%);
  color:var(--gray);font-size:15px;pointer-events:none;
}
.filter-select{
  padding:10px 14px;border:1px solid #d4dcc0;border-radius:8px;
  font-size:13px;font-family:inherit;background:var(--white);
  outline:none;cursor:pointer;color:var(--dark);
}
.refresh-btn{
  padding:10px 16px;background:var(--white);border:1px solid #d4dcc0;
  border-radius:8px;font-size:13px;cursor:pointer;font-family:inherit;
  color:var(--dark);transition:all .2s;display:flex;align-items:center;gap:6px;
}
.refresh-btn:hover{border-color:var(--olive);color:var(--olive)}
.refresh-btn.spinning svg{animation:spin .8s linear infinite}
@keyframes spin{to{transform:rotate(360deg)}}

/* Stats bar */
.stats{
  display:grid;grid-template-columns:repeat(4,1fr);gap:14px;margin-bottom:24px;
}
.stat-card{
  background:var(--white);border:1px solid rgba(0,0,0,.06);
  border-radius:10px;padding:16px 20px;
}
.stat-val{font-size:24px;font-weight:700;color:var(--dark);margin-bottom:2px}
.stat-label{font-size:11px;color:var(--gray);text-transform:uppercase;letter-spacing:.8px}
.stat-card.olive .stat-val{color:var(--olive)}
.stat-card.gold .stat-val{color:var(--gold)}

/* Plans grid */
.plans-grid{
  display:grid;
  grid-template-columns:repeat(auto-fill,minmax(320px,1fr));
  gap:16px;
}

/* Plan card */
.plan-card{
  background:var(--white);border:1px solid rgba(0,0,0,.07);
  border-radius:10px;overflow:hidden;transition:box-shadow .2s,transform .15s;
}
.plan-card:hover{box-shadow:0 4px 20px rgba(0,0,0,.1);transform:translateY(-2px)}

.card-header{
  background:var(--dark);padding:14px 18px;
  display:flex;align-items:flex-start;justify-content:space-between;gap:10px;
}
.card-nombre{
  font-size:15px;font-weight:600;color:var(--white);
  line-height:1.3;flex:1;
}
.card-modelo{
  font-size:10px;font-weight:700;padding:3px 8px;border-radius:12px;
  text-transform:uppercase;letter-spacing:.5px;flex-shrink:0;margin-top:2px;
}
.modelo-claude{background:rgba(255,255,255,.15);color:rgba(255,255,255,.8)}
.modelo-gemini{background:rgba(26,115,232,.3);color:#93c5fd}
.modelo-groq{background:rgba(245,80,54,.3);color:#fca5a5}
.modelo-otro{background:rgba(255,255,255,.1);color:rgba(255,255,255,.6)}

.card-body{padding:16px 18px}
.card-meta{display:flex;flex-direction:column;gap:5px;margin-bottom:14px}
.card-meta-row{
  display:flex;align-items:center;gap:7px;
  font-size:12px;color:var(--gray);
}
.card-meta-row svg{width:13px;height:13px;stroke:currentColor;fill:none;stroke-width:1.8;flex-shrink:0}
.card-meta-row strong{color:var(--dark)}

.card-actions{display:flex;gap:8px}
.btn-ver{
  flex:1;padding:9px;background:var(--olive-light);
  border:1px solid var(--border);border-radius:6px;
  font-size:12px;font-weight:600;color:var(--olive);
  text-decoration:none;text-align:center;
  transition:all .2s;cursor:pointer;
  display:flex;align-items:center;justify-content:center;gap:5px;
}
.btn-ver:hover{background:var(--olive);color:var(--white);border-color:var(--olive)}
.btn-editar{
  flex:1;padding:9px;background:#f0f9e8;
  border:1px solid rgba(143,168,50,.3);border-radius:6px;
  font-size:12px;font-weight:600;color:#5a7a1a;
  text-decoration:none;text-align:center;
  transition:all .2s;cursor:pointer;
  display:flex;align-items:center;justify-content:center;gap:5px;
}
.btn-editar:hover{background:#2d3a2e;color:var(--white);border-color:#2d3a2e}

/* Empty / loading */
.state-box{
  text-align:center;padding:80px 20px;color:var(--gray);
  grid-column:1/-1;
}
.state-box .icon{font-size:48px;margin-bottom:16px}
.state-box h3{font-size:18px;color:var(--dark);margin-bottom:8px}
.state-box p{font-size:13px;line-height:1.6;max-width:340px;margin:0 auto}

/* Skeleton */
.skeleton{
  background:linear-gradient(90deg,#f0f0f0 25%,#e0e0e0 50%,#f0f0f0 75%);
  background-size:200% 100%;
  animation:shimmer 1.5s infinite;
  border-radius:6px;
}
@keyframes shimmer{0%{background-position:200% 0}100%{background-position:-200% 0}}
.skel-card{
  background:var(--white);border:1px solid rgba(0,0,0,.07);
  border-radius:10px;overflow:hidden;
}
.skel-header{height:68px;background:#2d2d2d}
.skel-body{padding:16px 18px}
.skel-line{height:12px;margin-bottom:8px}
.skel-actions{display:flex;gap:8px;margin-top:14px}
.skel-btn{height:36px;flex:1;border-radius:6px}

/* Toast */
.toast{
  position:fixed;bottom:24px;right:24px;
  background:var(--dark);color:var(--white);
  padding:12px 20px;border-radius:8px;font-size:13px;
  box-shadow:0 4px 20px rgba(0,0,0,.2);
  transform:translateY(100px);opacity:0;
  transition:all .3s;z-index:999;
  display:flex;align-items:center;gap:8px;
}
.toast.show{transform:translateY(0);opacity:1}
.toast.success{border-left:3px solid var(--olive)}
.toast.error{border-left:3px solid var(--red)}
</style>
</head>
<body>

<!-- ══ LOGIN ══ -->
<div id="login-screen" style="display:none">
  <div class="login-box">
    <svg class="login-logo" width="48" height="48" viewBox="0 0 48 48" fill="none">
      <circle cx="24" cy="24" r="22" stroke="#b8935a" stroke-width="1.5"/>
      <text x="24" y="29" text-anchor="middle" fill="#b8935a" font-family="Georgia,serif" font-size="14" font-weight="600">CC</text>
    </svg>
    <h2>Centro Carvajal</h2>
    <p>Panel de administración · Planes generados</p>
    <div class="login-error" id="login-error">Contraseña incorrecta. Inténtalo de nuevo.</div>
    <input type="password" class="login-field" id="login-pass" placeholder="Contraseña" onkeydown="if(event.key==='Enter')doLogin()">
    <button class="login-btn" onclick="doLogin()">Ingresar</button>
    <button class="login-btn" onclick="mostrarRecuperar()" style="background:transparent;color:#8fa832;border:1px solid #8fa832;margin-top:8px">Olvidé mi contraseña</button>
    <div id="recuperar-wrap" style="display:none;margin-top:16px">
      <input type="email" class="login-field" id="recuperar-email" placeholder="Tu correo para recuperar acceso">
      <button class="login-btn" onclick="enviarRecuperacion()" style="background:#2d3a2e">Enviar link de recuperación</button>
      <div id="recuperar-msg" style="font-size:12px;color:#8fa832;margin-top:8px;display:none">✓ Si el correo existe, recibirás el link en unos minutos.</div>
    </div>
  </div>
</div>

<!-- ══ APP ══ -->
<div id="app" style="display:flex">

  <!-- Topbar -->
  <div class="topbar">
    <div class="topbar-left">
      <div class="topbar-logo">CC</div>
      <div>
        <div class="topbar-title">Planes Generados</div>
        <div class="topbar-sub">Centro Carvajal · Panel de Administración</div>
      </div>
    </div>
    <div class="topbar-right">
      <a class="nav-link" href="/catalogo">✎ Catálogo</a>
      <a class="nav-link" href="/panel">← Panel</a>
      <span id="user-email-display" style="font-size:11px;color:rgba(255,255,255,.4);margin-right:4px"></span>
      <div class="topbar-count" id="total-count">— planes</div>
      <button class="logout-btn" onclick="doLogout()">Cerrar sesión</button>
    </div>
  </div>

  <!-- Main -->
  <div class="main">

    <!-- Stats -->
    <div class="stats" id="stats-bar" style="display:none">
      <div class="stat-card">
        <div class="stat-val" id="stat-total">—</div>
        <div class="stat-label">Planes totales</div>
      </div>
      <div class="stat-card olive">
        <div class="stat-val" id="stat-mes">—</div>
        <div class="stat-label">Este mes</div>
      </div>
      <div class="stat-card gold">
        <div class="stat-val" id="stat-claude">—</div>
        <div class="stat-label">Con Claude</div>
      </div>
      <div class="stat-card">
        <div class="stat-val" id="stat-otros">—</div>
        <div class="stat-label">Gemini / Groq</div>
      </div>
    </div>

    <!-- Toolbar -->
    <div class="toolbar">
      <div class="search-wrap">
        <span class="search-icon">🔍</span>
        <input type="text" class="search-input" id="search-input" placeholder="Buscar por nombre de paciente..." oninput="filtrarPlanes()">
      </div>
      <select class="filter-select" id="filter-modelo" onchange="filtrarPlanes()">
        <option value="">Todos los modelos</option>
        <option value="claude">Claude</option>
        <option value="gemini">Gemini</option>
        <option value="groq">Groq</option>
      </select>
      <select class="filter-select" id="filter-orden" onchange="filtrarPlanes()">
        <option value="desc">Más recientes primero</option>
        <option value="asc">Más antiguos primero</option>
      </select>
      <button class="refresh-btn" id="refresh-btn" onclick="cargarPlanes()">
        <svg width="14" height="14" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><polyline points="23 4 23 10 17 10"/><path d="M20.49 15a9 9 0 1 1-2.12-9.36L23 10"/></svg>
        Actualizar
      </button>
    </div>

    <!-- Grid -->
    <div class="plans-grid" id="plans-grid">
      <!-- Skeletons mientras carga -->
      <div class="skel-card"><div class="skel-header"></div><div class="skel-body"><div class="skeleton skel-line" style="width:70%"></div><div class="skeleton skel-line" style="width:50%"></div><div class="skeleton skel-line" style="width:60%"></div><div class="skel-actions"><div class="skeleton skel-btn"></div><div class="skeleton skel-btn"></div></div></div></div>
      <div class="skel-card"><div class="skel-header"></div><div class="skel-body"><div class="skeleton skel-line" style="width:65%"></div><div class="skeleton skel-line" style="width:45%"></div><div class="skeleton skel-line" style="width:55%"></div><div class="skel-actions"><div class="skeleton skel-btn"></div><div class="skeleton skel-btn"></div></div></div></div>
      <div class="skel-card"><div class="skel-header"></div><div class="skel-body"><div class="skeleton skel-line" style="width:75%"></div><div class="skeleton skel-line" style="width:55%"></div><div class="skeleton skel-line" style="width:50%"></div><div class="skel-actions"><div class="skeleton skel-btn"></div><div class="skeleton skel-btn"></div></div></div></div>
    </div>
  </div>
</div>


  <!-- ── SECCIÓN CALENDARIO ── -->
  <div style="background:#f4f5ef;border-top:1px solid rgba(143,168,50,.2);padding:28px;margin-top:8px">
    <div style="max-width:1200px;margin:0 auto">
      <div style="margin-bottom:16px">
        <div style="font-size:11px;font-weight:700;text-transform:uppercase;letter-spacing:2px;color:var(--olive);margin-bottom:4px">Herramienta</div>
        <div style="font-size:18px;font-weight:700;color:var(--dark)">Generador de Calendario</div>
        <div style="font-size:12px;color:var(--gray);margin-top:3px">Genera un calendario de seguimiento listo para imprimir como PDF desde Chrome</div>
      </div>
      <div style="background:white;border:1px solid rgba(143,168,50,.2);border-radius:10px;padding:20px;display:flex;align-items:flex-end;gap:16px;flex-wrap:wrap">
        <div style="display:flex;flex-direction:column;gap:5px">
          <label style="font-size:11px;font-weight:600;text-transform:uppercase;letter-spacing:1px;color:var(--gray)">Mes de inicio</label>
          <input type="month" id="cal-desde" style="padding:9px 12px;border:1px solid #d4dcc0;border-radius:6px;font-size:13px;font-family:inherit;outline:none;color:var(--dark);min-width:160px">
        </div>
        <div style="display:flex;flex-direction:column;gap:5px">
          <label style="font-size:11px;font-weight:600;text-transform:uppercase;letter-spacing:1px;color:var(--gray)">Mes final</label>
          <input type="month" id="cal-hasta" style="padding:9px 12px;border:1px solid #d4dcc0;border-radius:6px;font-size:13px;font-family:inherit;outline:none;color:var(--dark);min-width:160px">
        </div>
        <button onclick="generarCalendario()" style="padding:10px 24px;background:var(--olive);color:white;border:none;border-radius:6px;font-size:13px;font-weight:600;cursor:pointer;font-family:inherit;display:flex;align-items:center;gap:8px;white-space:nowrap">
          Generar e Imprimir
        </button>
      </div>
      <div style="margin-top:10px;font-size:11px;color:var(--gray)">
        Al generar se abrira una nueva ventana lista para imprimir. Usa <strong>Ctrl+P</strong> &rarr; <strong>Guardar como PDF</strong> &rarr; Orientacion <strong>Horizontal</strong> &rarr; Sin margenes.
      </div>
    </div>
  </div>


<!-- Toast -->
<div class="toast" id="toast"></div>

</div>
<script>

// ── Auth — sesión por cookie httpOnly ───────────────────────
let currentUser = null;

async function initAuth() {
  try {
    const r = await fetch('/api/check-token');
    if (r.ok) {
      const d = await r.json();
      currentUser = d;
      // Verificar si hay redirect pendiente inyectado por el servidor
      const next = window._cv_next || new URLSearchParams(window.location.search).get('next');
      if (next && next.startsWith('/')) {
        window.location.href = next;
        return;
      }
      mostrarApp();
    } else {
      mostrarLogin();
    }
  } catch(e) {
    mostrarLogin();
  }
}

function getCookie(name) {
  const v = document.cookie.match('(^|;)\\s*' + name + '\\s*=\\s*([^;]+)');
  return v ? v.pop() : '';
}

function deleteCookie(name) {
  document.cookie = name + '=;expires=Thu, 01 Jan 1970 00:00:00 GMT;path=/';
}

function mostrarLogin() {
  document.getElementById('login-screen').style.display = 'flex';
  document.getElementById('app').style.display = 'none';
}

function mostrarApp() {
  document.getElementById('login-screen').style.display = 'none';
  document.getElementById('app').style.display = 'flex';
  if (currentUser) {
    const el = document.getElementById('user-email-display');
    if (el) el.textContent = currentUser.email;
    // Mostrar botón crear usuario solo al admin
    const btnCrear = document.getElementById('btn-crear-usuario');
    if (btnCrear && currentUser.rol === 'admin') btnCrear.style.display = 'inline-flex';
  }
  cargarPlanes();
}

async function doLogin() {
  const pass  = document.getElementById('login-pass').value;
  const errEl = document.getElementById('login-error');
  errEl.style.display = 'none';
  if (!pass) { errEl.style.display='block'; errEl.textContent='Ingresa la contraseña.'; return; }
  try {
    const r = await fetch('/api/login', {
      method: 'POST',
      headers: {'Content-Type': 'application/json'},
      body: JSON.stringify({password: pass})
    });
    const d = await r.json();
    if (d.ok) {
      currentUser = {email: d.email || 'admin@centrocarvajal.com', rol: d.rol};
      const next = window._cv_next || new URLSearchParams(window.location.search).get('next');
      if (next && next.startsWith('/')) { window.location.href = next; return; }
      mostrarApp();
    } else {
      errEl.style.display = 'block';
      errEl.textContent = d.error || 'Credenciales incorrectas.';
    }
  } catch(e) {
    errEl.style.display = 'block';
    errEl.textContent = 'Error de conexión. Intenta de nuevo.';
  }
}

async function doLogout() {
  await fetch('/api/logout', {method:'POST'});
  currentUser = null;
  mostrarLogin();
}

function mostrarRecuperar() {
  const w = document.getElementById('recuperar-wrap');
  w.style.display = w.style.display === 'none' ? 'block' : 'none';
}

async function enviarRecuperacion() {
  const email = document.getElementById('recuperar-email').value.trim();
  if (!email) return;
  await fetch('/api/recuperar-password', {
    method: 'POST',
    headers: {'Content-Type':'application/json'},
    body: JSON.stringify({email})
  });
  document.getElementById('recuperar-msg').style.display = 'block';
}

// ── Datos ────────────────────────────────────────────────────
let todosLosPlanes = [];

async function cargarPlanes() {
  const btn = document.getElementById('refresh-btn');
  btn.classList.add('spinning');

  try {
    const r    = await fetch('/api/planes');
    const data = await r.json();
    todosLosPlanes = data.planes || [];
    actualizarStats();
    filtrarPlanes();
    document.getElementById('stats-bar').style.display = 'grid';
  } catch(e) {
    mostrarToast('Error al cargar planes', 'error');
  } finally {
    btn.classList.remove('spinning');
  }
}

function actualizarStats() {
  const total  = todosLosPlanes.length;
  const ahora  = new Date();
  const mes    = todosLosPlanes.filter(p => {
    if (!p.fecha_raw) return false;
    const f = new Date(p.fecha_raw);
    return f.getMonth() === ahora.getMonth() && f.getFullYear() === ahora.getFullYear();
  }).length;
  const claude = todosLosPlanes.filter(p => p.modelo === 'claude').length;
  const otros  = total - claude;

  document.getElementById('stat-total').textContent  = total;
  document.getElementById('stat-mes').textContent     = mes;
  document.getElementById('stat-claude').textContent  = claude;
  document.getElementById('stat-otros').textContent   = otros;
  document.getElementById('total-count').textContent  = `${total} plan${total !== 1 ? 'es' : ''}`;
}

function filtrarPlanes() {
  const q      = document.getElementById('search-input').value.toLowerCase();
  const modelo = document.getElementById('filter-modelo').value;
  const orden  = document.getElementById('filter-orden').value;

  let lista = todosLosPlanes.filter(p => {
    const matchQ = !q || p.nombre.toLowerCase().includes(q);
    const matchM = !modelo || p.modelo === modelo;
    return matchQ && matchM;
  });

  if (orden === 'asc') lista = lista.slice().reverse();

  renderPlanes(lista);
}

function renderPlanes(lista) {
  const grid = document.getElementById('plans-grid');

  if (!lista.length) {
    grid.innerHTML = `<div class="state-box">
      <div class="icon">📋</div>
      <h3>No hay planes</h3>
      <p>No se encontraron planes con los filtros actuales.</p>
    </div>`;
    return;
  }

  grid.innerHTML = lista.map(p => {
    const modeloClass = {claude:'modelo-claude', gemini:'modelo-gemini', groq:'modelo-groq'}[p.modelo] || 'modelo-otro';
    const modeloLabel = (p.modelo || 'IA').charAt(0).toUpperCase() + (p.modelo || 'ia').slice(1);
    const borrador_url = `/borrador/${p.job_id}`;
    return `<div class="plan-card">
      <div class="card-header">
        <div class="card-nombre">${esc(p.nombre)}</div>
        <span class="card-modelo ${modeloClass}">${modeloLabel}</span>
      </div>
      <div class="card-body">
        <div class="card-meta">
          <div class="card-meta-row">
            <svg viewBox="0 0 24 24"><rect x="3" y="4" width="18" height="18" rx="2"/><line x1="16" y1="2" x2="16" y2="6"/><line x1="8" y1="2" x2="8" y2="6"/><line x1="3" y1="10" x2="21" y2="10"/></svg>
            <span>${esc(p.fecha)}</span>
          </div>
          <div class="card-meta-row">
            <svg viewBox="0 0 24 24"><path d="M14 2H6a2 2 0 0 0-2 2v16a2 2 0 0 0 2 2h12a2 2 0 0 0 2-2V8z"/><polyline points="14 2 14 8 20 8"/></svg>
            <span style="word-break:break-all;font-size:11px;color:#aaa">${esc(p.nombre_archivo)}</span>
          </div>
        </div>
        <div class="card-actions">
          <a href="${esc(p.url)}" target="_blank" class="btn-ver">
            <svg width="13" height="13" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><path d="M1 12s4-8 11-8 11 8 11 8-4 8-11 8-11-8-11-8z"/><circle cx="12" cy="12" r="3"/></svg>
            Ver plan
          </a>
          <a href="${esc(borrador_url)}" target="_blank" class="btn-editar">
            <svg width="13" height="13" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2"><path d="M11 4H4a2 2 0 0 0-2 2v14a2 2 0 0 0 2 2h14a2 2 0 0 0 2-2v-7"/><path d="M18.5 2.5a2.121 2.121 0 0 1 3 3L12 15l-4 1 1-4 9.5-9.5z"/></svg>
            Editar borrador
          </a>
        </div>
      </div>
    </div>`;
  }).join('');
}

function esc(s) {
  return String(s || '').replace(/&/g,'&amp;').replace(/</g,'&lt;').replace(/>/g,'&gt;').replace(/"/g,'&quot;');
}

function mostrarToast(msg, tipo='success') {
  const t = document.getElementById('toast');
  t.textContent = (tipo === 'success' ? '✓ ' : '✕ ') + msg;
  t.className = `toast ${tipo} show`;
  setTimeout(() => t.classList.remove('show'), 3000);
}


function generarCalendario() {
  var desde = document.getElementById('cal-desde').value;
  var hasta = document.getElementById('cal-hasta').value;
  if (!desde || !hasta) {
    mostrarToast('Selecciona mes de inicio y mes final', 'error');
    return;
  }
  var d0 = new Date(desde + '-01');
  var d1 = new Date(hasta + '-01');
  if (d0 > d1) {
    mostrarToast('El mes de inicio debe ser anterior al mes final', 'error');
    return;
  }

  var meses = ['Enero','Febrero','Marzo','Abril','Mayo','Junio',
               'Julio','Agosto','Septiembre','Octubre','Noviembre','Diciembre'];
  var diasSem = ['LU','MA','MI','JU','VI','SA','DO'];
  var colores = [
    {bg:'#c8e6c9',border:'#a5d6a7',label:'Nutrici\u00f3n'},
    {bg:'#bbdefb',border:'#90caf9',label:'Ejercicio'},
    {bg:'#f8bbd0',border:'#f48fb1',label:'Rutina facial'},
    {bg:'#fff9c4',border:'#ffe082',label:'Sue\u00f1o'}
  ];

  // Construir array con el HTML de cada mes
  var todosMeses = [];
  var cur = new Date(d0);
  while (cur <= d1) {
    var yr  = cur.getFullYear();
    var mo  = cur.getMonth();
    var diasEnMes = new Date(yr, mo + 1, 0).getDate();
    var primerDia = new Date(yr, mo, 1).getDay();
    primerDia = (primerDia === 0) ? 6 : primerDia - 1;

    var cabHtml = '';
    for (var di = 0; di < diasSem.length; di++) {
      cabHtml += '<div class="cal-dh">' + diasSem[di] + '</div>';
    }

    var celdasHtml = '';
    for (var e = 0; e < primerDia; e++) celdasHtml += '<div class="cal-day empty"></div>';
    for (var dia = 1; dia <= diasEnMes; dia++) {
      var dots = '';
      for (var ci = 0; ci < colores.length; ci++) {
        dots += '<div style="width:11px;height:11px;border-radius:2px;background:' +
                colores[ci].bg + ';border:1px solid ' + colores[ci].border +
                ';flex-shrink:0"></div>';
      }
      celdasHtml += '<div class="cal-day"><div class="cal-day-num">' + dia +
                    '</div><div class="cal-dots">' + dots + '</div></div>';
    }
    var total = primerDia + diasEnMes;
    var rest  = (7 - total % 7) % 7;
    for (var r = 0; r < rest; r++) celdasHtml += '<div class="cal-day empty"></div>';

    todosMeses.push('<div class="cal-month">' +
      '<div class="cal-mhdr">' + meses[mo] + ' ' + yr + '</div>' +
      '<div class="cal-days-hdr">' + cabHtml + '</div>' +
      '<div class="cal-days-grid">' + celdasHtml + '</div>' +
    '</div>');

    cur.setMonth(cur.getMonth() + 1);
  }

  // Leyenda compartida
  var leyenda = '';
  for (var li = 0; li < colores.length; li++) {
    leyenda += '<div style="display:flex;align-items:center;gap:5px;font-size:11px;color:#6b7280">' +
               '<div style="width:12px;height:12px;border-radius:2px;background:' + colores[li].bg +
               ';border:1px solid ' + colores[li].border + '"></div>' + colores[li].label + '</div>';
  }

  // Agrupar en páginas de 6 meses (3 cols x 2 filas) — cada una con su propio encabezado
  var MESES_POR_PAG = 6;
  var paginasHtml = '';
  for (var pi = 0; pi < todosMeses.length; pi += MESES_POR_PAG) {
    var grupo = todosMeses.slice(pi, pi + MESES_POR_PAG);
    // Rellenar con celdas vacías si el grupo no completa 6
    while (grupo.length < MESES_POR_PAG && grupo.length % 3 !== 0) {
      grupo.push('<div class="cal-month cal-month-empty"></div>');
    }
    var esUltima = (pi + MESES_POR_PAG >= todosMeses.length);
    paginasHtml +=
      '<div class="cal-page' + (esUltima ? '' : ' page-break') + '">' +
      '<div class="page-header">' +
      '<h1>Calendario de Seguimiento</h1>' +
      '<p>Centro Carvajal \u00b7 M\u00e9todo de Rejuvenecimiento Carvajal \u00b7 Marca cada d\u00eda al completar tu rutina</p>' +
      '</div>' +
      '<div class="cal-grid">' + grupo.join('') + '</div>' +
      '<div class="leyenda">' + leyenda + '</div>' +
      '<div class="footer">Centro Carvajal \u00b7 L\u00edderes en Medicina Est\u00e9tica en Panam\u00e1 \u00b7 centrocarvajal.com</div>' +
      '</div>';
  }

  var printHtml =
    '<!DOCTYPE html><html lang="es"><head><meta charset="UTF-8">' +
    '<title>Calendario de Seguimiento \u2014 Centro Carvajal</title>' +
    '<style>' +
    '@page{size:A4 landscape;margin:8mm}' +
    '*{margin:0;padding:0;box-sizing:border-box}' +
    'body{font-family:"Segoe UI",system-ui,sans-serif;background:#faf9f6;color:#1c1c1c}' +
    '.cal-page{width:100%;height:100vh;display:flex;flex-direction:column;justify-content:space-between}' +
    '.page-break{page-break-after:always}' +
    '.page-header{text-align:center;margin-bottom:10px;flex-shrink:0}' +
    '.page-header h1{font-family:Georgia,serif;font-size:17pt;color:#1c1c1c;margin-bottom:2px}' +
    '.page-header p{font-size:8.5pt;color:#6b7280}' +
    '.cal-grid{display:grid;grid-template-columns:repeat(3,1fr);gap:8px;flex:1;min-height:0}' +
    '.cal-month{border:1px solid rgba(143,168,50,.22);border-radius:8px;overflow:hidden;background:#fff;page-break-inside:avoid;display:flex;flex-direction:column}' +
    '.cal-month-empty{border:none;background:transparent}' +
    '.cal-mhdr{background:#1c1c1c;color:#8fa832;text-align:center;padding:6px 10px;font-family:Georgia,serif;font-size:10.5pt;font-weight:700;flex-shrink:0}' +
    '.cal-days-hdr{display:grid;grid-template-columns:repeat(7,1fr);background:rgba(143,168,50,.1);flex-shrink:0}' +
    '.cal-dh{text-align:center;font-size:6.5pt;font-weight:700;color:#8fa832;padding:3px 1px;text-transform:uppercase}' +
    '.cal-days-grid{display:grid;grid-template-columns:repeat(7,1fr);flex:1}' +
    '.cal-day{border-right:1px solid rgba(143,168,50,.1);border-bottom:1px solid rgba(143,168,50,.1);padding:2px 3px;background:#fff;min-height:30px}' +
    '.cal-day:nth-child(7n){border-right:none}' +
    '.cal-day.empty{background:#fafafa}' +
    '.cal-day-num{font-size:7pt;font-weight:700;color:#1c1c1c;margin-bottom:2px}' +
    '.cal-dots{display:flex;gap:2px;flex-wrap:wrap}' +
    '.leyenda{display:flex;gap:16px;justify-content:center;margin-top:8px;padding-top:6px;border-top:1px solid rgba(143,168,50,.2);flex-shrink:0}' +
    '.footer{text-align:center;margin-top:6px;font-size:7.5pt;color:#9aaa8a;flex-shrink:0}' +
    '@media print{body{background:white;-webkit-print-color-adjust:exact;print-color-adjust:exact}}' +
    '</style></head><body>' +
    paginasHtml +
    '<script>window.onload=function(){window.print();}</scr' + 'ipt>' +
    '</body></html>';

  var win = window.open('', '_blank');
  win.document.write(printHtml);
  win.document.close();
}

// Iniciar auth al cargar — verifica sesión y redirige si no está logueado
initAuth();
</script>
</body>
</html>""" 


# ════════════════════════════════════════════════════════════
# AUTH — usuarios, sesiones, recuperación de contraseña
# ════════════════════════════════════════════════════════════

USUARIOS_CLOUDINARY_ID = 'carvajal/config/usuarios'

def _cargar_usuarios():
    """Carga usuarios desde Cloudinary. Devuelve dict {email: {hash, rol}}."""
    if not CLOUDINARY_CLOUD_NAME:
        return {}
    try:
        # Cache-buster para evitar que el CDN sirva version antigua tras actualizar
        url = f'https://res.cloudinary.com/{CLOUDINARY_CLOUD_NAME}/raw/upload/{USUARIOS_CLOUDINARY_ID}.json?_={int(_time.time())}'
        r = req.get(url, timeout=10)
        if r.status_code == 200:
            return r.json()
    except Exception as e:
        print(f'[auth] Error cargando usuarios: {e}')
    return {}


def _guardar_usuarios(usuarios):
    """Guarda dict de usuarios en Cloudinary."""
    if not CLOUDINARY_CLOUD_NAME:
        return
    try:
        import tempfile
        with tempfile.NamedTemporaryFile(mode='w', suffix='.json', delete=False, encoding='utf-8') as tmp:
            json.dump(usuarios, tmp, ensure_ascii=False)
            tmp_path = tmp.name
        cloudinary.uploader.upload(
            tmp_path,
            folder='carvajal/config',
            public_id='usuarios',
            resource_type='raw',
            overwrite=True,
        )
        os.unlink(tmp_path)
        print('[auth] Usuarios guardados en Cloudinary')
    except Exception as e:
        print(f'[auth] Error guardando usuarios: {e}')


def _inicializar_usuarios():
    """Crea los usuarios iniciales si no existen."""
    usuarios = _cargar_usuarios()
    changed = False
    defaults = [
        ('centrocarvajal1@gmail.com', 'Carvajal2026!', 'staff'),
        ('isai.josue@gmail.com',      'Carvajal2026!', 'admin'),
    ]
    for email, pwd, rol in defaults:
        if email not in usuarios:
            hashed = bcrypt.hashpw(pwd.encode(), bcrypt.gensalt()).decode()
            usuarios[email] = {'hash': hashed, 'rol': rol}
            changed = True
            print(f'[auth] Usuario creado: {email} ({rol})')
    if changed:
        _guardar_usuarios(usuarios)
    return usuarios


def _limpiar_sesiones():
    """Elimina sesiones expiradas del dict en memoria."""
    now = _time.time()
    expiradas = [t for t, s in auth_sessions.items() if s['expires'] < now]
    for t in expiradas:
        del auth_sessions[t]


def _verificar_sesion(token):
    """Verifica si un token de sesión es válido y no expiró."""
    _limpiar_sesiones()
    sesion = auth_sessions.get(token)
    if not sesion:
        return None
    if sesion['expires'] < _time.time():
        del auth_sessions[token]
        return None
    return sesion


def _check_token():
    """Helper para endpoints protegidos — verifica cookie de sesión o header X-Token."""
    tok = request.cookies.get('cv_session', '') or request.headers.get('X-Token', '').strip()
    return _verificar_sesion(tok) is not None


def _get_sesion_from_cookie():
    """Lee token de cookie de sesión."""
    from flask import request as req_ctx
    token = req_ctx.cookies.get('cv_session', '')
    return _verificar_sesion(token)


@app.route('/api/login', methods=['POST'])
def api_login():
    data = request.get_json(silent=True) or {}
    email = data.get('email', '').strip().lower()
    pwd   = data.get('password', '').strip()
    # Login simple con solo contraseña para panel/planes: usa el primer usuario coincidente
    if not pwd:
        return jsonify({'ok': False, 'error': 'Contraseña requerida'}), 400
    usuarios = _cargar_usuarios()
    if not usuarios:
        usuarios = _inicializar_usuarios()
    admin_pwd = os.environ.get('ADMIN_PASSWORD', '').strip()
    if email:
        user = usuarios.get(email)
    elif admin_pwd and pwd == admin_pwd:
        # Login universal con ADMIN_PASSWORD para panel/planes/catalogo
        user = {'rol': 'admin'}
        email = 'admin@centrocarvajal.com'
    else:
        # Buscar primera coincidencia por contraseña
        user = None
        for u_email, u_data in usuarios.items():
            try:
                if bcrypt.checkpw(pwd.encode(), u_data.get('hash','').encode()):
                    user = u_data
                    email = u_email
                    break
            except Exception:
                continue
    if not user:
        return jsonify({'ok': False, 'error': 'Credenciales incorrectas'}), 401
    tok = secrets.token_hex(32)
    auth_sessions[tok] = {
        'email':   email,
        'rol':     user.get('rol', 'staff'),
        'expires': _time.time() + SESSION_DURATION,
    }
    from flask import make_response
    resp = make_response(jsonify({'ok': True, 'rol': user.get('rol', 'staff'), 'token': tok, 'email': email}))
    resp.set_cookie('cv_session', tok, max_age=SESSION_DURATION, httponly=True, samesite='Lax')
    print(f'[auth] Login: {email}')
    return resp


@app.route('/api/logout', methods=['POST'])
def api_logout():
    from flask import make_response
    tok = request.cookies.get('cv_session', '')
    if tok and tok in auth_sessions:
        del auth_sessions[tok]
    resp = make_response(jsonify({'ok': True}))
    resp.delete_cookie('cv_session')
    return resp


@app.route('/api/check-token', methods=['GET'])
def api_check_token():
    tok = request.cookies.get('cv_session', '') or request.headers.get('X-Token', '')
    sesion = _verificar_sesion(tok)
    if sesion:
        return jsonify({'ok': True, 'email': sesion['email'], 'rol': sesion['rol']})
    return jsonify({'ok': False}), 401


@app.route('/api/cambiar-password', methods=['POST'])
def api_cambiar_password():
    tok = request.cookies.get('cv_session', '')
    sesion = _verificar_sesion(tok)
    if not sesion:
        return jsonify({'ok': False, 'error': 'No autenticado'}), 401
    data = request.get_json(silent=True) or {}
    pwd_actual = data.get('actual', '').strip()
    pwd_nueva  = data.get('nueva', '').strip()
    if len(pwd_nueva) < 8:
        return jsonify({'ok': False, 'error': 'La contraseña debe tener al menos 8 caracteres'}), 400
    usuarios = _cargar_usuarios()
    user = usuarios.get(sesion['email'])
    if not user or not bcrypt.checkpw(pwd_actual.encode(), user['hash'].encode()):
        return jsonify({'ok': False, 'error': 'Contraseña actual incorrecta'}), 401
    usuarios[sesion['email']]['hash'] = bcrypt.hashpw(pwd_nueva.encode(), bcrypt.gensalt()).decode()
    _guardar_usuarios(usuarios)
    return jsonify({'ok': True})


@app.route('/api/recuperar-password', methods=['POST'])
def api_recuperar_password():
    data = request.get_json(silent=True) or {}
    email = data.get('email', '').strip().lower()
    usuarios = _cargar_usuarios()
    if email not in usuarios:
        # No revelar si el email existe o no
        return jsonify({'ok': True})
    tok = secrets.token_urlsafe(32)
    reset_tokens[tok] = {'email': email, 'expires': _time.time() + RESET_DURATION}
    base_url = os.environ.get('BASE_URL', 'https://metodo.centrocarvajal.com')
    link = f'{base_url}/reset-password?token={tok}'
    cuerpo = (
        f'<!DOCTYPE html><html><head><meta charset="UTF-8"></head>'
        f'<body style="background:#f0e8de;padding:20px;font-family:sans-serif">'
        f'<div style="max-width:500px;margin:0 auto;background:#fff;border:1px solid #ddd">'
        f'<div style="background:#1a1410;padding:20px 24px">'
        f'<div style="color:#b8935a;font-size:11px;letter-spacing:3px;text-transform:uppercase">Centro Carvajal · Recuperación de acceso</div>'
        f'</div>'
        f'<div style="padding:24px">'
        f'<p style="font-size:14px;color:#2d2020;margin-bottom:20px">Recibimos una solicitud para restablecer tu contraseña.</p>'
        f'<div style="text-align:center;margin:24px 0">'
        f'<a href="{link}" style="background:#1a1410;color:#b8935a;padding:13px 28px;border-radius:4px;text-decoration:none;font-size:14px;font-weight:500">Restablecer contraseña</a>'
        f'</div>'
        f'<p style="font-size:12px;color:#999;text-align:center">Este link expira en 1 hora. Si no solicitaste esto, ignora este correo.</p>'
        f'</div>'
        f'<div style="background:#1a1410;padding:12px 24px;text-align:center;font-size:10px;color:rgba(255,255,255,0.3)">Centro Carvajal · metodo.centrocarvajal.com</div>'
        f'</div></body></html>'
    )
    try:
        enviar_resend(f'Recuperación de contraseña — Centro Carvajal', cuerpo, email)
    except Exception as e:
        print(f'[auth] Error enviando reset email: {e}')
    return jsonify({'ok': True})


@app.route('/api/reset-password', methods=['POST'])
def api_reset_password():
    data = request.get_json(silent=True) or {}
    tok  = data.get('token', '').strip()
    pwd  = data.get('password', '').strip()
    if len(pwd) < 8:
        return jsonify({'ok': False, 'error': 'La contraseña debe tener al menos 8 caracteres'}), 400
    info = reset_tokens.get(tok)
    if not info or info['expires'] < _time.time():
        return jsonify({'ok': False, 'error': 'Link inválido o expirado'}), 400
    usuarios = _cargar_usuarios()
    email = info['email']
    if email not in usuarios:
        return jsonify({'ok': False, 'error': 'Usuario no encontrado'}), 404
    usuarios[email]['hash'] = bcrypt.hashpw(pwd.encode(), bcrypt.gensalt()).decode()
    _guardar_usuarios(usuarios)
    del reset_tokens[tok]
    print(f'[auth] Password reseteado: {email}')
    return jsonify({'ok': True})


@app.route('/reset-password', methods=['GET'])
def reset_password_page():
    """Página para restablecer contraseña desde el link del correo."""
    tok = request.args.get('token', '')
    info = reset_tokens.get(tok)
    valid = bool(info and info['expires'] >= _time.time())
    return f'''<!DOCTYPE html><html lang="es"><head><meta charset="UTF-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>Restablecer contraseña · Centro Carvajal</title>
<style>
*{{margin:0;padding:0;box-sizing:border-box}}
body{{font-family:'Segoe UI',sans-serif;background:#f4f5ef;min-height:100vh;display:flex;align-items:center;justify-content:center;padding:20px}}
.box{{background:#fff;border-radius:12px;padding:44px 40px;max-width:400px;width:100%;box-shadow:0 8px 40px rgba(0,0,0,.1);text-align:center}}
h2{{font-size:22px;color:#1a1410;margin-bottom:8px}}
p{{font-size:13px;color:#6b7280;margin-bottom:28px}}
input{{width:100%;padding:13px 16px;border:1px solid #d4dcc0;border-radius:8px;font-size:14px;margin-bottom:14px;font-family:inherit;outline:none}}
input:focus{{border-color:#8fa832}}
button{{width:100%;padding:14px;background:#1a1410;color:#b8935a;border:none;border-radius:8px;font-size:14px;font-weight:500;cursor:pointer;font-family:inherit}}
.msg{{margin-top:14px;font-size:13px;padding:10px;border-radius:6px;display:none}}
.msg.ok{{background:#f0f7e6;color:#3a6020;border:1px solid #c5d9a0}}
.msg.err{{background:#fef2f2;color:#c0392b;border:1px solid #fecaca}}
</style></head><body>
<div class="box">
{"<h2>Link inválido</h2><p>Este link de recuperación expiró o ya fue usado. Solicita uno nuevo desde el panel.</p>" if not valid else f'''
<h2>Nueva contraseña</h2>
<p>Ingresa tu nueva contraseña para continuar.</p>
<input type="password" id="p1" placeholder="Nueva contraseña (mín. 8 caracteres)">
<input type="password" id="p2" placeholder="Confirmar contraseña">
<button onclick="doReset()">Restablecer contraseña</button>
<div class="msg" id="msg"></div>
<script>
function doReset(){{
  const p1=document.getElementById("p1").value;
  const p2=document.getElementById("p2").value;
  const msg=document.getElementById("msg");
  if(p1.length<8){{msg.className="msg err";msg.style.display="block";msg.textContent="Mínimo 8 caracteres";return}}
  if(p1!==p2){{msg.className="msg err";msg.style.display="block";msg.textContent="Las contraseñas no coinciden";return}}
  fetch("/api/reset-password",{{method:"POST",headers:{{"Content-Type":"application/json"}},body:JSON.stringify({{token:"{tok}",password:p1}})}})
  .then(r=>r.json()).then(d=>{{
    if(d.ok){{msg.className="msg ok";msg.style.display="block";msg.textContent="Contraseña actualizada. Ya puedes cerrar esta ventana.";}}
    else{{msg.className="msg err";msg.style.display="block";msg.textContent=d.error||"Error al restablecer";}}
  }});
}}
</script>
'''}
</div></body></html>'''


@app.route('/api/crear-usuario', methods=['POST'])
def api_crear_usuario():
    """Solo el admin puede crear usuarios."""
    tok = request.cookies.get('cv_session', '')
    sesion = _verificar_sesion(tok)
    if not sesion or sesion.get('rol') != 'admin':
        return jsonify({'ok': False, 'error': 'Solo el administrador puede crear usuarios'}), 403
    data = request.get_json(silent=True) or {}
    email = data.get('email', '').strip().lower()
    pwd   = data.get('password', '').strip()
    rol   = data.get('rol', 'staff')
    if not email or not pwd:
        return jsonify({'ok': False, 'error': 'Email y contraseña requeridos'}), 400
    if len(pwd) < 8:
        return jsonify({'ok': False, 'error': 'Contraseña mínimo 8 caracteres'}), 400
    if rol not in ('admin', 'staff'):
        rol = 'staff'
    usuarios = _cargar_usuarios()
    if email in usuarios:
        return jsonify({'ok': False, 'error': 'El usuario ya existe'}), 409
    usuarios[email] = {
        'hash': bcrypt.hashpw(pwd.encode(), bcrypt.gensalt()).decode(),
        'rol': rol,
    }
    _guardar_usuarios(usuarios)
    print(f'[auth] Usuario creado por admin: {email} ({rol})')
    return jsonify({'ok': True})


@app.route('/api/planes', methods=['GET'])
def api_listar_planes():
    if not _check_token():
        return jsonify({'error': 'No autenticado'}), 401
    planes = []
    try:
        import cloudinary.api
        from datetime import datetime as dt
        resultado = cloudinary.api.resources(
            type='upload',
            resource_type='raw',
            prefix='carvajal/planes/',
            max_results=200,
            context=True,
        )
        for r in resultado.get('resources', []):
            public_id      = r.get('public_id', '')
            # secure_url ya incluye la extensión correcta
            plan_url       = r.get('secure_url', '')
            nombre_archivo = public_id.split('/')[-1] + '.html'

            # Extraer campos del nombre de archivo
            # Formato: Plan_Nombre_Apellido_YYYYMMDD_modelo(.html)
            base  = public_id.split('/')[-1]
            base  = re.sub(r'\.html$', '', base, flags=re.IGNORECASE)  # quitar .html si viene
            parts = base.split('_')
            modelo = 'claude'
            fecha_raw = ''
            nombre_parts = []
            for p in parts:
                p_clean = re.sub(r'\.html$', '', p, flags=re.IGNORECASE)  # limpiar .html en cada parte
                if p_clean.lower() in ('claude', 'gemini', 'groq'):
                    modelo = p_clean.lower()
                elif len(p_clean) == 8 and p_clean.isdigit():
                    fecha_raw = p_clean
                elif p_clean not in ('Plan',) and p_clean:
                    nombre_parts.append(p_clean)
            nombre_paciente = ' '.join(nombre_parts) if nombre_parts else base

            # Fecha legible desde nombre de archivo
            fecha_legible = ''
            if fecha_raw and len(fecha_raw) == 8:
                try:
                    fecha_legible = dt.strptime(fecha_raw, '%Y%m%d').strftime('%d/%m/%Y')
                except:
                    fecha_legible = fecha_raw

            # Fecha ISO de Cloudinary para stats "este mes"
            created_at = r.get('created_at', '')

            # job_id: primero desde context (si fue guardado), luego fallback al nombre
            ctx      = r.get('context', {}).get('custom', {})
            job_id_r = ctx.get('job_id', '') or base

            planes.append({
                'url'           : plan_url,
                'nombre_archivo': nombre_archivo,
                'nombre'        : nombre_paciente,
                'modelo'        : modelo,
                'fecha'         : fecha_legible or 'Sin fecha',
                'fecha_raw'     : fecha_raw,
                'job_id'        : job_id_r,
                'created_at'    : created_at,
            })
        planes.sort(key=lambda x: x.get('created_at', ''), reverse=True)
    except Exception as e:
        print(f'[api/planes] Error: {e}')
        return jsonify({'planes': [], 'error': str(e)})
    return jsonify({'planes': planes, 'total': len(planes)})


@app.route('/catalogo', methods=['GET'])
def admin_catalogo():
    return send_from_directory(os.path.dirname(os.path.abspath(__file__)), 'catalogo.html')


@app.route('/api/catalogo', methods=['GET', 'POST'])
def api_catalogo():
    if not _check_token():
        return jsonify({'error': 'No autenticado'}), 401
    if request.method == 'GET':
        return jsonify(_cargar_catalogo())
    data = request.get_json(silent=True)
    if not data or (not isinstance(data, dict) and not isinstance(data, list)):
        return jsonify({'error': 'Body inválido: se esperaba objeto o lista de tratamientos'}), 400
    try:
        _guardar_catalogo(data)
        total = len(data.get('tratamientos', [])) if isinstance(data, dict) else len(data)
        return jsonify({'ok': True, 'count': total})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


def generar_analisis_medico(data):
    """Llama a Claude para generar un análisis clínico técnico orientado al médico."""
    print("[analisis_medico] Generando análisis clínico con Claude...")
    t0 = time.time()

    # Construir resumen de datos clínicos relevantes
    campos = {
        'Nombre': data.get('nombre',''),
        'Edad': data.get('edad',''),
        'Sexo': data.get('sexo',''),
        'Fecha nacimiento': data.get('fechaNacimiento',''),
        'IMC': data.get('imc',''),
        'Estatura': data.get('estatura','') + ' cm',
        'Peso': data.get('peso','') + ' kg',
        'Condicion sistemica': data.get('condicionSistemica',''),
        'Otras condiciones': data.get('condiciones',''),
        'Medicamentos': data.get('medicamentos',''),
        'Cirugias': data.get('cirugias',''),
        'Fuma': data.get('fuma',''),
        'Alcohol': data.get('alcohol',''),
        'Embarazo': data.get('embarazo',''),
        'Lactancia': data.get('lactancia',''),
        'Anticonceptivos': data.get('anticonceptivos',''),
        'SOP': data.get('sop',''),
        'Menopausia': data.get('menopausia',''),
        'Perimenopausia': data.get('perimenopausia',''),
        'Sueno': data.get('sueno',''),
        'Hora duerme': data.get('horaDuerme',''),
        'Hora despierta': data.get('horaDespierta',''),
        'Cansancio diurno': data.get('cansancioDia',''),
        'Nivel estres': data.get('nivelEstres',''),
        'Actividad fisica': data.get('actFisica',''),
        'Intolerancias': data.get('intolerancias',''),
        'Sintomas digestivos': data.get('sintomasDigestivos',''),
        'Tipo piel': data.get('pielTipo',''),
        'Problemas piel': data.get('pielProblemas',''),
        'Exposicion solar': data.get('solar',''),
        'SPF': data.get('spf',''),
        'Historial estetico': data.get('historialEstetico',''),
        'Contraindicaciones': data.get('contraindications',''),
        'Antecedentes familiares': (data.get('antecedentesFam') or '') + ' ' + (data.get('antecedentesFamDet') or ''),
        'Satisfaccion actual': str(data.get('satisfaccion','')) + '/10',
        'Prioridad': data.get('prioridad',''),
        'Areas faciales': data.get('areasFaciales',''),
        'Areas corporales': data.get('areasCorporales',''),
    }
    def _str(v):
        if isinstance(v, list):  return ', '.join(str(i) for i in v) if v else ''
        if isinstance(v, dict):  return '; '.join(f'{k}: {val}' for k, val in v.items()) if v else ''
        return str(v) if v else ''

    resumen = '\n'.join(
        f'{k}: {_str(v)}'
        for k, v in campos.items()
        if _str(v).strip() not in ('', '0', 'No registrado', ' ')
    )

    system_prompt = """Eres un médico especialista en medicina estética y bienestar integral.
Recibes los datos del cuestionario de un paciente nuevo de Centro Carvajal, clínica de medicina estética en Panamá.
Tu tarea es generar un ANÁLISIS CLÍNICO TÉCNICO para el expediente médico interno — NO para el paciente.
El análisis debe ser en lenguaje médico técnico, objetivo y clínicamente relevante.

Estructura tu respuesta en estas secciones exactas, usando encabezados con ##:

## RESUMEN DEL PERFIL CLÍNICO
2-3 oraciones que resuman el cuadro general del paciente desde perspectiva clínica.

## HALLAZGOS RELEVANTES
Lista de hallazgos clínicamente significativos (metabólicos, hormonales, dermatológicos, nutricionales, etc.).

## POSIBLES CONDICIONES A EVALUAR
Si alguna combinación de síntomas o factores sugiere una condición subyacente, mencionarla con justificación clínica breve. Ser claro que son hipótesis diagnósticas que requieren evaluación, no diagnósticos definitivos.

## CONTRAINDICACIONES Y PRECAUCIONES PARA TRATAMIENTOS ESTÉTICOS
Lista específica de contraindicaciones absolutas y relativas según el perfil del paciente.

## RECOMENDACIONES PARA EL EQUIPO MÉDICO
Acciones sugeridas antes de iniciar tratamientos (exámenes, consultas con especialistas, ajustes de protocolo, etc.).

Sé conciso pero completo. Usa terminología médica adecuada."""

    try:
        resp = req.post(
            'https://api.anthropic.com/v1/messages',
            headers={
                'Content-Type': 'application/json',
                'x-api-key': CLAUDE_KEY,
                'anthropic-version': '2023-06-01',
            },
            json={
                'model': 'claude-opus-4-6',
                'max_tokens': 2000,
                'system': system_prompt,
                'messages': [{'role': 'user', 'content': f'Datos del paciente:\n{resumen}'}],
            },
            timeout=120
        )
        elapsed = round(time.time() - t0, 1)
        if resp.status_code != 200:
            error_text = resp.text
            print(f'[analisis_medico] ERROR {resp.status_code}: {error_text[:200]}')
            if resp.status_code == 400 and 'credit balance is too low' in error_text:
                print('[analisis_medico] Sin créditos Claude — fallback a Groq')
                if error_text not in _claude_sin_creditos:
                    _claude_sin_creditos.append(error_text[:500])
                try:
                    rg = req.post(
                        'https://api.groq.com/openai/v1/chat/completions',
                        headers={'Authorization': f'Bearer {GROQ_KEY}', 'Content-Type': 'application/json'},
                        json={
                            'model': 'llama-3.3-70b-versatile',
                            'messages': [
                                {'role': 'system', 'content': system_prompt},
                                {'role': 'user',   'content': f'Datos del paciente:\n{resumen}'}
                            ],
                            'max_tokens': 2000,
                            'temperature': 0.4,
                        },
                        timeout=120
                    )
                    if rg.status_code == 200:
                        txt_g = rg.json()['choices'][0]['message']['content'].strip()
                        print(f'[analisis_medico] Groq fallback OK — {len(txt_g)} chars')
                        return txt_g
                    print(f'[analisis_medico] Groq fallback error {rg.status_code}')
                except Exception as eg:
                    print(f'[analisis_medico] Groq fallback excepción: {eg}')
            return None
        txt = resp.json()['content'][0]['text'].strip()
        print(f'[analisis_medico] OK — {elapsed}s ({len(txt)} chars)')
        return txt
    except Exception as e:
        import traceback
        print(f'[analisis_medico] Excepcion: {e}')
        print(traceback.format_exc())
        return None


def generar_docx_cuestionario(data, plan_json=None, analisis_medico=None):
    """Genera .docx con análisis clínico + tabla plana de campos del formulario."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import re as _re

    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(2);   section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5); section.right_margin  = Cm(2.5)

    VERDE  = RGBColor(0x2d, 0x3a, 0x2e)
    OLIVE  = RGBColor(0x8f, 0xa8, 0x32)
    GOLD   = RGBColor(0xb8, 0x93, 0x5a)
    BLANCO = RGBColor(0xff, 0xff, 0xff)
    GRIS   = RGBColor(0x6b, 0x72, 0x80)

    def set_cell_bg(cell, color_hex):
        tc = cell._tc; tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), color_hex); tcPr.append(shd)

    # ── Limpiador completo de markdown ──────────────────────────
    def limpiar_markdown(texto):
        lineas_limpias = []
        for linea in texto.split('\n'):
            l = linea.strip()
            # Saltar líneas separadoras de tabla: |---|---| o :---:|:---:
            if _re.match(r'^\|[\s\-\|:]+\|?\s*$', l):
                continue
            # Saltar líneas que son solo guiones/iguales (separadores)
            if _re.match(r'^[-=]{3,}\s*$', l):
                continue
            # Filas de tabla markdown: |col1|col2| → extraer contenido
            if '|' in l:
                celdas = [c.strip() for c in l.split('|') if c.strip()]
                if celdas:
                    l = '  '.join(celdas)
                else:
                    continue
            # Quitar encabezados # ## ### (al inicio o en cualquier posición)
            l = _re.sub(r'#{1,6}\s*', '', l)
            # Quitar **negrita** y __negrita__
            l = _re.sub(r'\*\*(.+?)\*\*', r'\1', l)
            l = _re.sub(r'__(.+?)__',     r'\1', l)
            # Quitar *cursiva* y _cursiva_
            l = _re.sub(r'\*(.+?)\*', r'\1', l)
            l = _re.sub(r'_(.+?)_',   r'\1', l)
            # Quitar asteriscos y guiones bajos sueltos restantes
            l = l.replace('**', '').replace('__', '').replace('*', '')
            # Quitar viñetas markdown al inicio (-, •, >)
            l = _re.sub(r'^[-•>]\s+', '', l.strip())
            # Limpiar espacios múltiples
            l = _re.sub(r'  +', ' ', l).strip()
            if l:
                lineas_limpias.append(l)
        return '\n'.join(lineas_limpias)

    # ── Encabezado ───────────────────────────────────────────────
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('CENTRO CARVAJAL')
    r.bold = True; r.font.size = Pt(16); r.font.color.rgb = VERDE
    p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run('Cuestionario del Paciente — Método de Rejuvenecimiento Carvajal')
    r2.font.size = Pt(10); r2.font.color.rgb = GOLD
    doc.add_paragraph()

    # ── Análisis clínico (sección IA) ────────────────────────────
    if analisis_medico:
        p_title = doc.add_paragraph()
        r_title = p_title.add_run('ANÁLISIS CLÍNICO — USO INTERNO DEL EQUIPO MÉDICO')
        r_title.bold = True; r_title.font.size = Pt(11); r_title.font.color.rgb = BLANCO
        p_title.paragraph_format.space_before = Pt(0)
        p_title.paragraph_format.space_after  = Pt(0)
        pPr = p_title._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1c1c1c'); pPr.append(shd)
        p_title.paragraph_format.left_indent  = Cm(0.4)
        p_title.paragraph_format.right_indent = Cm(0.4)

        secciones = _re.split(r'\n## ', analisis_medico)
        for i, seccion in enumerate(secciones):
            if not seccion.strip():
                continue
            lineas = seccion.strip().split('\n')
            if not lineas:
                continue

            # Título de subsección — limpiar cualquier símbolo
            titulo_sec = limpiar_markdown(lineas[0]).strip()
            p_sec = doc.add_paragraph()
            p_sec.paragraph_format.space_before = Pt(8)
            p_sec.paragraph_format.space_after  = Pt(2)
            r_sec = p_sec.add_run(titulo_sec)
            r_sec.bold = True; r_sec.font.size = Pt(10); r_sec.font.color.rgb = OLIVE

            # Cuerpo de la subsección
            cuerpo = limpiar_markdown('\n'.join(lineas[1:])).strip()
            for linea in cuerpo.split('\n'):
                linea = linea.strip()
                if not linea:
                    continue
                p_l = doc.add_paragraph()
                run = p_l.add_run(linea)
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(0x1c, 0x1c, 0x1c)
                p_l.paragraph_format.space_before = Pt(1)
                p_l.paragraph_format.space_after  = Pt(1)

        doc.add_paragraph()

    # ── Tabla plana de campos del formulario ─────────────────────
    def _titulo_seccion(texto):
        p_s = doc.add_paragraph()
        r_s = p_s.add_run(texto)
        r_s.bold = True; r_s.font.size = Pt(10); r_s.font.color.rgb = BLANCO
        p_s.paragraph_format.space_before = Pt(6)
        p_s.paragraph_format.space_after  = Pt(0)
        pPr_s = p_s._p.get_or_add_pPr()
        shd_s = OxmlElement('w:shd')
        shd_s.set(qn('w:val'), 'clear'); shd_s.set(qn('w:color'), 'auto')
        shd_s.set(qn('w:fill'), '2d3a2e'); pPr_s.append(shd_s)
        p_s.paragraph_format.left_indent  = Cm(0.3)
        p_s.paragraph_format.right_indent = Cm(0.3)

    def _nueva_tabla():
        tb = doc.add_table(rows=0, cols=2)
        tb.style = 'Table Grid'
        tb.alignment = WD_TABLE_ALIGNMENT.LEFT
        return tb

    def campo(tb, label, value):
        """Fila estándar: etiqueta | valor (blanco si no respondió)."""
        row = tb.add_row()
        c0, c1 = row.cells[0], row.cells[1]
        c0.text = label
        set_cell_bg(c0, 'f4f5ef')
        p0 = c0.paragraphs[0]; p0.runs[0].bold = True
        p0.runs[0].font.size = Pt(9); p0.runs[0].font.color.rgb = VERDE
        val = value
        if isinstance(val, list): val = ', '.join(str(v) for v in val if v)
        elif val is None: val = ''
        else: val = str(val)
        c1.text = val
        if c1.paragraphs[0].runs:
            c1.paragraphs[0].runs[0].font.size = Pt(10)

    def campo_yn(tb, label, value):
        """Fila Sí/No: muestra la selección con color destacado."""
        row = tb.add_row()
        c0, c1 = row.cells[0], row.cells[1]
        c0.text = label
        set_cell_bg(c0, 'f4f5ef')
        p0 = c0.paragraphs[0]; p0.runs[0].bold = True
        p0.runs[0].font.size = Pt(9); p0.runs[0].font.color.rgb = VERDE
        val = str(value).strip() if value else ''
        es_si = val.lower() in ['si', 'sí', 'yes', 'sí']
        p1 = c1.paragraphs[0]
        run = p1.add_run('● ' + (val if val else ''))
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = OLIVE if es_si else RGBColor(0x44, 0x44, 0x44)

    # Preparar historial estético
    hist    = data.get('historialEstetico', [])
    det_map = data.get('historialDetalle', {}) or {}
    hist_str = ''
    if hist:
        partes = []
        for trat in (hist if isinstance(hist, list) else [hist]):
            det   = det_map.get(trat, {}) if isinstance(det_map, dict) else {}
            fecha = det.get('fecha', '') if det else ''
            zona  = det.get('zona',  '') if det else ''
            txt   = trat
            if fecha: txt += f' — última sesión: {fecha}'
            if zona:  txt += f' — zona: {zona}'
            partes.append(txt)
        hist_str = '\n'.join(partes)

    contra    = data.get('contraindications', {})
    contra_si = [k for k, v in contra.items() if v == 'Si'] if isinstance(contra, dict) else []

    # ═══════════════════════════════════════════════════════════
    # TABLA UNIFICADA — snake_case con headers de sección
    # ═══════════════════════════════════════════════════════════
    from docx.enum.table import WD_ROW_HEIGHT_RULE

    tb = doc.add_table(rows=0, cols=2)
    tb.style = 'Table Grid'
    tb.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Anchos de columna: 40% clave / 60% valor
    from docx.shared import Inches
    for cell in tb.columns[0].cells if False else []:
        pass  # widths set via XML below

    def _sec(titulo):
        """Fila de sección: celda fusionada con fondo oscuro."""
        row = tb.add_row()
        row.height = Cm(0.75)
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        merged = row.cells[0].merge(row.cells[1])
        merged.text = titulo
        p = merged.paragraphs[0]
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(9)
        p.runs[0].font.color.rgb = BLANCO
        set_cell_bg(merged, '2d3a2e')
        p.paragraph_format.left_indent = Cm(0.3)

    def _fila(key, val):
        """Fila de dato: clave snake_case | valor."""
        if isinstance(val, list):
            val = ', '.join(str(x) for x in val if x)
        elif val is None:
            val = ''
        row = tb.add_row()
        row.height = Cm(0.80)
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        c0, c1 = row.cells[0], row.cells[1]
        c0.text = str(key)
        c1.text = str(val)
        set_cell_bg(c0, 'f4f5ef')
        p0 = c0.paragraphs[0]
        if p0.runs:
            p0.runs[0].font.size = Pt(9)
            p0.runs[0].font.color.rgb = VERDE
        p1 = c1.paragraphs[0]
        if p1.runs:
            p1.runs[0].font.size = Pt(10)

    # Preparar variables derivadas (misma lógica que antes)
    contra   = data.get('contraindications', {}) or {}
    hist_lst = data.get('historialEstetico', []) or []
    hist_low = [h.lower() for h in (hist_lst if isinstance(hist_lst, list) else [])]
    det_m    = data.get('historialDetalle', {}) or {}
    sint_lst = data.get('sintomasDigestivos', []) or []
    sint_low = ' '.join(s.lower() for s in sint_lst)
    piel_raw = data.get('pielProblemas', []) or []
    piel_low = ' '.join((piel_raw if isinstance(piel_raw, list) else [])).lower()
    areas_f  = data.get('areasFaciales', []) or []
    areas_c  = data.get('areasCorporales', []) or []
    areas_low= ' '.join((areas_f if isinstance(areas_f, list) else []) +
                         (areas_c if isinstance(areas_c, list) else [])).lower()

    def _yn_c(label):
        return 'SÍ' if contra.get(label, 'No') == 'Si' else 'NO'
    def _yn_h(nombre):
        return 'SÍ' if any(nombre.lower() in h for h in hist_low) else 'NO'
    def _fh(label):
        det = det_m.get(label, {}) or {}
        return det.get('fecha', '') if isinstance(det, dict) else ''
    def _zh(label):
        det = det_m.get(label, {}) or {}
        return det.get('zona', '') if isinstance(det, dict) else ''
    def _yn_s(kw):
        return 'SÍ' if kw.lower() in sint_low else 'NO'
    def _yn_p(kw):
        return 'SÍ' if any(k in piel_low for k in ([kw] if isinstance(kw, str) else kw)) else 'NO'
    def _g(k):
        v2 = data.get(k, '')
        if isinstance(v2, list): return ', '.join(str(x) for x in v2 if x)
        return str(v2) if v2 is not None else ''

    cond  = _g('condicionSistemica')
    fuma_v = _g('fuma')

    # ── 1. DATOS PERSONALES ─────────────────────────────────────
    _sec('1. DATOS PERSONALES')
    _fila('nombre_completo',          _g('nombre'))
    _fila('cedula',                   _g('cedula'))
    _fila('fecha_nacimiento',         _g('fechaNacimiento'))
    _fila('edad',                     _g('edad'))
    _fila('sexo',                     _g('sexo'))
    _fila('celular',                  _g('celular'))
    _fila('email',                    _g('email'))
    _fila('direccion',                _g('direccion'))
    _fila('ocupacion',                _g('ocupacion'))
    _fila('horario_laboral',          _g('horarioLaboral'))
    _fila('num_hijos',                _g('numHijos'))
    _fila('referencia_recomendacion', _g('comoConociste'))
    _fila('fecha_firma',              _g('fecha'))

    # ── 2. CONTACTO DE EMERGENCIA ────────────────────────────────
    _sec('2. CONTACTO DE EMERGENCIA')
    _fila('emergencia_nombre',   _g('contactoEmergencia'))
    _fila('emergencia_relacion', _g('contactoRelacion'))
    _fila('emergencia_telefono', _g('contactoTel'))

    # ── 3. MEDIDAS CORPORALES ────────────────────────────────────
    _sec('3. MEDIDAS CORPORALES')
    _fila('estatura', _g('estatura'))
    _fila('peso',     _g('peso'))
    _fila('imc',      _g('imc'))

    # ── 4. SALUD GENERAL ─────────────────────────────────────────
    _sec('4. SALUD GENERAL')
    _fila('sufre_enfermedad',  'NO' if cond in ['Sin enfermedades','','No'] else 'SÍ')
    _fila('enfermedad_detalle','' if cond in ['Sin enfermedades','','No'] else cond)
    _fila('embarazada',        _yn_c('Embarazo'))
    _fila('embarazo_semanas',  '')
    _fila('lactancia',         _yn_c('Lactancia'))
    _fila('sop',               _yn_c('SOP'))
    _fila('anticonceptivos',   _yn_c('Anticonceptivos'))
    _fila('menopausia',        _yn_c('Menopausia'))
    _fila('perimenopausia',    _g('perimenopausia') or 'NO')
    _fila('fuma',              'SÍ' if fuma_v.lower() not in ['no',''] else 'NO')
    _fila('fuma_frecuencia',   fuma_v.replace('Si - ','').replace('Sí - ','') if ' - ' in fuma_v else '')
    _fila('alcohol',           _g('alcohol'))
    _fila('toma_medicamentos', 'SÍ' if _g('medicamentos') not in ['Ninguna','No',''] else 'NO')
    _fila('medicamento_cual',  _g('medicamentos'))
    _fila('otras_condiciones', _g('condiciones'))
    _fila('horas_sueno',       _g('sueno'))
    _fila('hora_acuesta',      _g('horaDuerme'))
    _fila('hora_levanta',      _g('horaDespierta'))
    _fila('cansancio_dia',     _g('cansancioDia'))
    _fila('nivel_estres',      _g('nivelEstres'))

    # ── 5. SÍNTOMAS DIGESTIVOS ───────────────────────────────────
    _sec('5. SÍNTOMAS DIGESTIVOS E INTOLERANCIAS')
    _fila('hinchazon_abdominal',       _yn_s('hinchazon'))
    _fila('gases_flatulencias',        _yn_s('gases'))
    _fila('nauseas_malestar',          _yn_s('nauseas'))

    # ── 6. ALERGIAS ──────────────────────────────────────────────
    _sec('6. ALERGIAS')
    _fila('alergia_medicamentos', 'SÍ' if 'medicamento' in _g('alergias').lower() else 'NO')

    # ── 7. CIRUGÍAS Y ANTECEDENTES ───────────────────────────────
    _sec('7. CIRUGÍAS Y ANTECEDENTES')
    _fila('cirugias',              'SÍ' if _g('cirugias') not in ['Ninguna','No',''] else 'NO')
    _fila('cirugias_detalle',      _g('cirugias'))
    _fila('antecedentes_familiares','SÍ' if _g('antecedentesFam') not in ['Ninguno','No',''] else 'NO')
    _fila('antecedentes_detalle',  _g('antecedentesFamDet'))

    # Función local: normaliza Sí/No/No respondido → SÍ/NO
    def _yn_norm(key):
        val = _g(key).strip()
        if not val or val.lower() in ['no', 'no respondido', '']:
            return 'NO'
        return 'SÍ' if val.lower() in ['sí', 'si', 'yes'] else val

    _fila('infecciones_cutaneas',          _yn_norm('infeccionesCutaneas'))
    _fila('infecciones_cutaneas_detalle',  _g('infeccionesDet'))
    _fila('dispositivos_medicos',          _yn_norm('dispositivosMedicos'))
    _fila('dispositivos_medicos_detalle',  _g('dispositivosDet'))

    # ── 8. ALIMENTACIÓN ──────────────────────────────────────────
    _sec('8. ALIMENTACIÓN Y NUTRICIÓN')
    _fila('observaciones_alimentarias', _g('notasAlimentacion'))

    # ── 9. PIEL Y RUTINA ─────────────────────────────────────────
    _sec('9. PIEL, RUTINA Y PROTECCIÓN SOLAR')
    _fila('tipo_piel_mixta',          _g('pielTipo') if 'mixta' in _g('pielTipo').lower() else _g('pielTipo'))
    _fila('prob_manchas_solares',     'SÍ' if _yn_p('manchas') == 'SÍ' else 'NO')
    _fila('prob_pecas',               'SÍ' if _yn_p('peca') == 'SÍ' else 'NO')
    _fila('prob_lineas_finas',        'SÍ' if _yn_p(['línea','linea','arruga']) == 'SÍ' else 'NO')
    _fila('prob_flacidez',            'SÍ' if _yn_p('flacidez') == 'SÍ' else 'NO')
    _fila('prob_perdida_luminosidad', 'SÍ' if _yn_p(['luminosidad','opac']) == 'SÍ' else 'NO')
    _fila('exposicion_solar',         _g('solar'))
    _fila('protector_solar',          _g('usaProtectorSolar'))
    _fila('protector_marca',          _g('protectorMarca'))
    _fila('protector_fps',            _g('spf').split()[0] if _g('spf') else '')
    _fila('protector_hora',           _g('protectorHora'))
    _fila('protector_reaplica',       _g('reaplicaSolar'))
    _fila('rutina_diaria_cuidado',    _g('tieneRutina'))
    _fila('rutina_manana',            _g('rutinaManana'))
    _fila('rutina_tarde',             _g('rutinaTarde'))
    _fila('rutina_noche',             _g('rutinaNoche'))
    _fila('rutina_limpieza',          'Limpieza'   if any(x in (_g('rutinaManana')+_g('rutinaNoche')).lower() for x in ['limpiad','jabon','jabón','espuma']) else '')
    _fila('rutina_hidratacion',       'Hidratación' if any(x in (_g('rutinaManana')+_g('rutinaNoche')).lower() for x in ['hidrat','crema']) else '')
    _fila('rutina_retinol',           'Retinol'    if 'retinol' in (_g('rutinaManana')+_g('rutinaNoche')).lower() else '')
    _fila('productos_cosmeticos_frecuentes', _g('productosFrecuentes'))
    _fila('actividad_fisica',         _g('actFisica'))

    # ── 10. HISTORIAL ESTÉTICO ───────────────────────────────────
    _sec('10. HISTORIAL ESTÉTICO Y TRATAMIENTOS PREVIOS')
    _fila('tratamientos_previos',     'SÍ' if hist_lst else 'NO')
    _fila('tratamiento_otro',         _yn_h('otro'))
    _fila('complicaciones_esteticas', 'SÍ' if _g('complicacionesDet') else 'NO')
    _fila('complicaciones_esteticas_desc', _g('complicacionesDet'))

    # ── 11. OBJETIVOS ────────────────────────────────────────────
    _sec('11. OBJETIVOS Y EXPECTATIVAS')
    _fila('area_arrugas',         'Arrugas'         if 'arrugas'  in areas_low else '')
    _fila('area_manchas',         'Manchas'         if 'manchas'  in areas_low else '')
    _fila('area_flacidez_facial', 'Flacidez facial' if 'flacidez' in areas_low else '')
    _fila('area_ojeras',          'Ojeras'          if 'ojeras'   in areas_low else '')
    _fila('area_celulitis',       'Celulitis'       if 'celulitis' in areas_low else '')
    _fila('grasa_zona',           'Grasa localizada' if 'grasa'   in areas_low else '')
    _fila('prioridad_principal',  _g('prioridad'))
    _fila('satisfaccion',         _g('satisfaccion'))
    _fila('entiende_sesiones',    _g('entiendeSesiones'))

    # ── 12. DECLARACIONES ────────────────────────────────────────
    _sec('12. DECLARACIONES Y CONSENTIMIENTO')
    _fila('declara_veracidad',  'Acepto')
    _fila('declara_riesgo',     'Acepto')
    _fila('declara_cambios',    'Acepto')
    _fila('declara_resultados', 'Acepto')
    _fila('autoriza_contacto',  'Acepto')

    doc.add_paragraph()
    p_pie = doc.add_paragraph(); p_pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rp = p_pie.add_run('Centro Carvajal · Líderes en Medicina Estética en Panamá · centrocarvajal.com')
    rp.font.size = Pt(8); rp.font.color.rgb = GRIS


    nombre_safe = re.sub(r'[^a-zA-Z0-9]', '_', data.get('nombre', 'Paciente'))
    docx_path = f'/tmp/Cuestionario_{nombre_safe}_{uuid.uuid4().hex[:6]}.docx'
    doc.save(docx_path)
    return docx_path


def worker(job_id, data, faltantes, fotos=None, modelo='claude', session_id=None):
    try:
        jobs[job_id] = {'status': 'working', 'msg': 'Generando plan con IA (puede tomar 1-2 min)...'}

        # session_id ya fue guardado en /enviar antes de lanzar el worker.
        # generar_plan_ia lo usará para actualizar el progreso entre pasos.
        plan_json = generar_plan_ia(data, job_id, modelo=modelo, session_id=session_id)
        if 'error' in plan_json:
            jobs[job_id] = {'status': 'error', 'msg': plan_json['error']}
            return

        jobs[job_id] = {'status': 'working', 'msg': 'Construyendo borrador y plan...'}
        print(f'[worker] Renderizando plan HTML...')
        html = render_plan(plan_json, data)
        print(f'[worker] Plan HTML OK ({len(html)} chars)')
        borrador_html = render_borrador(plan_json, data, job_id)
        print(f'[worker] Borrador HTML OK')

        nombre    = data.get('nombre', 'Paciente')
        if modelo == 'gemini':
            sufijo_modelo = '_gemini'
        elif modelo == 'groq':
            sufijo_modelo = '_groq'
        else:
            sufijo_modelo = '_claude'
        html_name = 'Plan_' + re.sub(r'[^a-zA-Z0-9]', '_', nombre) + '_' + datetime.now().strftime('%Y%m%d') + sufijo_modelo + '.html'
        html_path = os.path.join(PLANES_DIR, html_name)
        with open(html_path, 'w', encoding='utf-8') as f:
            f.write(html)

        base_url = os.environ.get('BASE_URL', 'https://metodo.centrocarvajal.com')
        html_url_local = f'{base_url}/planes_generados/{html_name}'

        # Subir plan final a Cloudinary
        print(f'[worker] Subiendo plan a Cloudinary...')
        html_url_cdn = subir_plan_cloudinary(html_path, html_name, job_id=job_id)
        print(f'[worker] Cloudinary plan: {html_url_cdn}')
        html_url = html_url_cdn if html_url_cdn else html_url_local

        # Subir borrador editable a Cloudinary
        jobs[job_id] = {'status': 'working', 'msg': 'Guardando borrador editable...'}
        subir_borrador_cloudinary(borrador_html, job_id)
        borrador_url = f'{base_url}/borrador/{job_id}'

        jobs[job_id] = {'status': 'working', 'msg': 'Enviando correos...'}
        fecha_hoy = datetime.now().strftime('%d/%m/%Y a las %H:%M')

        # Generar análisis clínico con Claude (para el médico)
        analisis_medico = None
        try:
            time.sleep(2)  # pausa tras las 3 secciones del plan
            analisis_medico = generar_analisis_medico(data)
        except Exception as e:
            print(f'[worker] Error analisis medico: {e}')

        # Generar .docx del cuestionario para adjuntar
        docx_cuestionario = None
        try:
            docx_cuestionario = generar_docx_cuestionario(data, plan_json=plan_json, analisis_medico=analisis_medico)
            print(f'[worker] Cuestionario .docx: {docx_cuestionario}')
        except Exception as e:
            print(f'[worker] Error docx cuestionario: {e}')

        # ── CORREO 2: plan generado — borrador editable + .docx con análisis ──
        # Las fotos ya se enviaron en el correo 1 (formulario inmediato)
        adjuntos_plan = []
        if docx_cuestionario and os.path.exists(docx_cuestionario):
            adjuntos_plan.append(docx_cuestionario)

        enviar_resend(
            f'\u2705 Plan IA generado \u2014 {nombre} ({fecha_hoy})',
            email_plan_completo(data, borrador_url, html_url, faltantes),
            MAIL_TO,
            adjuntos_extra=adjuntos_plan,
            cc=MAIL_CC or None
        )
        print(f'[worker] Correo 2 (plan) enviado: {nombre}')

        # ── Alerta si Claude se quedó sin créditos durante este job ──
        if _claude_sin_creditos:
            error_raw = _claude_sin_creditos[-1]
            try:
                enviar_resend(
                    f'⚠️ Claude sin créditos — se usó Groq como fallback ({nombre})',
                    email_alerta_creditos(error_raw, nombre),
                    MAIL_TO,
                    cc=MAIL_CC or None
                )
                print(f'[worker] Alerta de créditos enviada')
            except Exception as e:
                print(f'[worker] Error enviando alerta créditos: {e}')
            _claude_sin_creditos.clear()

        # Limpiar temporales
        for f_path in adjuntos_plan:
            try: os.unlink(f_path)
            except: pass

        jobs[job_id] = {
            'status'      : 'done',
            'nombre'      : nombre,
            'html_url'    : html_url,
            'html_name'   : html_name,
            'borrador_url': borrador_url,
            'faltantes'   : faltantes,
        }

        # No se elimina la sesión — expira automáticamente a las 72h (SESION_TTL_HORAS)
        # Esto permite recuperar datos si algo falla en los días siguientes

    except Exception as e:
        import traceback
        print(f'[worker] EXCEPCION: {e}')
        print(traceback.format_exc())
        jobs[job_id] = {'status': 'error', 'msg': str(e)}
        # Enviar correo de error con botón para regenerar desde el panel
        try:
            nombre = data.get('nombre', 'el paciente')
            _enviar_correo_error_plan(nombre, session_id, str(e))
        except Exception as e2:
            print(f'[worker] Error enviando correo de fallo: {e2}')


# ════════════════════════════════════════════════════════════
# LEER DOCX
# ════════════════════════════════════════════════════════════

def leer_docx(path):
    try:
        from docx import Document as _Doc
        doc = _Doc(path)
        pares = []
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                if len(cells) >= 2:
                    key = cells[0].lower().strip()
                    val = cells[1].strip()
                    if re.match(r'^[a-z][a-z0-9_]+$', key):
                        pares.append(f'{key}\t{val}')
        return '\n'.join(pares) if pares else None
    except Exception as e:
        print(f'leer_docx error: {e}')
        return None




# ════════════════════════════════════════════════════════════
# MAPEAR FORMULARIO WEB → data dict
# ════════════════════════════════════════════════════════════

def _sanitizar_texto(texto, max_len=None):
    """Limpia texto libre de caracteres peligrosos para prevenir inyecciones.
    - Elimina caracteres de control (excepto salto de línea y tab)
    - Recorta a max_len si se especifica
    - Devuelve str limpio"""
    if not isinstance(texto, str):
        texto = str(texto) if texto is not None else ''
    # Eliminar caracteres de control peligrosos (excepto \n, \t, \r)
    import unicodedata
    texto = ''.join(
        c for c in texto
        if unicodedata.category(c) != 'Cc' or c in ('\n', '\t', '\r')
    )
    # Limitar longitud si se especifica
    if max_len and len(texto) > max_len:
        texto = texto[:max_len]
    return texto.strip()


def _sanitizar_form(f):
    """Recorre el dict del formulario y sanitiza todos los valores de texto."""
    if not isinstance(f, dict):
        return f
    resultado = {}
    for key, val in f.items():
        if isinstance(val, str):
            resultado[key] = _sanitizar_texto(val)
        elif isinstance(val, list):
            resultado[key] = [
                _sanitizar_texto(v) if isinstance(v, str) else v
                for v in val
            ]
        elif isinstance(val, dict):
            resultado[key] = _sanitizar_form(val)
        else:
            resultado[key] = val
    return resultado


def _mapear_formulario(f):
    """Convierte el JSON de index.html al mismo dict que usan generar_plan_ia() y render_plan()."""

    def s(key, default=''):
        val = f.get(key, default)
        if isinstance(val, list):
            return ', '.join(str(v) for v in val if v)
        return str(val).strip() if val else default

    def lst(key):
        val = f.get(key, [])
        return val if isinstance(val, list) else []

    est = s('estatura')
    pes = s('peso')
    imc = 'No registrado'
    if est and pes:
        try: imc = f'{float(pes) / (float(est)/100)**2:.1f}'
        except: pass

    contra_raw = f.get('contraindications', {})
    contra = {k: ('Si' if str(v).lower() in ['si','sí','yes'] else 'No') for k, v in contra_raw.items()}

    sintomas_map = {
        'hinchazon_abdominal': 'Hinchazon',
        'gases':               'Gases',
        'nauseas':             'Nauseas',
    }
    sintomas = [sintomas_map[k] for k in lst('sintomasDigestivos') if k in sintomas_map]
    faciales   = lst('areasFaciales')
    corporales = lst('areasCorporales')
    rutina_m = s('rutinaManana')
    rutina_n = s('rutinaNoche')
    rutina_facial = (rutina_m + ' | ' + rutina_n).strip(' |') if (rutina_m or rutina_n) else 'No tiene rutina facial'

    return {
        'nombre':              s('nombre'),
        'cedula':              s('cedula'),
        'direccion':           s('direccion'),
        'edad':                s('edad'),
        'fechaNacimiento':     s('fechaNacimiento'),
        'sexo':                s('sexo'),
        'ocupacion':           s('ocupacion'),
        'horarioLaboral':      s('horarioLaboral'),
        'email':               s('email'),
        'celular':             s('celular'),
        'comoConociste':       s('comoConociste'),
        'fecha':               datetime.now().strftime('%d de %B, %Y'),
        'estatura':            est or None,
        'peso':                pes or None,
        'imc':                 imc,
        'pielTipo':            s('pielTipo'),
        'pielProblemas':       lst('pielProblemas'),
        'rutinaFacial':        rutina_facial,
        'rutinaManana':        rutina_m,
        'rutinaNoche':         rutina_n,
        'productosFrecuentes': s('productosFrecuentes'),
        'solar':               s('solar'),
        'spf':                 s('spf'),
        'actFisica':           s('actFisica'),
        'sueno':               s('sueno'),
        'horaDespierta':       s('horaDespierta'),
        'horaDuerme':          s('horaDuerme'),
        'cansancioDia':        'Si' if s('cansancioDia').lower() in ['si','sí','yes'] else 'No',
        'fuma':                s('fuma'),
        'alcohol':             s('alcohol'),
        'condicionSistemica':  s('condicionSistemica') or 'Sin enfermedades',
        'condiciones':         s('condiciones'),
        'medicamentos':        s('medicamentos'),
        'cirugias':            s('cirugias') or 'Ninguna',
        'antecedentesFam':     'Ninguno',
        'alergias':            s('alergias') or 'Ninguna',
        'alergiasDetalle':     f.get('alergiasDetalle', {}),
        'contraindications':   contra,
        'embarazo':            s('embarazo'),
        'lactancia':           s('lactancia'),
        'anticonceptivos':     s('anticonceptivos'),
        'sop':                 s('sop'),
        'menopausia':          s('menopausia'),
        'perimenopausia':      s('perimenopausia'),
        'antecedentesFam':     s('antecedentesFam'),
        'antecedentesFamDet':  s('antecedentesFamDet'),
        'comoConociste':       s('comoConociste'),
        'areasFaciales':       faciales,
        'areasCorporales':     corporales,
        'prioridad':           s('prioridad'),
        'satisfaccion':        s('satisfaccion'),
        'historialEstetico':   lst('historialEstetico'),
        'historialDetalle':    f.get('historialDetalle', {}),
        'entiendeSesiones':    'Si' if s('entiendeSesiones').lower() in ['si','sí','yes'] else 'No',
        'sintomasDigestivos':  sintomas,
        'notasAlimentacion':   s('notasAlimentacion'),
        'nivelEstres':         s('nivelEstres'),
        'numHijos':            s('numHijosVal') or s('numHijos'),
        # Contacto emergencia
        'contactoEmergencia':  s('contactoEmergencia'),
        'contactoRelacion':    s('contactoRelacion'),
        'contactoTel':         s('contactoTel'),
        # Solar detallado
        'usaProtectorSolar':   s('usaProtectorSolar'),
        'protectorMarca':      s('protectorMarca'),
        'protectorHora':       s('protectorHora'),
        'reaplicaSolar':       s('reaplicaSolar'),
        # Rutina
        'tieneRutina':         s('tieneRutina'),
        # Historial estético extra
        'laserActualDet':      s('laserActualDet'),
        'complicacionesDet':   s('complicacionesDet'),
        # Contraindicaciones clínicas adicionales
        'infeccionesCutaneas': s('infeccionesCutaneas'),
        'dispositivosMedicos': s('dispositivosMedicos'),
        # Rutina tarde (campo opcional, puede estar vacío si el formulario no lo tiene)
        'rutinaTarde':         s('rutinaTarde'),
        'notasStaff':          '',
    }

# ════════════════════════════════════════════════════════════
# PARSEAR CUESTIONARIO
# ════════════════════════════════════════════════════════════

def parsear_cuestionario(texto):
    raw = {}
    for linea in texto.split('\n'):
        linea = linea.strip()
        if not linea: continue
        if '\t' in linea:
            parts = linea.split('\t', 1)
            key, val = parts[0], parts[1]
        else:
            m = re.match(r'^([a-z][a-z0-9_]+)\s+(.+)$', linea, re.I)
            if m: key, val = m.group(1), m.group(2)
            else: continue
        key, val = key.lower().strip(), val.strip()
        if key and val: raw[key] = val

    def v(k):
        val = raw.get(k)
        if not val or val == '0' or val.lower() in ['ninguna','ninguno','aun no','n/a','nada']:
            return None
        return val

    def si(k):
        return raw.get(k, 'NO').upper() in ['SI', 'SÍ', 'S', 'YES']

    faltantes = []
    if not (raw.get('peso') and raw.get('peso') != '0'):
        faltantes.append('Peso')
    if not (raw.get('altura') or raw.get('estatura') or raw.get('talla')):
        faltantes.append('Estatura')
    if not (raw.get('piel_tipo') or raw.get('tipo_piel') or raw.get('tipo_piel_mixta')):
        faltantes.append('Tipo de piel')
    if not raw.get('prioridad_principal'):
        faltantes.append('Prioridad principal')
    if not (raw.get('satisfaccion') or raw.get('satisfaccion_actual')):
        faltantes.append('Satisfaccion (1-10)')
    if not raw.get('expectativas'):
        faltantes.append('Expectativas')
    if not raw.get('exposicion_solar'):
        faltantes.append('Exposicion solar')

    historial = []
    det_hist = {}
    for k, l in [('botox','Botox'),('rellenos','Rellenos'),('hilos','Hilos PDO'),('peeling','Peeling'),
                 ('laser','Laser'),('microblading','Microblading'),('mesoterapia','Mesoterapia'),
                 ('radiofrecuencia','Radiofrecuencia'),('criolipolisis','Criolipólisis')]:
        if si('tratamiento_' + k):
            historial.append(l)
            det = {}
            fecha = v(k + '_fecha')
            if fecha: det['fecha'] = fecha
            if k == 'rellenos' and v('rellenos_zonas'): det['zona'] = v('rellenos_zonas')
            if k == 'laser' and v('laser_tipo'): det['zona'] = v('laser_tipo')
            if det: det_hist[l] = det

    contra = {}
    for c, l in [('embarazada','Embarazo'),('lactancia','Lactancia'),('anticonceptivos','Anticonceptivos'),
                 ('sop','SOP'),('menopausia','Menopausia'),('fuma','Tabaquismo'),
                 ('alergia_lidocaina','Alergia lidocaina'),('alergia_penicilina','Alergia penicilina'),
                 ('alergia_yodo','Alergia yodo'),('alergia_aines','Alergia AINEs'),
                 ('alergia_latex','Alergia latex')]:
        contra[l] = 'Si' if si(c) else 'No'

    sintomas = []
    for c, l in [('hinchazon_abdominal','Hinchazon'),('gases_flatulencias','Gases'),
                 ('estrenimiento_diarrea','Estrenimiento'),('cansancio_comidas','Cansancio tras comer'),
                 ('digestion_lenta','Digestion lenta'),('nauseas_malestar','Nauseas')]:
        if si(c): sintomas.append(l)

    intol = []
    if si('sintomas_lacteos'):    intol.append('Lacteos')
    if si('sintomas_gluten'):     intol.append('Gluten')
    if si('sintomas_procesados'): intol.append('Procesados')

    faciales, corporales = [], []
    if raw.get('area_flacidez_facial'): faciales.append('Flacidez facial')
    if raw.get('area_grasa'):           corporales.append('Grasa localizada')
    if raw.get('area_arrugas'):         faciales.append('Arrugas')
    if raw.get('area_manchas'):         faciales.append('Manchas')
    if raw.get('area_ojeras'):          faciales.append('Ojeras')
    if raw.get('area_celulitis'):       corporales.append('Celulitis')
    prio = (v('prioridad_principal') or '').lower()
    if 'rostro' in prio or 'cara' in prio: faciales.append('Rejuvenecimiento facial')
    if 'cabello' in prio or 'pelo' in prio: faciales.append('Caida de cabello')
    if 'mancha' in prio or 'melasma' in prio: faciales.append('Manchas')
    faciales   = list(dict.fromkeys(faciales))
    corporales = list(dict.fromkeys(corporales))

    ale_list = []
    for c, l in [('alergia_lidocaina','Lidocaina'),('alergia_penicilina','Penicilina'),
                 ('alergia_yodo','Yodo'),('alergia_aines','AINEs'),('alergia_latex','Latex'),
                 ('alergia_aloe','Aloe'),('alergia_fragancias','Fragancias')]:
        if si(c): ale_list.append(l)
    if si('alergia_medicamentos'): ale_list.append(v('medicamentos_cuales') or 'Medicamentos')

    est = v('altura') or v('talla') or v('estatura')
    pes = v('peso')
    imc = 'No registrado'
    if est and pes:
        try: imc = f'{float(pes) / (float(est)/100)**2:.1f}'
        except: pass

    # Proteinas — incluir todos los campos individuales
    proteinas_partes = [v('proteina_pollo'), v('proteina_pescado'), v('proteina_mariscos'),
                        v('proteina_res'), v('proteina_cerdo'), v('proteina_huevos'),
                        v('proteina_legumbres'), v('proteinas_otras'), v('proteinas_frecuentes')]
    proteinas_str = ', '.join(filter(None, proteinas_partes))

    # Carbohidratos — incluir campos individuales
    carbs_partes = [v('carb_arroz_blanco'), v('carb_pasta'), v('carb_pan'), v('carb_batata')]
    carbs_str = ', '.join(filter(None, carbs_partes)) or v('carbohidratos') or ''

    data = {
        'nombre':              v('nombre_completo') or '',
        'cedula':              v('cedula') or '',
        'fechaNacimiento':     v('fecha_nacimiento') or '',
        'celular':             v('celular') or '',
        'direccion':           v('direccion') or '',
        'edad':                v('edad') or '',
        'sexo':                v('sexo') or '',
        'ocupacion':           v('ocupacion') or v('profesion_trabajo') or '',
        'actLaboral':          v('nivel_actividad_laboral') or v('actividad_laboral') or '',
        'horarioLaboral':      v('horario_laboral') or '',
        'comoConociste':       v('referencia_recomendacion') or v('como_conocio_clinica') or '',
        'email':               v('email') or '',
        'fecha':               datetime.now().strftime('%d de %B, %Y'),
        'estatura':            est,
        'peso':                pes,
        'imc':                 imc,
        'contactoEmergencia':  v('emergencia_nombre') or '',
        'contactoRelacion':    v('emergencia_relacion') or '',
        'contactoTel':         v('emergencia_telefono') or '',
        'pielTipo':            v('piel_tipo') or v('tipo_piel') or v('tipo_piel_mixta') or '',
        'pielProblemas':       faciales,
        'rutinaFacial':        (v('rutina_manana') or '') + ' | ' + (v('rutina_noche') or '') if si('rutina_diaria_cuidado') else 'No tiene rutina facial',
        'rutinaManana':        v('rutina_manana') or '',
        'rutinaNoche':         v('rutina_noche') or '',
        'productosFrecuentes': v('productos_cosmeticos_frecuentes') or v('productos_frecuentes') or '',
        'solar':               v('exposicion_solar') or v('usa_solar') or '',
        'spf':                 (v('protector_fps') or v('spf_factor') or '') + (' ' + (v('protector_marca') or '')).rstrip(),
        'actFisica':           v('actividad_fisica') or '',
        'sueno':               (lambda _h, _c: _h if 'hora' in _h.lower() and not _c else (_h + (' horas ' if _h else '') + _c).strip())(v('horas_sueno') or '', v('calidad_sueno') or ''),
        'horaDespierta':       v('hora_levanta') or v('hora_despierta') or '',
        'horaDuerme':          v('hora_acuesta') or v('hora_duerme') or '',
        'cansancioDia':        'Si' if si('cansancio_dia') or si('cansancio_diurno') else 'No',
        'nivelEstres':         v('nivel_estres') or '',
        'numHijos':            v('num_hijos') or v('numero_hijos') or '',
        'fuma':                ('Si - ' + (v('fuma_frecuencia') or v('fuma_cantidad') or '').strip()).rstrip(' -') if si('fuma') else 'No',
        'alcohol':             'Si' if si('alcohol') or (v('consume_alcohol') or '').lower() not in ['no','no respondido',''] else 'No',
        'condicionSistemica':  (v('enfermedad_detalle') or v('condicion_sistemica') or 'Si') if si('sufre_enfermedad') else (v('condicion_sistemica') or 'Sin enfermedades'),
        'condiciones':         v('otras_condiciones') or v('condiciones_medicas') or '',
        'medicamentos':        v('medicamento_cual') or v('medicamentos') or '',
        'cirugias':            (v('cirugias_detalle') or v('cirugias_previas') or 'Si') if si('cirugias') else (v('cirugias_previas') or 'Ninguna'),
        'antecedentesFam':     (v('antecedentes_detalle') or 'Si') if si('antecedentes_familiares') else 'Ninguno',
        'antecedentesFamDet':  v('antecedentes_detalle') or v('detalle_antecedentes') or '',
        'alergias':            ', '.join(ale_list) if ale_list else 'Ninguna',
        'contraindications':   contra,
        'embarazo':            'Si' if si('embarazada') else 'No',
        'lactancia':           v('lactancia') or ('Si' if si('lactancia') else 'No'),
        'anticonceptivos':     v('anticonceptivos') or ('Si' if si('anticonceptivos') else 'No'),
        'sop':                 v('sop') or ('Si' if si('sop') else 'No'),
        'menopausia':          v('menopausia') or ('Si' if si('menopausia') else 'No'),
        'perimenopausia':      v('perimenopausia') or ('Si' if si('perimenopausia') else 'No'),
        'alergia_lidocaina':   'Si' if si('alergia_lidocaina') else 'No',
        'alergia_penicilina':  'Si' if si('alergia_penicilina') else 'No',
        'alergia_yodo':        'Si' if si('alergia_yodo') else 'No',
        'alergia_aines':       'Si' if si('alergia_aines') else 'No',
        'alergia_latex':       'Si' if si('alergia_latex') else 'No',
        'alergia_aloe':        'Si' if si('alergia_aloe') else 'No',
        'alergia_fragancias':  'Si' if si('alergia_fragancias') else 'No',
        'evacuacion':          v('evacuacion') or '',
        'sintLacteos':         'Si' if si('sintomas_lacteos') else 'No',
        'sintGluten':          'Si' if si('sintomas_gluten') else 'No',
        'sintProcesados':      'Si' if si('sintomas_procesados') else 'No',
        'usaProtectorSolar':   'Si' if si('protector_solar') or si('usa_protector_solar') else 'No',
        'protectorMarca':      v('protector_marca') or '',
        'protectorHora':       v('protector_hora') or '',
        'reaplicaSolar':       'Si' if si('protector_reaplica') else 'No',
        'tieneRutina':         'Si' if si('rutina_diaria_cuidado') else 'No',
        'laserActualDet':      v('laser_actual_detalle') or '',
        'complicacionesDet':   v('complicaciones_esteticas_desc') or '',
        'proteinasEvitar':     v('proteinas_evitar') or '',
        'carbosEvitar':        v('carbohidratos_evitar') or '',
        'verdurasEvitar':      v('verduras_evitar') or '',
        'grasasEvitar':        v('grasas_evitar') or '',
        'grasasEvitarPorque':  v('grasas_evitar_porque') or '',
        'areasFaciales':       faciales,
        'areasCorporales':     corporales,
        'prioridad':           v('prioridad_principal') or '',
        'expectativas':        v('expectativas') or '',
        'satisfaccion':        (v('satisfaccion') or v('satisfaccion_actual') or '').replace('/10','').strip(),
        'historialEstetico':   list(dict.fromkeys(historial)),
        'historialDetalle':    det_hist,
        'laserActivo':         'Si' if si('laser_actual') or si('laser_activo') else 'No',
        'intolerancias':       intol,
        'sintomasDigestivos':  sintomas,
        'proteinas':           proteinas_str,
        'carbohidratos':       carbs_str,
        'verduras':            v('verduras_consume') or v('verduras') or '',
        'frutas':              v('carb_frutas') or v('frutas') or '',
        'alimentosEvitar':     ', '.join(filter(None, [v('proteinas_evitar'), v('carbohidratos_evitar'), v('verduras_evitar'), v('alimentos_evitar')])),
        'postres':             v('postres_favoritos') or v('postres_dulces') or '',
        'bebidas':             (v('bebidas_azucaradas_cuales') or 'Si') if si('bebidas_azucaradas') else (v('bebidas_habituales') or 'No'),
        'notasAlimentacion':   v('observaciones_alimentarias') or v('notas_alimentacion') or '',
        'notasStaff':          '',
    }
    return data, faltantes


# ════════════════════════════════════════════════════════════
# GENERAR PLAN CON CLAUDE API — 3 LLAMADAS SEGMENTADAS
# ════════════════════════════════════════════════════════════

import time

def _datos_paciente(d):
    est = d.get('estatura', '')
    pes = d.get('peso', '')
    imc = d.get('imc', 'No registrado')
    contra_activas = [k for k, v in d.get('contraindications', {}).items() if v == 'Si']
    contra_txt = 'CONTRAINDICACIONES ACTIVAS: ' + ', '.join(contra_activas) if contra_activas else 'Sin contraindicaciones activas.'
    return f"""DATOS DEL PACIENTE:
Nombre: {d['nombre']} | Edad: {d['edad']} | Sexo: {d['sexo']}
Ocupacion: {d['ocupacion']} | Horario: {d['horarioLaboral']}
Fecha evaluacion: {d['fecha']}
Estatura: {est}cm | Peso: {pes}kg | IMC: {imc}

PIEL: {d['pielTipo']} | Problemas: {', '.join(d['pielProblemas'])}
Rutina manana: {d['rutinaManana']} | Noche: {d['rutinaNoche']}
Productos: {d['productosFrecuentes']} | Solar: {d['solar']} | SPF: {d['spf']}

HABITOS: Act.fisica: {d['actFisica']} | Sueno: {d['sueno']}
Fuma: {d['fuma']} | Alcohol: {d['alcohol']}

SALUD: {contra_txt}
Condicion sistemica: {d['condicionSistemica']}
Condiciones: {d['condiciones']}
Medicamentos: {d['medicamentos']}
Cirugias: {d['cirugias']}
Alergias: {d['alergias']}

OBJETIVOS: Faciales: {', '.join(d['areasFaciales'])} | Corporales: {', '.join(d['areasCorporales'])}
Prioridad (palabras del paciente): "{d['prioridad']}"
Satisfaccion: {d['satisfaccion']}/10
Historial estetico: {', '.join(d['historialEstetico']) or 'Ninguno'}

ALIMENTACION:
Notas: {d['notasAlimentacion']}

CONTEXTO PERSONAL ADICIONAL:
Numero de hijos: {d.get('numHijos','No especificado')}
Nivel de estres (1-10): {d.get('nivelEstres','No especificado')}"""


def _llamar_claude(num, total, system_prompt, user_msg, max_tok=6000):
    MAX_REINTENTOS = 3
    ESPERAS = [10, 25, 45]  # backoff escalonado en segundos para error 529

    for intento in range(MAX_REINTENTOS + 1):
        if intento > 0:
            espera = ESPERAS[intento - 1]
            print(f"[{num}/{total}] Reintento {intento}/{MAX_REINTENTOS} — esperando {espera}s...")
            time.sleep(espera)

        print(f"[{num}/{total}] Iniciando{'...' if intento == 0 else f' (intento {intento + 1})'}")
        t0 = time.time()
        resp = req.post(
            'https://api.anthropic.com/v1/messages',
            headers={
                'Content-Type': 'application/json',
                'x-api-key': CLAUDE_KEY,
                'anthropic-version': '2023-06-01',
            },
            json={
                'model': 'claude-sonnet-4-6',
                'max_tokens': max_tok,
                'system': system_prompt,
                'messages': [{'role': 'user', 'content': user_msg}],
            },
            timeout=300
        )
        elapsed = round(time.time() - t0, 1)

        # 529 = overloaded → reintentar con backoff
        if resp.status_code == 529 and intento < MAX_REINTENTOS:
            print(f"[{num}/{total}] 529 Overloaded — reintentando...")
            continue

        if resp.status_code != 200:
            error_text = resp.text
            print(f"[{num}/{total}] ERROR {resp.status_code}: {error_text[:400]}")
            if resp.status_code == 400 and 'credit balance is too low' in error_text:
                print(f"[{num}/{total}] Sin créditos Claude — fallback automático a Groq")
                if error_text not in _claude_sin_creditos:
                    _claude_sin_creditos.append(error_text[:500])
                return _llamar_groq(num, total, system_prompt, user_msg, max_tok)
            return None, f'Error API Claude ({resp.status_code}): {error_text[:300]}'

        rj = resp.json()
        txt = rj['content'][0]['text']
        stop_reason = rj.get('stop_reason', '')
        usage = rj.get('usage', {})
        tok_in  = usage.get('input_tokens', '?')
        tok_out = usage.get('output_tokens', '?')
        print(f"[{num}/{total}] OK — {elapsed}s | input: {tok_in} tokens | output: {tok_out} tokens | stop: {stop_reason}")

        if stop_reason == 'max_tokens':
            return None, f'Llamada {num} truncada por limite de tokens.'

        txt = re.sub(r'^```json\s*', '', txt.strip())
        txt = re.sub(r'^```\s*', '', txt)
        txt = re.sub(r'```\s*$', '', txt).strip()

        try:
            result = json.loads(txt)
            print(f"[{num}/{total}] JSON OK — claves: {list(result.keys())}")
            return result, None
        except Exception as e:
            print(f"[{num}/{total}] JSON invalido: {e} | inicio: {txt[:300]}")
            return None, f'JSON invalido en llamada {num}: {str(e)[:150]}'

    return None, f'Llamada {num} falló tras {MAX_REINTENTOS} reintentos (overloaded)'



def _llamar_gemini(num, total, system_prompt, user_msg, max_tok=8000):
    """Llama a Gemini 2.0 Flash con el mismo contrato que _llamar_claude."""
    print(f"[Gemini {num}/{total}] Iniciando...")
    t0 = time.time()
    model = 'gemini-2.5-flash-lite'
    url = f'https://generativelanguage.googleapis.com/v1beta/models/{model}:generateContent?key={GEMINI_KEY}'
    payload = {
        'system_instruction': {'parts': [{'text': system_prompt}]},
        'contents': [{'role': 'user', 'parts': [{'text': user_msg}]}],
        'generationConfig': {
            'maxOutputTokens': max_tok,
            'temperature': 0.7,
        }
    }
    try:
        resp = req.post(url, json=payload, timeout=300)
        elapsed = round(time.time() - t0, 1)
        if resp.status_code != 200:
            print(f"[Gemini {num}/{total}] ERROR {resp.status_code}: {resp.text[:300]}")
            return None, f'Error API Gemini ({resp.status_code}): {resp.text[:200]}'
        rj = resp.json()
        txt = rj['candidates'][0]['content']['parts'][0]['text']
        print(f"[Gemini {num}/{total}] OK — {elapsed}s")
        txt = re.sub(r'^```json\s*', '', txt.strip())
        txt = re.sub(r'^```\s*', '', txt)
        txt = re.sub(r'```\s*$', '', txt).strip()
        try:
            result = json.loads(txt)
            print(f"[Gemini {num}/{total}] JSON OK — claves: {list(result.keys())}")
            return result, None
        except Exception as e:
            print(f"[Gemini {num}/{total}] JSON invalido: {e} | inicio: {txt[:300]}")
            return None, f'JSON invalido en llamada Gemini {num}: {str(e)[:150]}'
    except Exception as e:
        return None, f'Error llamando Gemini: {str(e)[:200]}'



def _llamar_groq(num, total, system_prompt, user_msg, max_tok=8000):
    """Llama a Groq (Llama 3.3 70B) con el mismo contrato que _llamar_claude."""
    print(f"[Groq {num}/{total}] Iniciando...")
    t0 = time.time()
    url = 'https://api.groq.com/openai/v1/chat/completions'

    # Prefijo que fuerza respuestas detalladas y completas
    prefijo = (
        "INSTRUCCION CRITICA DE CALIDAD: Eres un experto medico y de salud integral. "
        "Debes ser EXTREMADAMENTE detallado, especifico y personalizado. "
        "PROHIBIDO usar frases genericas o vagas. "
        "Cada campo del JSON debe tener minimo 2-3 oraciones ricas en contenido clinico. "
        "El menu semanal debe tener comidas COMPLETAS y VARIADAS para cada dia. "
        "Los protocolos deben tener pasos numerados con tiempos y cantidades exactas. "
        "Las recomendaciones deben mencionar el nombre del paciente y sus condiciones especificas. "
        "Completa ABSOLUTAMENTE TODOS los campos sin omitir ninguno. "
        "La calidad de tu respuesta es lo mas importante — tomaté el tiempo necesario.\n\n"
    )

    payload = {
        'model': 'llama-3.3-70b-versatile',
        'messages': [
            {'role': 'system', 'content': prefijo + system_prompt},
            {'role': 'user',   'content': user_msg}
        ],
        'max_tokens': max_tok,
        'temperature': 0.4,  # más bajo = más preciso y consistente
    }
    try:
        resp = req.post(
            url,
            headers={'Authorization': f'Bearer {GROQ_KEY}', 'Content-Type': 'application/json'},
            json=payload,
            timeout=300
        )
        elapsed = round(time.time() - t0, 1)
        if resp.status_code != 200:
            print(f"[Groq {num}/{total}] ERROR {resp.status_code}: {resp.text[:300]}")
            return None, f'Error API Groq ({resp.status_code}): {resp.text[:200]}'
        rj = resp.json()
        txt = rj['choices'][0]['message']['content']
        print(f"[Groq {num}/{total}] OK — {elapsed}s")
        txt = re.sub(r'^```json\s*', '', txt.strip())
        txt = re.sub(r'^```\s*', '', txt)
        txt = re.sub(r'```\s*$', '', txt).strip()
        try:
            result = json.loads(txt)
            print(f"[Groq {num}/{total}] JSON OK — claves: {list(result.keys())}")
            return result, None
        except Exception as e:
            print(f"[Groq {num}/{total}] JSON invalido: {e} | inicio: {txt[:300]}")
            return None, f'JSON invalido en llamada Groq {num}: {str(e)[:150]}'
    except Exception as e:
        return None, f'Error llamando Groq: {str(e)[:200]}'



# ════════════════════════════════════════════════════════════
# CATÁLOGO DE TRATAMIENTOS — editable desde /catalogo
# ════════════════════════════════════════════════════════════

def _ruta_catalogo():
    return os.path.join(os.path.dirname(__file__), 'catalogo_tratamientos.json')


def _descargar_catalogo_cloudinary():
    """Descarga el catálogo desde Cloudinary si existe. Devuelve dict o None."""
    if not CLOUDINARY_CLOUD_NAME:
        return None
    url = f'https://res.cloudinary.com/{CLOUDINARY_CLOUD_NAME}/raw/upload/carvajal/catalogo_tratamientos.json'
    try:
        r = req.get(url, timeout=15)
        if r.status_code == 200:
            return r.json()
        print(f'[catalogo] Cloudinary status {r.status_code}')
    except Exception as e:
        print(f'[catalogo] Error descargando: {e}')
    return None


def _cargar_catalogo():
    """Carga el catálogo: primero local, si falla descarga de Cloudinary."""
    ruta = _ruta_catalogo()
    try:
        if os.path.exists(ruta):
            with open(ruta, 'r', encoding='utf-8') as f:
                return json.load(f)
    except Exception as e:
        print(f'[catalogo] Error leyendo local: {e}')
    remoto = _descargar_catalogo_cloudinary()
    if remoto:
        try:
            with open(ruta, 'w', encoding='utf-8') as f:
                json.dump(remoto, f, ensure_ascii=False, indent=2)
        except Exception as e:
            print(f'[catalogo] Error guardando local remoto: {e}')
        return remoto
    # Fallback vacío con estructura válida
    return {'version': '1', 'ultima_actualizacion': '', 'notas': '', 'tratamientos': []}


def _guardar_catalogo(data):
    """Guarda catálogo localmente y lo sube a Cloudinary.
    Acepta el objeto completo o una lista de tratamientos."""
    ruta = _ruta_catalogo()
    if isinstance(data, list):
        data = {'version': '1', 'ultima_actualizacion': '', 'notas': '', 'tratamientos': data}
    if 'version' not in data:
        data['version'] = '1'
    if 'tratamientos' not in data:
        data['tratamientos'] = []
    with open(ruta, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    if CLOUDINARY_CLOUD_NAME:
        try:
            resultado = cloudinary.uploader.upload(
                ruta,
                folder='carvajal',
                public_id='catalogo_tratamientos',
                resource_type='raw',
                overwrite=True,
                invalidate=True,
            )
            print(f'[catalogo] Subido a Cloudinary: {resultado.get("secure_url", "")[:80]}')
        except Exception as e:
            print(f'[catalogo] Error subiendo a Cloudinary: {e}')


def _catalogo_a_texto(catalogo):
    """Convierte el JSON del catálogo al texto que espera la IA."""
    lineas = [
        'BASE DE CONOCIMIENTO CLINICO - TRATAMIENTOS CENTRO CARVAJAL',
        'Usa UNICAMENTE estos tratamientos con precios y datos clinicos completos.\n'
    ]
    for t in catalogo.get('tratamientos', []):
        if not t.get('activo', True):
            continue
        precios = ' / '.join(
            f"{p['nombre']} ${p['valor']}" for p in t.get('precios', [])
            if p.get('nombre') and p.get('valor') is not None
        ) or 'Consultar'
        combinar = ', '.join(t.get('combinar_con', [])) or '—'
        no_combinar = ', '.join(t.get('no_combinar', [])) or '—'
        lineas.append(
            f"[{t['nombre']}] Precios: {precios} | "
            f"Problemas: {','.join(t.get('problemas', []))} | "
            f"Zonas: {','.join(t.get('zonas', []))} | "
            f"Grado: {','.join(t.get('grado', []))} | "
            f"Sesiones: {t.get('sesiones','')} | "
            f"Intervalo: {t.get('intervalo','')} | "
            f"Recovery: {t.get('recovery','')} | "
            f"Combinar: {combinar} | "
            f"NO combinar: {no_combinar} | "
            f"Orden: {t.get('orden','')} | "
            f"CONTRAINDICACIONES: {','.join(t.get('contraindicaciones', []))}"
        )
    return '\n'.join(lineas)


def _sincronizar_catalogo_inicial():
    """Al arrancar, asegura que el JSON local sea la última versión disponible."""
    print('[catalogo] Sincronizando catálogo inicial...')
    remoto = _descargar_catalogo_cloudinary()
    if remoto:
        try:
            with open(_ruta_catalogo(), 'w', encoding='utf-8') as f:
                json.dump(remoto, f, ensure_ascii=False, indent=2)
            print('[catalogo] Catálogo sincronizado desde Cloudinary')
            return
        except Exception as e:
            print(f'[catalogo] Error guardando catálogo remoto: {e}')
    print('[catalogo] Usando catálogo local')



def generar_plan_ia(d, job_id=None, modelo='claude', session_id=None):
    datos = _datos_paciente(d)
    t_total = time.time()

    def actualizar(msg, pct=None):
        if job_id and job_id in jobs:
            jobs[job_id]['msg'] = msg
            if pct is not None:
                jobs[job_id]['pct'] = pct

    def guardar_progreso(parcial):
        """Guarda el JSON parcial acumulado en Cloudinary si hay session_id."""
        if session_id and CLOUDINARY_CLOUD_NAME:
            try:
                guardar_sesion_cloudinary(session_id, {**d, '__plan_parcial__': parcial})
                print(f'[generar_plan_ia] Progreso guardado: {list(parcial.keys())}')
            except Exception as e:
                print(f'[generar_plan_ia] Error guardando progreso: {e}')


    SYS1 = '''Eres el generador de contenido para planes del METODO CARVAJAL.
Devuelve UNICAMENTE JSON valido sin explicaciones ni markdown.
Genera SOLO estas 3 claves: portada, diagnostico, rutina.
{"portada":{"titulo_pilares":"Los 5 Pilares - [condicion]","intro":"3-4 lineas calido motivador","pilares_resumen":[{"num":1,"emoji":"\ud83e\udd57","titulo":"Titulo adaptado","descripcion":"2-3 lineas"},{"num":2,"emoji":"\ud83c\udfc3","titulo":"Titulo","descripcion":"2-3 lineas"},{"num":3,"emoji":"\ud83e\udde0","titulo":"Titulo","descripcion":"2-3 lineas"},{"num":4,"emoji":"\ud83d\ude34","titulo":"Titulo","descripcion":"2-3 lineas"},{"num":5,"emoji":"\u2728","titulo":"Titulo","descripcion":"2-3 lineas"}]},"diagnostico":{"nota_medica":"Nota alertas criticas","filas":[{"area":"Antropometria","estado":"datos","hallazgos":"analisis","alerta":"normal"},{"area":"Salud Digestiva","estado":"...","hallazgos":"...","alerta":"normal"},{"area":"Sueno y Energia","estado":"...","hallazgos":"...","alerta":"normal"},{"area":"Evaluacion Cutanea","estado":"...","hallazgos":"...","alerta":"normal"},{"area":"Salud Capilar","estado":"...","hallazgos":"...","alerta":"normal"},{"area":"Prioridad Principal","estado":"palabras textuales","hallazgos":"...","alerta":"normal"},{"area":"Condiciones Medicas","estado":"...","hallazgos":"...","alerta":"normal"},{"area":"Rutina Facial Actual","estado":"...","hallazgos":"...","alerta":"normal"},{"area":"Estilo de Vida","estado":"...","hallazgos":"...","alerta":"normal"}]},"rutina":{"nota":"nota rutina","items":[{"hora":"07:00","actividad":"descripcion","pilar":"Nutricion"}]}}
REGLA: Maximo 8 items en rutina. Tips hiperspecificos con nombre y profesion.'''

    SYS2 = '''Eres el generador de contenido para planes del METODO CARVAJAL.
Devuelve UNICAMENTE JSON valido sin explicaciones ni markdown.
Genera SOLO estas 3 claves: pilar1, pilar2, pilar3.
{"pilar1":{"titulo":"Nutricion adaptada","objetivo":"2-3 lineas","frase_motivacional":"frase corta","frase_posicion":"inicio","permitidos":["item"],"evitar":["item"],"menu":[{"dia":"Lunes","desayuno":"...","almuerzo":"...","cena":"...","snack":"..."},{"dia":"Martes","desayuno":"...","almuerzo":"...","cena":"...","snack":"..."},{"dia":"Miercoles","desayuno":"...","almuerzo":"...","cena":"...","snack":"..."},{"dia":"Jueves","desayuno":"...","almuerzo":"...","cena":"...","snack":"..."},{"dia":"Viernes","desayuno":"...","almuerzo":"...","cena":"...","snack":"..."},{"dia":"Sabado","desayuno":"...","almuerzo":"...","cena":"...","snack":"..."},{"dia":"Domingo","desayuno":"...","almuerzo":"...","cena":"...","snack":"..."}],"compras":[{"categoria":"Proteinas","emoji":"\ud83e\udd69","items":["i1","i2","i3","i4","i5"]},{"categoria":"Carbohidratos","emoji":"\ud83c\udf3e","items":["i1","i2","i3","i4"]},{"categoria":"Vegetales","emoji":"\ud83e\udd66","items":["i1","i2","i3","i4","i5"]},{"categoria":"Frutas","emoji":"\ud83c\udf4e","items":["i1","i2","i3","i4"]},{"categoria":"Grasas","emoji":"\ud83e\udd51","items":["i1","i2","i3"]},{"categoria":"Otros","emoji":"\ud83e\uddf4","items":["i1","i2","i3","i4"]}],"suplementacion":["Sup1: dosis"],"tips":[{"texto":"tip especifico con nombre"}]},"pilar2":{"titulo":"Actividad Fisica","objetivo":"objetivo","frase_motivacional":"frase","frase_posicion":"medio","plan_semanal":"plan dia a dia","adaptaciones":"adaptaciones","tips":[{"texto":"tip"}]},"pilar3":{"titulo":"Bienestar Mental","objetivo":"objetivo","frase_motivacional":"frase","frase_posicion":"final","tecnicas":["t1","t2","t3","t4","t5"],"tips":[{"texto":"tip"}]}}
REGLAS: Respetar intolerancias. Tips con nombre, profesion, horario real.'''

    catalogo_json = _cargar_catalogo()
    CATALOGO = _catalogo_a_texto(catalogo_json)

    if modelo == 'gemini':
        _llamar = _llamar_gemini
        tok1, tok2, tok3 = 6000, 12000, 10000
        nombre_modelo = 'Gemini'
    elif modelo == 'groq':
        _llamar = _llamar_groq
        tok1, tok2, tok3 = 8000, 12000, 10000  # más tokens = más detalle
        nombre_modelo = 'Groq (Llama)'
    else:
        _llamar = _llamar_claude
        tok1, tok2, tok3 = 8000, 12000, 10000
        nombre_modelo = 'Claude'

    actualizar(f'Sección 1/3 — Portada, diagnóstico y rutina diaria... ({nombre_modelo})', 15)
    r1, err = _llamar(1, 3, SYS1, datos, max_tok=tok1)
    if err: return {'error': err}
    guardar_progreso(r1)  # ← sesión guardada tras paso 1

    if modelo == 'claude':
        time.sleep(2)  # pausa entre secciones para evitar overloaded (529)
    elif modelo in ('gemini', 'groq'):
        time.sleep(5)  # pausa para evitar rate limit en APIs con cuota baja

    actualizar(f'Sección 2/3 — Nutrición, ejercicio y bienestar mental... ({nombre_modelo})', 45)
    r2, err = _llamar(2, 3, SYS2, datos, max_tok=tok2)
    if err: return {'error': err}
    guardar_progreso({**r1, **r2})  # ← sesión actualizada tras paso 2

    if modelo == 'claude':
        time.sleep(2)
    elif modelo in ('gemini', 'groq'):
        time.sleep(5)

    SYS3 = '''Eres el generador de contenido para planes del METODO CARVAJAL.
Devuelve UNICAMENTE JSON valido sin explicaciones ni markdown.
Genera SOLO estas 3 claves: pilar4, pilar5, compromiso.
{"pilar4":{"titulo":"Optimizacion del Sueno","objetivo":"objetivo personalizado 2-3 lineas","frase_motivacional":"frase corta","frase_posicion":"inicio","protocolo":["paso1","paso2","paso3","paso4","paso5"],"reglas":["regla higiene sueno 1","regla2","regla3"],"tips":[{"texto":"tip especifico con nombre"}]},"pilar5":{"titulo":"Tratamientos Esteticos","objetivo":"objetivo personalizado","frase_motivacional":"frase","frase_posicion":"medio","bimestres":[{"periodo":"Bimestre 1","titulo":"titulo enfoque","tratamientos":[{"nombre":"Nombre tratamiento","sesiones":"N sesiones","inversion":"$XXX","beneficio":"beneficio concreto"}],"total":000},{"periodo":"Bimestre 2","titulo":"titulo","tratamientos":[{"nombre":"...","sesiones":"...","inversion":"$XXX","beneficio":"..."}],"total":000},{"periodo":"Bimestre 3","titulo":"titulo","tratamientos":[{"nombre":"...","sesiones":"...","inversion":"$XXX","beneficio":"..."}],"total":000},{"periodo":"Bimestre 4","titulo":"titulo","tratamientos":[{"nombre":"...","sesiones":"...","inversion":"$XXX","beneficio":"..."}],"total":000},{"periodo":"Bimestre 5","titulo":"titulo","tratamientos":[{"nombre":"...","sesiones":"...","inversion":"$XXX","beneficio":"..."}],"total":000},{"periodo":"Bimestre 6","titulo":"titulo","tratamientos":[{"nombre":"...","sesiones":"...","inversion":"$XXX","beneficio":"..."}],"total":000}],"total_anual":0000,"notas_criticas":["nota importante si aplica"],"rutina_am":[{"paso":1,"producto":"producto AM","descripcion":"como y cuando"}],"rutina_pm":[{"paso":1,"producto":"producto PM","descripcion":"como y cuando"}],"tips":[{"texto":"tip"}]},"compromiso":{"parrafo":"parrafo motivacional de cierre 3-4 lineas","resultados":[{"texto":"resultado esperable 1"},{"texto":"resultado esperable 2"},{"texto":"resultado esperable 3"}],"proximos_pasos":["paso concreto 1","paso concreto 2","paso concreto 3"]}}
REGLAS: Usar UNICAMENTE tratamientos del catalogo provisto. Verificar contraindicaciones. Calcular totales reales. Tips con nombre, profesion, horario real.''' + '\nCATALOGO DE TRATAMIENTOS:\n' + CATALOGO

    actualizar(f'Sección 3/3 — Sueño, tratamientos y plan de compromiso... ({nombre_modelo})', 75)
    r3, err = _llamar(3, 3, SYS3, datos, max_tok=tok3)
    if err: return {'error': err}

    resultado = {}
    resultado.update(r1)
    resultado.update(r2)
    resultado.update(r3)

    t_elapsed = round(time.time() - t_total, 1)
    print(f"[Total] {t_elapsed}s | claves: {list(resultado.keys())}")
    return resultado

# ════════════════════════════════════════════════════════════
# RENDER PLAN — llenar plantilla HTML
# ════════════════════════════════════════════════════════════

PLANTILLA_PLAN_HTML = """<!DOCTYPE html>
<html lang="es">
<head>
<meta charset="UTF-8">
<title>Plan Carvajal · {{NOMBRE}}</title>
<link href="https://fonts.googleapis.com/css2?family=Cormorant+Garamond:ital,wght@0,300;0,400;0,600;1,300;1,400&family=Inter:wght@300;400;500;600;700;800&display=swap" rel="stylesheet">
<style>
:root{--olive:#8fa832;--olive-light:rgba(143,168,50,0.10);--olive-border:rgba(143,168,50,0.22);--dark:#1c1c1c;--gold:#b8935a;--cream:#faf9f6;--green:#2d3a2e;--white:#fff;--gray:#6b7280;--border:rgba(143,168,50,0.18)}
*{margin:0;padding:0;box-sizing:border-box}
body{font-family:"Inter",sans-serif;background:#d8d8d8;color:var(--dark);font-size:10.5pt;line-height:1.6}
h1,h2,h3,h4{font-family:"Cormorant Garamond",serif;line-height:1.2}
strong{font-weight:600}

/* ── TOPBAR ── */
#topbar-cv{position:fixed;top:0;left:0;right:0;z-index:9999;background:var(--green);padding:8px 20px;display:flex;align-items:center;justify-content:space-between;box-shadow:0 2px 12px rgba(0,0,0,0.3)}
#topbar-cv .tb-info strong{color:#fff;font-size:12px;display:block}
#topbar-cv .tb-info span{color:rgba(255,255,255,0.45);font-size:10px}
.tb-btns{display:flex;gap:8px}
.tb-btn{font-family:"Inter",sans-serif;font-size:11px;font-weight:500;cursor:pointer;padding:6px 14px;border-radius:4px;border:none;letter-spacing:.03em;transition:all .18s}
.tb-outline{background:transparent;border:1px solid rgba(255,255,255,.35);color:rgba(255,255,255,.85)}
.tb-outline:hover{border-color:#fff;color:#fff}
.tb-green{background:var(--olive);color:#fff}
.tb-green:hover{background:#7a9428}
.tb-green:disabled{background:#9ca3af;cursor:not-allowed}
#toast-cv{position:fixed;bottom:20px;right:20px;background:var(--green);color:#fff;padding:10px 18px;border-radius:5px;font-size:11px;font-family:"Inter",sans-serif;opacity:0;transition:opacity .3s;z-index:9998;pointer-events:none}
#toast-cv.show{opacity:1}
body{padding-top:44px}
[contenteditable]{outline:none;border-radius:2px;transition:background .15s}
[contenteditable]:hover:not(:focus){background:rgba(143,168,50,0.08)}
[contenteditable]:focus{background:rgba(143,168,50,0.13);box-shadow:0 0 0 1px rgba(143,168,50,0.35)}

@page{size:A4;margin:0}
@media print{
  #topbar-cv,#toast-cv,.no-print{display:none!important}
  body{background:white;padding-top:0}
  .page{box-shadow:none!important;margin:0!important;page-break-after:always}
  .page:last-child{page-break-after:auto}
  .pilar-card,.rutina-row,.suppl-item,.result-item,.chk-list li,.rut-step,.bim-body,.mini-card{page-break-inside:avoid}
}

/* ── PÁGINAS ── */
.page{width:210mm;min-height:297mm;margin:0 auto 8mm;background:white;position:relative;overflow:hidden;box-shadow:0 4px 40px rgba(0,0,0,0.15);display:flex;flex-direction:column}

/* PORTADA */
.cover{width:100%;height:297mm;background:white;display:flex;flex-direction:column;position:relative;overflow:hidden}
.cover-stripe{height:4px;background:var(--olive);width:100%;flex-shrink:0}
.cover-accent{position:absolute;top:0;right:0;width:3px;height:100%;background:linear-gradient(to bottom,var(--olive),transparent)}
.cover-top{padding:22px 44px;display:flex;align-items:center;justify-content:space-between;position:relative;z-index:2;background:#f7f7f7}
.cover-logo{height:72px;width:auto}
.cover-date-block{text-align:right}
.cdb-label{font-size:6.5pt;font-weight:600;text-transform:uppercase;letter-spacing:2.5px;color:var(--olive);margin-bottom:3px}
.cdb-date{font-size:9pt;font-weight:500;color:var(--gray)}
.cover-rule{height:1px;background:var(--olive-border);margin:16px 44px 0;position:relative;z-index:2}
.cover-title-block{padding:26px 44px 0;position:relative;z-index:2}
.cover-eyebrow{font-size:7pt;font-weight:600;text-transform:uppercase;letter-spacing:4px;color:var(--olive);margin-bottom:14px;display:flex;align-items:center;gap:10px}
.cover-eyebrow::after{content:"";flex:1;height:1px;background:var(--olive-border)}
.cover-main-title{font-family:"Inter",sans-serif;line-height:0.95;letter-spacing:-2px;margin-bottom:0}
.cover-main-title .t-light{display:block;font-weight:300;color:var(--gray);font-size:18pt;letter-spacing:-0.5px;margin-bottom:4px}
.cover-main-title .t-bold{display:block;font-weight:800;color:var(--olive);font-size:42pt;letter-spacing:-2.5px;line-height:0.92}
.cover-main-title .t-dark{display:block;font-weight:800;color:var(--dark);font-size:42pt;letter-spacing:-2.5px;line-height:0.92}
.cover-subtitle{font-family:"Inter",sans-serif;font-size:8.5pt;font-weight:400;color:var(--gray);letter-spacing:3px;text-transform:uppercase;margin-top:14px;padding-top:14px;border-top:1px solid var(--olive-border);display:inline-block}
.cover-patient-block{margin:24px 44px 0;padding:18px 22px;background:var(--dark);border-radius:6px;display:flex;align-items:center;gap:18px;position:relative;z-index:2}
.cpb-bar{width:3px;height:48px;background:var(--olive);border-radius:2px;flex-shrink:0}
.cpb-name{font-family:"Cormorant Garamond",serif;font-size:21pt;font-weight:600;color:white;line-height:1.1;margin-bottom:2px}
.cpb-meta{font-size:8pt;color:rgba(255,255,255,0.4)}
.cpb-data{margin-left:auto;display:flex;border-left:1px solid rgba(255,255,255,0.08);padding-left:18px}
.cpb-data-item{text-align:center;padding:0 16px;border-right:1px solid rgba(255,255,255,0.08)}
.cpb-data-item:last-child{border-right:none}
.cpdi-label{font-size:6pt;font-weight:600;text-transform:uppercase;letter-spacing:1.5px;color:var(--olive);margin-bottom:3px}
.cpdi-val{font-size:9.5pt;font-weight:600;color:white}
.cover-letter{margin:20px 44px 0;padding:16px 20px;background:var(--olive-light);border-left:3px solid var(--olive);border-radius:0 6px 6px 0;position:relative;z-index:2}
.cl-text{font-family:"Cormorant Garamond",serif;font-size:10.5pt;font-style:italic;line-height:1.75;color:var(--dark)}
.cover-pilares{margin:18px 44px 0;display:flex;flex-direction:column;gap:8px;position:relative;z-index:2}
.cp-row{display:flex;gap:8px}
.cp-row .cp-item{flex:1}
.cp-wide{flex:3!important}
.cp-narrow{flex:2!important}
.cp-item{background:white;border:1px solid var(--olive-border);border-top:3px solid var(--olive);border-radius:0 0 6px 6px;padding:20px 14px;text-align:center;min-height:90px;display:flex;flex-direction:column;align-items:center;justify-content:center}
.cp-icon{font-size:18pt;margin-bottom:7px}
.cp-num{font-size:6.5pt;color:var(--olive);font-weight:600;text-transform:uppercase;letter-spacing:1px;margin-bottom:4px}
.cp-label{font-size:7.5pt;font-weight:500;color:var(--dark);line-height:1.35}
.cover-spacer{flex:1}
.cover-footer{background:var(--dark);padding:12px 44px;display:flex;align-items:center;justify-content:space-between;position:relative;z-index:2;flex-shrink:0}
.cf-brand{font-size:8pt;font-weight:600;color:var(--olive)}
.cf-sub{font-size:7pt;color:rgba(255,255,255,0.28);margin-top:1px}
.cf-right{font-size:7pt;color:rgba(255,255,255,0.28);text-align:right;line-height:1.7}

/* PÁGINAS INTERNAS */
.page-header{background:var(--dark);padding:10px 28px;display:flex;align-items:center;justify-content:space-between;flex-shrink:0}
.ph-left{display:flex;align-items:center;gap:9px}
.ph-dot{width:6px;height:6px;border-radius:50%;background:var(--olive)}
.ph-title{font-size:8.5pt;color:var(--olive);font-weight:600}
.ph-right{font-size:6.5pt;color:rgba(255,255,255,0.22);text-transform:uppercase;letter-spacing:1px}
.content{padding:20px 28px;flex:1}
.content-sm{padding:14px 28px;flex:1}
.sec-label{font-size:6.5pt;font-weight:600;text-transform:uppercase;letter-spacing:3px;color:var(--olive);margin-bottom:4px}
.sec-title{font-family:"Cormorant Garamond",serif;font-size:19pt;font-weight:600;color:var(--dark);margin-bottom:15px;line-height:1.15}
.sec-title-sm{font-family:"Cormorant Garamond",serif;font-size:13pt;font-weight:600;color:var(--dark);margin-bottom:10px;padding-bottom:5px;border-bottom:1px solid var(--olive-border)}
.nota-medica{background:#fffbee;border-left:3px solid var(--gold);border-radius:0 6px 6px 0;padding:10px 14px;margin-bottom:13px;font-size:7.5pt;color:#7a5a10;line-height:1.6}
.nota-medica strong{color:var(--gold);display:block;margin-bottom:2px;font-size:7pt;text-transform:uppercase;letter-spacing:1px}
.diag-table{width:100%;border-collapse:collapse;font-size:7.5pt}
.diag-table th{background:var(--dark);color:var(--olive);padding:7px 11px;text-align:left;font-size:6.5pt;text-transform:uppercase;letter-spacing:1px;font-weight:600}
.diag-table td{padding:7px 11px;border-bottom:1px solid rgba(143,168,50,0.09);vertical-align:top;line-height:1.5}
.diag-table tr:nth-child(even) td{background:rgba(143,168,50,0.025)}
.diag-left{width:105px;font-weight:600;font-size:7.5pt;color:var(--dark);background:rgba(45,58,46,0.04)!important}
.badge-w{display:inline-block;background:#fef3cd;color:#856404;font-size:6.5pt;font-weight:600;padding:1px 5px;border-radius:6px}
.badge-c{display:inline-block;background:#f8d7da;color:#721c24;font-size:6.5pt;font-weight:600;padding:1px 5px;border-radius:6px}
.diag-val-strong{font-weight:600;font-size:8pt;color:var(--dark);display:block;margin-bottom:1px}
.diag-val-sub{font-size:7pt;color:var(--gray);line-height:1.5}
.pilar-card{display:flex;gap:11px;align-items:flex-start;padding:10px 13px;background:white;border:1px solid var(--olive-border);border-left:3px solid var(--olive);border-radius:0 6px 6px 0;margin-bottom:7px;page-break-inside:avoid}
.pilar-icon{font-size:14pt;flex-shrink:0;margin-top:2px}
.pilar-title{font-weight:600;font-size:8.5pt;color:var(--dark);margin-bottom:2px}
.pilar-desc{font-size:7.5pt;color:var(--gray);line-height:1.5}
.rutina-row{display:flex;align-items:flex-start;gap:11px;padding:7px 0;border-bottom:1px solid rgba(143,168,50,0.07);page-break-inside:avoid}
.rutina-hora{font-weight:700;color:var(--olive);min-width:40px;font-size:8.5pt;flex-shrink:0}
.rutina-text{flex:1;line-height:1.5;color:var(--dark);font-size:8pt}
.rtag{font-size:6pt;font-weight:600;padding:1px 7px;border-radius:3px;white-space:nowrap;flex-shrink:0;margin-top:2px;display:inline-block}
.rtag-n{background:#e8f5e9;color:#2e7d32}
.rtag-a{background:#e3f2fd;color:#1565c0}
.rtag-s{background:#f3e5f5;color:#4a148c}
.rtag-e{background:#fff8e1;color:#f57f17}
.rtag-h{background:#e8f5e9;color:#1b5e20}
.rtag-m{background:#f3e8ff;color:#4a148c}
.perm-evit{display:grid;grid-template-columns:1fr 1fr;gap:11px;margin-bottom:12px}
.pe-box{background:white;border:1px solid var(--border);border-radius:6px;padding:11px 13px}
.pe-title{font-size:7pt;font-weight:600;text-transform:uppercase;letter-spacing:1.5px;margin-bottom:7px;padding-bottom:5px;border-bottom:1px solid var(--border)}
.pe-title.p{color:#2e7d32}.pe-title.e{color:#c62828}
.pe-item{font-size:7.5pt;color:var(--dark);padding:2px 0;display:flex;align-items:flex-start;gap:6px;line-height:1.4}
.pe-dot{width:5px;height:5px;border-radius:50%;flex-shrink:0;margin-top:4px}
.pe-dot.p{background:#2e7d32}.pe-dot.e{background:#c62828}
.menu-table{width:100%;border-collapse:collapse;font-size:7pt}
.menu-table th{background:var(--olive);color:white;padding:5px 7px;text-align:left;font-size:6pt;text-transform:uppercase;letter-spacing:0.8px;font-weight:600}
.menu-table td{padding:5px 7px;border-bottom:1px solid rgba(143,168,50,0.09);vertical-align:top;line-height:1.4}
.menu-table .dia{font-weight:700;color:var(--olive);background:rgba(143,168,50,0.05)!important}
.suppl-item{display:flex;align-items:flex-start;gap:9px;padding:6px 0;border-bottom:1px solid var(--border);page-break-inside:avoid}
.suppl-item:last-child{border-bottom:none}
.suppl-bullet{width:18px;height:18px;border-radius:50%;background:var(--olive-light);border:1.5px solid var(--olive);display:flex;align-items:center;justify-content:center;font-size:7pt;font-weight:700;color:var(--olive);flex-shrink:0;margin-top:1px}
.suppl-name{font-weight:600;color:var(--dark);font-size:8pt}
.suppl-desc{font-size:7pt;color:var(--gray)}
.bim-header{background:var(--dark);color:var(--olive);padding:8px 13px;font-size:9pt;font-weight:600;border-radius:5px 5px 0 0;display:flex;align-items:center;justify-content:space-between}
.bim-body{border-radius:0 0 5px 5px;overflow:hidden;margin-bottom:9px;border:1px solid var(--border);border-top:none;page-break-inside:avoid}
.bim-table{width:100%;border-collapse:collapse;font-size:7pt}
.bim-table th{background:var(--olive-light);color:var(--olive);padding:5px 9px;text-align:left;font-size:6pt;text-transform:uppercase;letter-spacing:0.8px;font-weight:700}
.bim-table td{padding:6px 9px;border-bottom:1px solid var(--border);vertical-align:top;line-height:1.4}
.bim-table tr:last-child td{border-bottom:none}
.bim-total{background:var(--olive-light);padding:7px 13px;text-align:right;font-weight:700;color:var(--olive);font-size:8pt;border-top:1px solid var(--olive-border)}
.total-anual-box{background:var(--dark);border-radius:6px;padding:20px;text-align:center;margin:14px 0}
.ta-label{font-size:6.5pt;letter-spacing:3px;text-transform:uppercase;color:rgba(143,168,50,0.7);margin-bottom:5px}
.ta-amount{font-family:"Cormorant Garamond",serif;font-size:34pt;font-weight:600;color:var(--olive)}
.ta-sub{font-size:7pt;color:rgba(255,255,255,0.25);margin-top:3px}
.rut-grid{display:grid;grid-template-columns:1fr 1fr;gap:11px;margin-bottom:11px}
.rut-box{background:white;border:1px solid var(--border);border-radius:6px;overflow:hidden}
.rut-header{padding:8px 13px;font-size:8.5pt;font-weight:600;display:flex;align-items:center;gap:6px}
.rut-header.am{background:#fff8e1;color:#e65100;border-bottom:1px solid rgba(245,127,23,0.12)}
.rut-header.pm{background:#e8eaf6;color:#283593;border-bottom:1px solid rgba(57,73,171,0.12)}
.rut-step{display:flex;align-items:flex-start;gap:8px;padding:6px 11px;border-bottom:1px solid rgba(0,0,0,0.04);font-size:7.5pt;page-break-inside:avoid}
.rut-step:last-child{border-bottom:none}
.rut-num{width:17px;height:17px;border-radius:50%;background:var(--olive-light);color:var(--olive);display:flex;align-items:center;justify-content:center;font-size:6.5pt;font-weight:700;flex-shrink:0;margin-top:1px}
.rut-prod{font-weight:600;color:var(--dark);display:block;margin-bottom:1px}
.rut-desc{font-size:7pt;color:var(--gray)}
.cards-2{display:grid;grid-template-columns:1fr 1fr;gap:10px;margin-bottom:11px}
.cards-3{display:grid;grid-template-columns:repeat(3,1fr);gap:9px;margin-bottom:11px}
.mini-card{background:white;border:1px solid var(--border);border-radius:6px;padding:11px;page-break-inside:avoid}
.mc-icon{font-size:14pt;margin-bottom:5px}
.mc-title{font-size:8pt;font-weight:600;color:var(--dark);margin-bottom:3px}
.mc-text{font-size:7.5pt;color:var(--gray);line-height:1.5}
.objetivo-box{background:var(--green);border-radius:6px;padding:13px 17px;margin-bottom:12px}
.obj-label{font-size:6.5pt;font-weight:600;text-transform:uppercase;letter-spacing:2px;color:var(--olive);margin-bottom:4px}
.obj-text{font-size:8pt;color:rgba(255,255,255,0.85);line-height:1.65}
.quote-box{background:var(--olive-light);border-left:3px solid var(--olive);border-radius:0 6px 6px 0;padding:12px 16px;margin:11px 0}
.quote-text{font-family:"Cormorant Garamond",serif;font-style:italic;font-size:10pt;color:var(--olive);line-height:1.6}
.comp-box{background:var(--green);border-radius:6px;padding:17px 21px;margin-bottom:13px}
.comp-text{font-size:8pt;color:rgba(255,255,255,0.85);line-height:1.75}
.result-grid{display:grid;grid-template-columns:1fr 1fr;gap:7px;margin-bottom:13px}
.result-item{display:flex;align-items:flex-start;gap:8px;padding:8px 11px;background:white;border:1px solid var(--border);border-radius:5px;font-size:7.5pt;line-height:1.4;page-break-inside:avoid}
.result-check{width:17px;height:17px;background:var(--olive);border-radius:50%;display:flex;align-items:center;justify-content:center;flex-shrink:0;font-size:7.5pt;color:white;font-weight:700}
.chk-list{list-style:none}
.chk-list li{display:flex;align-items:flex-start;gap:9px;padding:8px 0;border-bottom:1px solid rgba(143,168,50,0.08);font-size:7.5pt;line-height:1.5;page-break-inside:avoid}
.chk-list li:last-child{border-bottom:none}
.chk-ico{width:17px;height:17px;border:1.5px solid var(--olive);border-radius:3px;display:flex;align-items:center;justify-content:center;flex-shrink:0;margin-top:1px}
.chk-ico svg{width:9px;height:9px;stroke:var(--olive);fill:none;stroke-width:2.5}
.divider-olive{height:2px;background:var(--olive);margin:12px 0;border-radius:1px;width:36px}
.page-footer{background:var(--dark);padding:7px 28px;display:flex;align-items:center;justify-content:space-between;font-size:6.5pt;color:rgba(255,255,255,0.25);flex-shrink:0}
.pf-brand{color:var(--olive);font-weight:600}
</style>
</head>
<body>

<!-- ══ TOPBAR (no se imprime) ══ -->
<div id="topbar-cv" class="no-print">
  <div class="tb-info">
    <strong>{{NOMBRE}}</strong>
    <span>Plan editable · {{FECHA}}</span>
  </div>
  <div class="tb-btns">
    <button class="tb-btn tb-outline" onclick="window.print()">🖨 Imprimir / PDF</button>
    <button class="tb-btn tb-green" id="btn-guardar" onclick="guardarCambios()">💾 Guardar cambios</button>
  </div>
</div>
<div id="toast-cv"></div>

<!-- ══ PORTADA ══ -->
<div class="page">
<div class="cover">
  <div class="cover-stripe"></div>
  <div class="cover-accent"></div>
  <div class="cover-top">
    <img src="https://centrocarvajal.com/wp-content/uploads/2023/05/logo-carvajal-300x147.jpg" class="cover-logo" alt="Centro Carvajal">
    <div class="cover-date-block">
      <div class="cdb-label">Fecha del plan</div>
      <div class="cdb-date" contenteditable="true">{{FECHA}}</div>
    </div>
  </div>
  <div class="cover-rule"></div>
  <div class="cover-title-block">
    <div class="cover-eyebrow">Programa de Salud, Bienestar y Desarrollo de Imagen</div>
    <div class="cover-main-title">
      <span class="t-light">Método de</span>
      <span class="t-bold">Rejuvenecimiento</span>
      <span class="t-dark">Carvajal</span>
    </div>
    <div class="cover-subtitle">Plan Integral Personalizado</div>
  </div>
  <div class="cover-patient-block">
    <div class="cpb-bar"></div>
    <div>
      <div class="cpb-name" contenteditable="true">{{NOMBRE}}</div>
      <div class="cpb-meta" contenteditable="true">{{EDAD}} años · {{OCUPACION}}</div>
    </div>
    <div class="cpb-data">
      <div class="cpb-data-item"><div class="cpdi-label">IMC</div><div class="cpdi-val" contenteditable="true">{{IMC}}</div></div>
      <div class="cpb-data-item"><div class="cpdi-label">Satisfacción</div><div class="cpdi-val" contenteditable="true">{{SATISFACCION}}/10</div></div>
      <div class="cpb-data-item"><div class="cpdi-label">Condición</div><div class="cpdi-val" contenteditable="true">{{CONDICION_CORTA}}</div></div>
    </div>
  </div>
  <div class="cover-letter">
    <div class="cl-text" contenteditable="true">{{INTRO}}</div>
  </div>
  <div class="cover-pilares">
    {{PILARES_PORTADA}}
  </div>
  <div class="cover-spacer"></div>
  <div class="cover-footer">
    <div>
      <div class="cf-brand">Centro Carvajal · Líderes en Medicina Estética en Panamá</div>
      <div class="cf-sub">centrocarvajal.com · Tel: 263-8134 · 209-4284 · @centrocarvajal</div>
    </div>
    <div class="cf-right">Revisado por el equipo médico</div>
  </div>
</div>
</div>

<!-- ══ PÁG 2: DIAGNÓSTICO ══ -->
<div class="page">
  <div class="page-header">
    <div class="ph-left"><div class="ph-dot"></div><div class="ph-title">Diagnóstico y Análisis Inicial</div></div>
    <div class="ph-right">Método de Rejuvenecimiento Carvajal · {{NOMBRE}}</div>
  </div>
  <div class="content">
    <div class="sec-label">Sección 01</div>
    <div class="sec-title">Diagnóstico Integral</div>
    <div style="border-radius:6px;overflow:hidden;border:1px solid var(--border)">
      <table class="diag-table">
        <thead><tr><th style="width:105px">Área</th><th>Hallazgos Clave</th></tr></thead>
        <tbody>{{DIAGNOSTICO_FILAS}}</tbody>
      </table>
    </div>
    <div style="margin-top:14px">
      <div class="sec-label" style="margin-bottom:7px">Los 5 Pilares del Plan</div>
      {{PILARES_CARDS}}
    </div>
  </div>
  <div class="page-footer"><span><span class="pf-brand">Centro Carvajal · Líderes en Medicina Estética en Panamá</span></span><span>Página 2</span></div>
</div>

<!-- ══ PÁG 3: RUTINA + P1 NUTRICIÓN ══ -->
<div class="page">
  <div class="page-header">
    <div class="ph-left"><div class="ph-dot"></div><div class="ph-title">Rutina Diaria · Pilar 1: Nutrición</div></div>
    <div class="ph-right">Método de Rejuvenecimiento Carvajal · {{NOMBRE}}</div>
  </div>
  <div class="content">
    <div class="sec-label">Sección 02</div>
    <div class="sec-title">Rutina Diaria Ideal</div>
    {{RUTINA_FILAS}}
    <div style="margin-top:16px">
      <div class="divider-olive"></div>
      <div class="sec-label" style="margin-bottom:7px">Pilar 1 · Sección 03</div>
      <div class="sec-title">{{P1_TITULO}}</div>
      <div class="objetivo-box">
        <div class="obj-label">Objetivo Focal</div>
        <div class="obj-text" contenteditable="true">{{P1_OBJETIVO}}</div>
      </div>
      <div class="perm-evit">
        <div class="pe-box">
          <div class="pe-title p">✓ Alimentos Recomendados</div>
          {{P1_PERMITIDOS}}
        </div>
        <div class="pe-box">
          <div class="pe-title e">✕ Limitar o Evitar</div>
          {{P1_EVITAR}}
        </div>
      </div>
    </div>
  </div>
  <div class="page-footer"><span><span class="pf-brand">Centro Carvajal · Líderes en Medicina Estética en Panamá</span></span><span>Página 3</span></div>
</div>

<!-- ══ PÁG 4: MENÚ + SUPL + P2 ACTIVIDAD ══ -->
<div class="page">
  <div class="page-header">
    <div class="ph-left"><div class="ph-dot"></div><div class="ph-title">Menú Semanal · Suplementación · Actividad Física</div></div>
    <div class="ph-right">Método de Rejuvenecimiento Carvajal · {{NOMBRE}}</div>
  </div>
  <div class="content-sm">
    <div class="sec-label">Pilar 1 · Continuación</div>
    <div class="sec-title" style="font-size:15pt;margin-bottom:9px">Menú Semanal</div>
    <div style="border-radius:6px;overflow:hidden;border:1px solid var(--border)">
      <table class="menu-table">
        <thead><tr><th style="width:52px">Día</th><th>Desayuno</th><th>Almuerzo</th><th>Cena</th><th>Snack</th></tr></thead>
        <tbody>{{P1_MENU}}</tbody>
      </table>
    </div>
    <div style="margin-top:12px">
      <div class="sec-label" style="margin-bottom:6px">Suplementación Recomendada</div>
      <div style="background:white;border:1px solid var(--border);border-radius:6px;padding:11px">
        {{P1_SUPLEMENTACION}}
      </div>
    </div>
    <div style="margin-top:12px">
      <div class="divider-olive"></div>
      <div class="sec-label" style="margin-bottom:5px">Pilar 2 · Sección 04</div>
      <div class="sec-title" style="font-size:14pt;margin-bottom:9px">{{P2_TITULO}}</div>
      <div class="objetivo-box" style="padding:11px 15px;margin-bottom:9px">
        <div class="obj-label">Objetivo Focal</div>
        <div class="obj-text" style="font-size:7.5pt" contenteditable="true">{{P2_OBJETIVO}}</div>
      </div>
      <div class="cards-2">
        <div class="mini-card">
          <div class="mc-title">Plan Semanal</div>
          <div class="mc-text" contenteditable="true">{{P2_PLAN}}</div>
        </div>
        <div class="mini-card">
          <div class="mc-title">Adaptaciones Específicas</div>
          <div class="mc-text" contenteditable="true">{{P2_ADAPTACIONES}}</div>
        </div>
      </div>
    </div>
  </div>
  <div class="page-footer"><span><span class="pf-brand">Centro Carvajal · Líderes en Medicina Estética en Panamá</span></span><span>Página 4</span></div>
</div>

<!-- ══ PÁG 5: P3 BIENESTAR + P4 SUEÑO ══ -->
<div class="page">
  <div class="page-header">
    <div class="ph-left"><div class="ph-dot"></div><div class="ph-title">Pilar 3: Bienestar Mental · Pilar 4: Sueño</div></div>
    <div class="ph-right">Método de Rejuvenecimiento Carvajal · {{NOMBRE}}</div>
  </div>
  <div class="content">
    <div class="sec-label">Pilar 3 · Sección 05</div>
    <div class="sec-title">{{P3_TITULO}}</div>
    <div class="objetivo-box">
      <div class="obj-label">Objetivo Focal</div>
      <div class="obj-text" contenteditable="true">{{P3_OBJETIVO}}</div>
    </div>
    <div class="cards-3">{{P3_TECNICAS_CARDS}}</div>
    <div class="quote-box">
      <div class="quote-text" contenteditable="true">{{P3_FRASE}}</div>
    </div>
    <div style="margin-top:14px">
      <div class="divider-olive"></div>
      <div class="sec-label" style="margin-bottom:5px">Pilar 4 · Sección 06</div>
      <div class="sec-title">{{P4_TITULO}}</div>
      <div class="objetivo-box" style="margin-bottom:12px">
        <div class="obj-label">Objetivo</div>
        <div class="obj-text" contenteditable="true">{{P4_OBJETIVO}}</div>
      </div>
      <div class="cards-2">
        <div class="mini-card">
          <div class="mc-title">Protocolo Nocturno</div>
          <div class="mc-text" contenteditable="true">{{P4_PROTOCOLO_TEXT}}</div>
        </div>
        <div class="mini-card">
          <div class="mc-title">Reglas Clave</div>
          <div class="mc-text" contenteditable="true">{{P4_REGLAS_TEXT}}</div>
        </div>
      </div>
    </div>
  </div>
  <div class="page-footer"><span><span class="pf-brand">Centro Carvajal · Líderes en Medicina Estética en Panamá</span></span><span>Página 5</span></div>
</div>

<!-- ══ PÁG 6: TRATAMIENTOS BIM 1-4 ══ -->
<div class="page">
  <div class="page-header">
    <div class="ph-left"><div class="ph-dot"></div><div class="ph-title">Pilar 5 · Tratamientos — Bimestres 1 al 4</div></div>
    <div class="ph-right">Método de Rejuvenecimiento Carvajal · {{NOMBRE}}</div>
  </div>
  <div class="content-sm">
    <div class="sec-label">Pilar 5 · Sección 07</div>
    <div class="sec-title" style="font-size:15pt">{{P5_TITULO}}</div>
    <div class="objetivo-box" style="padding:11px 15px;margin-bottom:11px">
      <div class="obj-label">Objetivo</div>
      <div class="obj-text" style="font-size:7.5pt" contenteditable="true">{{P5_OBJETIVO}}</div>
    </div>
    {{P5_BIMESTRES_A}}
  </div>
  <div class="page-footer"><span><span class="pf-brand">Centro Carvajal · Líderes en Medicina Estética en Panamá</span></span><span>Página 6</span></div>
</div>

<!-- ══ PÁG 7: TRATAMIENTOS 5-6 + RUTINA FACIAL ══ -->
<div class="page">
  <div class="page-header">
    <div class="ph-left"><div class="ph-dot"></div><div class="ph-title">Tratamientos Bim. 5-6 · Total · Rutina Facial</div></div>
    <div class="ph-right">Método de Rejuvenecimiento Carvajal · {{NOMBRE}}</div>
  </div>
  <div class="content-sm">
    {{P5_BIMESTRES_B}}
    <div class="total-anual-box">
      <div class="ta-label">Inversión Total del Plan</div>
      <div class="ta-amount">{{P5_TOTAL_ANUAL}}</div>
      <div class="ta-sub">12 meses · 6 bimestres · Plan integral personalizado · Centro Carvajal</div>
    </div>
    <div>
      <div class="sec-label" style="margin-bottom:5px">Rutina de Cuidado en Casa</div>
      <div class="rut-grid">
        <div class="rut-box">
          <div class="rut-header am">☀️ Rutina Mañana</div>
          {{P5_RUTINA_AM}}
        </div>
        <div class="rut-box">
          <div class="rut-header pm">🌙 Rutina Noche</div>
          {{P5_RUTINA_PM}}
        </div>
      </div>
      {{P5_NOTAS_CRITICAS}}
    </div>
  </div>
  <div class="page-footer"><span><span class="pf-brand">Centro Carvajal · Líderes en Medicina Estética en Panamá</span></span><span>Página 7</span></div>
</div>

<!-- ══ PÁG 8: COMPROMISO ══ -->
<div class="page">
  <div class="page-header">
    <div class="ph-left"><div class="ph-dot"></div><div class="ph-title">Compromiso · Resultados · Próximos Pasos</div></div>
    <div class="ph-right">Método de Rejuvenecimiento Carvajal · {{NOMBRE}}</div>
  </div>
  <div class="content">
    <div class="sec-label">Sección Final</div>
    <div class="sec-title">Compromiso y Seguimiento</div>
    <div class="comp-box">
      <div class="comp-text" contenteditable="true">{{COMP_PARRAFO}}</div>
    </div>
    <div class="sec-title-sm">Resultados Esperados a 12 Meses</div>
    <div class="result-grid">{{COMP_RESULTADOS}}</div>
    <div class="sec-title-sm" style="margin-top:12px">Próximos Pasos Inmediatos</div>
    <div style="background:white;border:1px solid var(--border);border-radius:6px;padding:11px">
      <ul class="chk-list">{{COMP_PASOS}}</ul>
    </div>
    <div class="quote-box" style="margin-top:12px">
      <div class="quote-text" contenteditable="true">{{COMP_FRASE}}</div>
    </div>
  </div>
  <div style="background:var(--dark);padding:12px 28px;text-align:center;flex-shrink:0">
    <div style="color:var(--olive);font-size:9pt;font-weight:600;margin-bottom:3px">Centro Carvajal · Líderes en Medicina Estética en Panamá</div>
    <div style="font-size:7pt;color:rgba(255,255,255,0.28);line-height:1.7">centrocarvajal.com · Tel: 263-8134 &amp; 209-4284 · @centrocarvajal · Panamá<br>Revisado y validado por el equipo médico de Centro Carvajal.</div>
  </div>
</div>

<script>
const JOB_ID = "{{JOB_ID}}";

async function guardarCambios() {
  const btn = document.getElementById("btn-guardar");
  btn.disabled = true; btn.textContent = "⏳ Guardando...";
  try {
    const resp = await fetch("/guardar/" + JOB_ID, {
      method: "POST",
      headers: {"Content-Type": "text/html; charset=utf-8"},
      body: document.documentElement.outerHTML
    });
    const data = await resp.json();
    if (data.ok) { showToast("✓ Guardado en Cloudinary"); btn.textContent = "✓ Guardado"; }
    else { showToast("Error: " + (data.error || "?"), 4000); btn.textContent = "💾 Guardar cambios"; }
  } catch(e) { showToast("Error de conexión", 4000); btn.textContent = "💾 Guardar cambios"; }
  finally { setTimeout(() => { btn.disabled = false; btn.textContent = "💾 Guardar cambios"; }, 2800); }
}

function showToast(msg, dur=2500) {
  const t = document.getElementById("toast-cv");
  t.textContent = msg; t.classList.add("show");
  setTimeout(() => t.classList.remove("show"), dur);
}
</script>
</body>
</html>"""

def render_plan(j, d, job_id=''):
    tpl = PLANTILLA_PLAN_HTML

    def esc(s): return str(s).replace('&','&amp;').replace('<','&lt;').replace('>','&gt;').replace('"','&quot;')

    nombre = d.get('nombre', '')

    # ── Diagnóstico ──
    badge_map = {'warning':'<span class="badge-w">⚠</span> ','critical':'<span class="badge-c">✕ Crítico</span> ','normal':''}
    diag_html = ''.join(
        f'<tr><td class="diag-left">{badge_map.get(f.get("alerta","normal"),"")}{esc(f.get("area",""))}</td><td><span class="diag-val-strong" contenteditable="true">{esc(f.get("estado",""))}</span><span class="diag-val-sub" contenteditable="true">{esc(f.get("hallazgos",""))}</span></td></tr>'
        for f in j.get('diagnostico', {}).get('filas', [])
    )

    # ── Portada: pilares masonry ──
    pr = j.get('portada', {}).get('pilares_resumen', [{}]*5)
    def get_p(i): return pr[i] if i < len(pr) else {}
    p = [get_p(i) for i in range(5)]
    portada_pilares = (
        f'<div class="cp-row">'
        f'<div class="cp-item cp-wide"><div class="cp-icon">{p[0].get("emoji","🥗")}</div><div class="cp-num">Pilar 1</div><div class="cp-label" contenteditable="true">{esc(p[0].get("titulo",""))}</div></div>'
        f'<div class="cp-item cp-narrow"><div class="cp-icon">{p[1].get("emoji","🏃")}</div><div class="cp-num">Pilar 2</div><div class="cp-label" contenteditable="true">{esc(p[1].get("titulo",""))}</div></div>'
        f'</div>'
        f'<div class="cp-row">'
        f'<div class="cp-item"><div class="cp-icon">{p[2].get("emoji","🧠")}</div><div class="cp-num">Pilar 3</div><div class="cp-label" contenteditable="true">{esc(p[2].get("titulo",""))}</div></div>'
        f'<div class="cp-item"><div class="cp-icon">{p[3].get("emoji","😴")}</div><div class="cp-num">Pilar 4</div><div class="cp-label" contenteditable="true">{esc(p[3].get("titulo",""))}</div></div>'
        f'<div class="cp-item"><div class="cp-icon">{p[4].get("emoji","✨")}</div><div class="cp-num">Pilar 5</div><div class="cp-label" contenteditable="true">{esc(p[4].get("titulo",""))}</div></div>'
        f'</div>'
    )

    # ── Pág 2: pilares cards ──
    pilares_cards = ''.join(
        f'<div class="pilar-card"><div class="pilar-icon">{pi.get("emoji","")}</div><div><div class="pilar-title" contenteditable="true">Pilar {pi.get("num","")} · {esc(pi.get("titulo",""))}</div><div class="pilar-desc" contenteditable="true">{esc(pi.get("descripcion",""))}</div></div></div>'
        for pi in j.get('portada', {}).get('pilares_resumen', [])
    )

    # ── Rutina ──
    tag_css = {'Nutricion':'rtag-n','Nutrición':'rtag-n','Sueno':'rtag-s','Sueño':'rtag-s','Actividad':'rtag-a','Mental':'rtag-m','Estetico':'rtag-e','Estético':'rtag-e','Salud':'rtag-h','Imagen':'rtag-e','Trabajo':'rtag-n'}
    rutina_html = ''.join(
        f'<div class="rutina-row"><div class="rutina-hora">{esc(r["hora"])}</div><div class="rutina-text" contenteditable="true">{esc(r["actividad"])}</div><div class="rtag {tag_css.get(r["pilar"],"rtag-n")}">{esc(r["pilar"])}</div></div>'
        for r in j.get('rutina', {}).get('items', [])
    )

    # ── P1 Nutrición ──
    p1 = j.get('pilar1', {})
    p1_perm = ''.join(f'<div class="pe-item"><div class="pe-dot p"></div><span contenteditable="true">{esc(i)}</span></div>' for i in p1.get('permitidos',[]))
    p1_evit = ''.join(f'<div class="pe-item"><div class="pe-dot e"></div><span contenteditable="true">{esc(i)}</span></div>' for i in p1.get('evitar',[]))
    p1_menu = ''.join(
        f'<tr><td class="dia">{esc(m.get("dia",""))}</td><td contenteditable="true">{esc(m.get("desayuno",""))}</td><td contenteditable="true">{esc(m.get("almuerzo",""))}</td><td contenteditable="true">{esc(m.get("cena",""))}</td><td contenteditable="true">{esc(m.get("snack",""))}</td></tr>'
        for m in p1.get('menu',[])
    )
    p1_supl = ''.join(
        f'<div class="suppl-item"><div class="suppl-bullet">{i+1}</div><div><div class="suppl-name" contenteditable="true">{esc(s)}</div></div></div>'
        for i, s in enumerate(p1.get('suplementacion',[]))
    )

    # ── P2 Actividad ──
    p2 = j.get('pilar2', {})

    # ── P3 Bienestar ──
    p3 = j.get('pilar3', {})
    icons3 = ['🧘','📵','📓','🌿','🛁','👩‍⚕️']
    p3_cards = ''.join(
        f'<div class="mini-card"><div class="mc-icon">{icons3[i] if i < len(icons3) else "✦"}</div><div class="mc-title" contenteditable="true">{esc(t)}</div></div>'
        for i, t in enumerate(p3.get('tecnicas',[]))
    )

    # ── P4 Sueño ──
    p4 = j.get('pilar4', {})
    p4_proto_text = '<br>'.join(f'<strong>{i+1}.</strong> {esc(s)}' for i, s in enumerate(p4.get('protocolo',[])))
    p4_reglas_text = '<br>'.join(f'• {esc(r)}' for r in p4.get('reglas',[]))

    # ── P5 Tratamientos (split bimestres A=1-4, B=5-6) ──
    p5 = j.get('pilar5', {})
    bimestres = p5.get('bimestres',[])

    def render_bim(bim):
        rows = ''.join(
            f'<tr><td contenteditable="true"><strong>{esc(t.get("nombre",""))}</strong></td><td contenteditable="true">{esc(t.get("sesiones",""))}</td><td contenteditable="true"><strong>{esc(t.get("inversion",""))}</strong></td><td contenteditable="true">{esc(t.get("beneficio",""))}</td></tr>'
            for t in bim.get('tratamientos',[])
        )
        total = bim.get('total',0)
        return (
            f'<div class="bim-header">{esc(bim.get("periodo",""))} · {esc(bim.get("titulo",""))}<span style="font-size:7pt;color:rgba(143,168,50,0.6)">Bimestre {bim.get("bimestre","")}</span></div>'
            f'<div class="bim-body"><table class="bim-table"><thead><tr><th>Tratamiento</th><th style="width:75px">Sesiones</th><th style="width:65px">Inversión</th><th>Beneficio</th></tr></thead><tbody>{rows}</tbody></table>'
            f'<div class="bim-total">💰 Inversión: ${total:,}</div></div>'
        )

    # Enumerate bimestres to add bimestre number
    for idx, bim in enumerate(bimestres):
        bim['bimestre'] = idx + 1

    half = (len(bimestres) + 1) // 2  # primera mitad en pagina 6, resto en pagina 7
    p5_bim_a = ''.join(render_bim(b) for b in bimestres[:half])
    p5_bim_b = ''.join(render_bim(b) for b in bimestres[half:])

    p5_am = ''.join(
        f'<div class="rut-step"><div class="rut-num">{s.get("paso","")}</div><div><div class="rut-prod" contenteditable="true">{esc(s.get("producto",""))}</div><div class="rut-desc" contenteditable="true">{esc(s.get("descripcion",""))}</div></div></div>'
        for s in p5.get('rutina_am',[])
    )
    p5_pm = ''.join(
        f'<div class="rut-step"><div class="rut-num">{s.get("paso","")}</div><div><div class="rut-prod" contenteditable="true">{esc(s.get("producto",""))}</div><div class="rut-desc" contenteditable="true">{esc(s.get("descripcion",""))}</div></div></div>'
        for s in p5.get('rutina_pm',[])
    )
    notas = p5.get('notas_criticas',[])
    p5_notas = ''
    if notas:
        items = ''.join(f'<span contenteditable="true">{esc(n)}</span><br>' for n in notas)
        p5_notas = f'<div class="nota-medica" style="margin-top:9px"><strong>⚠ Notas Críticas</strong>{items}</div>'

    total_anual = p5.get('total_anual', 0)

    # ── Compromiso ──
    comp = j.get('compromiso', {})
    comp_res = ''.join(
        f'<div class="result-item"><div class="result-check">✓</div><span contenteditable="true">{esc(r["texto"])}</span></div>'
        for r in comp.get('resultados',[])
    )
    comp_pasos = ''.join(
        f'<li><div class="chk-ico"><svg viewBox="0 0 24 24"><polyline points="20 6 9 17 4 12"/></svg></div><span contenteditable="true">{esc(p_)}</span></li>'
        for p_ in comp.get('proximos_pasos',[])
    )
    # Closing quote from any pilar frase
    comp_frase = p5.get('frase_motivacional') or p3.get('frase_motivacional') or ''

    # Condición corta para portada
    cond = d.get('condicionSistemica','') or ''
    cond_corta = cond[:20] + '…' if len(cond) > 20 else cond

    replacements = {
        '{{JOB_ID}}': job_id,
        '{{NOMBRE}}': esc(nombre),
        '{{EDAD}}': esc(d.get('edad','')),
        '{{OCUPACION}}': esc(d.get('ocupacion','')),
        '{{FECHA}}': esc(d.get('fecha','')),
        '{{IMC}}': esc(d.get('imc','N/A')),
        '{{SATISFACCION}}': esc(str(d.get('satisfaccion','?'))),
        '{{CONDICION_CORTA}}': esc(cond_corta),
        '{{INTRO}}': esc(j.get('portada',{}).get('intro','')),
        '{{PILARES_PORTADA}}': portada_pilares,
        '{{DIAGNOSTICO_FILAS}}': diag_html,
        '{{PILARES_CARDS}}': pilares_cards,
        '{{RUTINA_FILAS}}': rutina_html,
        '{{P1_TITULO}}': esc(p1.get('titulo','Nutrición')),
        '{{P1_OBJETIVO}}': esc(p1.get('objetivo','')),
        '{{P1_PERMITIDOS}}': p1_perm,
        '{{P1_EVITAR}}': p1_evit,
        '{{P1_MENU}}': p1_menu,
        '{{P1_SUPLEMENTACION}}': p1_supl,
        '{{P2_TITULO}}': esc(p2.get('titulo','Actividad Física')),
        '{{P2_OBJETIVO}}': esc(p2.get('objetivo','')),
        '{{P2_PLAN}}': esc(p2.get('plan_semanal','')),
        '{{P2_ADAPTACIONES}}': esc(p2.get('adaptaciones','')),
        '{{P3_TITULO}}': esc(p3.get('titulo','Bienestar Mental')),
        '{{P3_OBJETIVO}}': esc(p3.get('objetivo','')),
        '{{P3_TECNICAS_CARDS}}': p3_cards,
        '{{P3_FRASE}}': esc(p3.get('frase_motivacional','')),
        '{{P4_TITULO}}': esc(p4.get('titulo','Optimización del Sueño')),
        '{{P4_OBJETIVO}}': esc(p4.get('objetivo','')),
        '{{P4_PROTOCOLO_TEXT}}': p4_proto_text,
        '{{P4_REGLAS_TEXT}}': p4_reglas_text,
        '{{P5_TITULO}}': esc(p5.get('titulo','Tratamientos')),
        '{{P5_OBJETIVO}}': esc(p5.get('objetivo','')),
        '{{P5_BIMESTRES_A}}': p5_bim_a,
        '{{P5_BIMESTRES_B}}': p5_bim_b,
        '{{P5_TOTAL_ANUAL}}': f'${total_anual:,}',
        '{{P5_RUTINA_AM}}': p5_am,
        '{{P5_RUTINA_PM}}': p5_pm,
        '{{P5_NOTAS_CRITICAS}}': p5_notas,
        '{{COMP_PARRAFO}}': esc(comp.get('parrafo','')),
        '{{COMP_RESULTADOS}}': comp_res,
        '{{COMP_PASOS}}': comp_pasos,
        '{{COMP_FRASE}}': esc(comp_frase),
    }
    for k, v in replacements.items():
        tpl = tpl.replace(k, v)
    return tpl


def generar_calendario():
    mes_names = ['Enero','Febrero','Marzo','Abril','Mayo','Junio','Julio','Agosto','Septiembre','Octubre','Noviembre','Diciembre']
    dias_sem  = ['Lu','Ma','Mi','Ju','Vi','Sa','Do']
    hoy = date.today()
    yr  = hoy.year
    mi  = hoy.month
    cal_html = ''
    row_html = ''
    for m in range(12):
        mes     = ((mi - 1 + m) % 12) + 1
        anio_m  = yr + (mi - 1 + m) // 12
        dias_en = calendar.monthrange(anio_m, mes)[1]
        primer  = date(anio_m, mes, 1).isoweekday()
        grid    = '<div class="cal-grid">'
        for d in dias_sem: grid += f'<div class="cal-dh">{d}</div>'
        for _ in range(1, primer): grid += '<div class="cal-d cal-empty"></div>'
        for dia in range(1, dias_en + 1):
            grid += f'<div class="cal-d"><div class="n">{dia}</div><div class="dots"><div class="dot nu"></div><div class="dot ac"></div><div class="dot me"></div><div class="dot su"></div></div></div>'
        grid += '</div>'
        row_html += f'<div class="cal-month"><div class="cal-mhdr">{mes_names[mes-1]} {anio_m}</div>{grid}</div>'
        if (m + 1) % 3 == 0:
            cal_html += f'<div class="cal-row">{row_html}</div>'
            row_html = ''
    if row_html: cal_html += f'<div class="cal-row">{row_html}</div>'
    return cal_html


# ════════════════════════════════════════════════════════════
# EMAILS con Resend
# ════════════════════════════════════════════════════════════

def email_alerta_creditos(error_raw, nombre_paciente=''):
    """HTML del correo de alerta cuando Claude se queda sin créditos."""
    fecha = datetime.now().strftime('%d/%m/%Y a las %H:%M')
    paciente_line = f' durante el procesamiento de <strong>{nombre_paciente}</strong>' if nombre_paciente else ''
    return f'''<!DOCTYPE html><html><head><meta charset="UTF-8"></head>
<body style="background:#f0e8de;padding:20px;font-family:sans-serif">
<div style="max-width:620px;margin:0 auto;background:#fff;border:1px solid #ddd">
  <div style="background:#1a1410;padding:20px 24px">
    <div style="color:#e05252;font-size:11px;letter-spacing:3px;text-transform:uppercase">Centro Carvajal · Alerta del Sistema</div>
    <div style="color:#fff;font-size:18px;margin-top:4px">⚠️ API de Claude sin créditos</div>
    <div style="color:rgba(255,255,255,.4);font-size:11px;margin-top:2px">{fecha}</div>
  </div>
  <div style="padding:24px">
    <p style="font-size:13px;color:#3d2e20;margin-bottom:16px">
      El sistema detectó un error de créditos al llamar a la API de Claude{paciente_line}.
      El plan y el análisis clínico fueron generados usando <strong>Groq (Llama 3.3 70B)</strong> como fallback automático.
    </p>
    <div style="background:#fff5f5;border:1px solid #fca5a5;border-radius:6px;padding:14px 18px;margin-bottom:20px">
      <div style="font-size:10px;font-weight:700;color:#991b1b;text-transform:uppercase;letter-spacing:1px;margin-bottom:8px">Error exacto recibido de Anthropic:</div>
      <code style="font-size:10px;color:#7f1d1d;word-break:break-all;display:block">{error_raw}</code>
    </div>
    <div style="background:#fffbf0;border:1px solid #e8d89a;border-radius:6px;padding:14px 18px;margin-bottom:20px">
      <div style="font-size:11px;font-weight:700;color:#8a7030;text-transform:uppercase;letter-spacing:1px;margin-bottom:8px">¿Qué significa este error?</div>
      <p style="font-size:12px;color:#5a4a20;margin:0">
        La cuenta de Anthropic se quedó sin créditos. El sistema no puede usar Claude hasta que se recargue el saldo.<br><br>
        <strong>Acción requerida:</strong> Ingresar a
        <a href="https://console.anthropic.com" style="color:#b8935a">console.anthropic.com</a>
        → Plans &amp; Billing → recargar créditos.
      </p>
    </div>
    <div style="background:#f0f7f0;border:1px solid #a8d5a2;border-radius:6px;padding:14px 18px">
      <div style="font-size:11px;font-weight:700;color:#2d6a2d;text-transform:uppercase;letter-spacing:1px;margin-bottom:6px">✅ El plan fue entregado igualmente</div>
      <p style="font-size:12px;color:#1a4a1a;margin:0">
        El sistema usó Groq automáticamente y el plan fue generado y enviado sin interrupciones al paciente.
      </p>
    </div>
  </div>
  <div style="background:#1a1410;padding:12px 24px;text-align:center;font-size:10px;color:rgba(255,255,255,0.3)">Centro Carvajal · centrocarvajal.com</div>
</div></body></html>'''


def enviar_resend(asunto, cuerpo, to, adjunto_path=None, adjunto_name=None, adjuntos_extra=None, cc=None):
    if not RESEND_KEY:
        print('RESEND_KEY no configurado')
        return
    payload = {
        'from': f'Centro Carvajal <{MAIL_FROM}>',
        'to': [to],
        'subject': asunto,
        'html': cuerpo,
    }
    if cc:
        payload['cc'] = cc if isinstance(cc, list) else [cc]
    attachments = []
    if adjunto_path and adjunto_name and os.path.exists(adjunto_path):
        import base64
        with open(adjunto_path, 'rb') as f:
            attachments.append({
                'filename': adjunto_name,
                'content': base64.b64encode(f.read()).decode(),
            })
    for fp in (adjuntos_extra or []):
        if os.path.exists(fp):
            import base64
            with open(fp, 'rb') as f:
                attachments.append({
                    'filename': os.path.basename(fp),
                    'content': base64.b64encode(f.read()).decode(),
                })
    if attachments:
        payload['attachments'] = attachments
    try:
        print(f'[resend] Enviando a to={to} cc={payload.get("cc")} from={payload.get("from")} asunto={payload.get("subject","")[:50]}')
        r = req.post('https://api.resend.com/emails',
            headers={'Authorization': f'Bearer {RESEND_KEY}', 'Content-Type': 'application/json'},
            json=payload, timeout=30)
        print(f'[resend] Respuesta {r.status_code}: {r.text[:200]}')
    except Exception as e:
        print(f'Resend error: {e}')


def email_formulario(d, faltantes, borrador_url=''):
    nombre = d.get('nombre', '')
    rows = ''.join(f'<tr><td style="color:#b8935a;font-weight:600;padding:8px 16px;font-size:12px;text-transform:uppercase;letter-spacing:1px;width:140px">{k}</td><td style="padding:8px 16px;font-size:13px">{v}</td></tr>'
        for k, v in [('Nombre', nombre), ('Edad', d.get('edad','')), ('Ocupacion', d.get('ocupacion','')),
                     ('Medicamentos', d.get('medicamentos','')), ('Cirugias', d.get('cirugias','')),
                     ('Prioridad', d.get('prioridad','')), ('Satisfaccion', d.get('satisfaccion','') + '/10')])
    faltantes_html = ''
    if faltantes:
        items = ''.join(f'<li style="font-size:12px;color:#6a5a20;padding:2px 0">{f}</li>' for f in faltantes)
        faltantes_html = f'<div style="background:#fffbf0;border:1px solid #e8d89a;border-radius:4px;padding:14px 18px;margin:16px 24px"><div style="font-size:11px;font-weight:700;color:#8a7030;margin-bottom:8px;text-transform:uppercase;letter-spacing:1px">Campos sin datos</div><ul style="padding-left:16px;margin:0">{items}</ul></div>'
    borrador_btn = f'<div style="text-align:center;margin:20px 24px"><a href="{borrador_url}" style="background:#8fa832;color:#fff;padding:13px 28px;border-radius:4px;text-decoration:none;font-size:14px;font-weight:500;display:inline-block">✏️ Revisar y editar borrador del plan</a><p style="font-size:11px;color:#999;margin-top:8px">Una vez editado, genera el PDF final desde el borrador.</p></div>' if borrador_url else ''
    return f'<!DOCTYPE html><html><head><meta charset="UTF-8"></head><body style="background:#f0e8de;padding:20px;font-family:sans-serif"><div style="max-width:600px;margin:0 auto;background:#fff;border:1px solid #ddd"><div style="background:#1a1410;padding:20px 24px"><div style="color:#b8935a;font-size:11px;letter-spacing:3px;text-transform:uppercase">Centro Carvajal · Nuevo Plan IA</div><div style="color:#fff;font-size:18px;margin-top:4px">{nombre}</div></div>{borrador_btn}<table style="width:100%;border-collapse:collapse">{rows}</table>{faltantes_html}<div style="background:#1a1410;padding:12px 24px;text-align:center;font-size:10px;color:rgba(255,255,255,0.3)">Centro Carvajal · centrocarvajal.com</div></div></body></html>'


def email_plan(nombre, html_url, fecha):
    return f'<!DOCTYPE html><html><head><meta charset="UTF-8"></head><body style="background:#f0e8de;padding:20px;font-family:sans-serif"><div style="max-width:600px;margin:0 auto;background:#fff;border:1px solid #ddd"><div style="background:#1a1410;padding:20px 24px"><div style="color:#b8935a;font-size:11px;letter-spacing:3px;text-transform:uppercase">Centro Carvajal · Plan Generado</div><div style="color:#fff;font-size:18px;margin-top:4px">{nombre}</div><div style="color:rgba(255,255,255,0.4);font-size:11px;margin-top:2px">{fecha}</div></div><div style="padding:24px"><p style="font-size:13px;color:#3d2e20;margin-bottom:16px">El plan personalizado de <strong>{nombre}</strong> ha sido generado exitosamente.</p><div style="text-align:center;margin:20px 0"><a href="{html_url}" style="background:#b8935a;color:#fff;padding:14px 32px;border-radius:4px;text-decoration:none;font-size:14px;font-weight:500">Ver Plan Completo</a></div><p style="font-size:11px;color:#999;text-align:center">O copia este link: {html_url}</p></div><div style="background:#1a1410;padding:12px 24px;text-align:center;font-size:10px;color:rgba(255,255,255,0.3)">Centro Carvajal · centrocarvajal.com</div></div></body></html>'



def email_formulario_inmediato(d, fotos=None):
    """Correo 1 — sale inmediatamente al recibir el formulario.
    Contiene TODOS los campos útiles del paciente con claves correctas.
    El .docx va adjunto. Es el respaldo completo ante cualquier fallo de la IA."""
    nombre = d.get('nombre', '')
    fecha  = datetime.now().strftime('%d/%m/%Y a las %H:%M')

    def row(k, v):
        if not v or str(v) in ('', '[]', 'None', '0', 'No registrado'): return ''
        if isinstance(v, list): v = ', '.join(str(i) for i in v if i)
        if isinstance(v, dict):
            parts = [f'{dk}: {dv}' for dk, dv in v.items() if dv and dv not in ('No','')]
            v = ' | '.join(parts) if parts else ''
        if not v: return ''
        return (f'<tr>'
                f'<td style="color:#b8935a;font-weight:600;padding:6px 14px;font-size:11px;'
                f'text-transform:uppercase;letter-spacing:1px;width:180px;vertical-align:top;'
                f'border-bottom:1px solid #f0e8de">{k}</td>'
                f'<td style="padding:6px 14px;font-size:13px;color:#2d2020;'
                f'border-bottom:1px solid #f0e8de">{v}</td>'
                f'</tr>')

    def seccion(titulo):
        return (f'<tr><td colspan="2" style="background:#1a1410;color:#b8935a;font-size:10px;'
                f'font-weight:700;letter-spacing:2px;text-transform:uppercase;padding:8px 14px">'
                f'{titulo}</td></tr>')

    contra = d.get('contraindications', {})
    contra_si = [k for k, v in contra.items() if v == 'Si'] if isinstance(contra, dict) else []

    # Alergias individuales (solo Medicamentos/Alimentos/Otro quedan, van en 'alergias' agregado)
    alergias_ind = []

    # Historial estético con detalle
    hist = d.get('historialEstetico', [])
    hist_det = d.get('historialDetalle', {})
    hist_str = ''
    if hist:
        partes = []
        for t in (hist if isinstance(hist, list) else [hist]):
            det = hist_det.get(t, {}) if isinstance(hist_det, dict) else {}
            s = t
            if det.get('fecha'): s += f' (última: {det["fecha"]})'
            if det.get('zona'):  s += f' — zona: {det["zona"]}'
            partes.append(s)
        hist_str = ' | '.join(partes)

    rows = ''.join([
        seccion('Datos personales'),
        row('Nombre',            nombre),
        row('Cédula',            d.get('cedula','')),
        row('Fecha nacimiento',  d.get('fechaNacimiento','')),
        row('Email',             d.get('email','')),
        row('Teléfono',          d.get('celular','')),
        row('Dirección',         d.get('direccion','')),
        row('Edad',              d.get('edad','')),
        row('Sexo',              d.get('sexo','')),
        row('Ocupación',         d.get('ocupacion','')),
        row('Horario laboral',   d.get('horarioLaboral','')),
        row('Hijos',             d.get('numHijos','')),
        row('Cómo nos conoció',  d.get('comoConociste','')),
        row('Contacto emergencia', d.get('contactoEmergencia','')),
        row('Relación',          d.get('contactoRelacion','')),
        row('Tel. emergencia',   d.get('contactoTel','')),

        seccion('Medidas'),
        row('Estatura',          str(d.get('estatura','')) + ' cm' if d.get('estatura') else ''),
        row('Peso',              str(d.get('peso','')) + ' kg' if d.get('peso') else ''),
        row('IMC',               d.get('imc','')),

        seccion('Salud general'),
        row('Condición sistémica', d.get('condicionSistemica','')),
        row('Condiciones',         d.get('condiciones','')),
        row('Medicamentos',        d.get('medicamentos','')),
        row('Cirugías',            d.get('cirugias','')),
        row('Fuma',                d.get('fuma','')),
        row('Alcohol',             d.get('alcohol','')),
        row('Antec. familiares',   d.get('antecedentesFam','')),
        row('Detalle antec.',      d.get('antecedentesFamDet','')),

        seccion('Contraindicaciones y alergias'),
        row('Contraindicaciones',  ', '.join(contra_si) if contra_si else ''),
        row('Embarazo',            d.get('embarazo','')),
        row('Lactancia',           d.get('lactancia','')),
        row('Anticonceptivos',     d.get('anticonceptivos','')),
        row('SOP',                 d.get('sop','')),
        row('Menopausia',          d.get('menopausia','')),
        row('Perimenopausia',      d.get('perimenopausia','')),
        row('Alergias',            d.get('alergias','')),
        row('Alergias específicas',', '.join(alergias_ind) if alergias_ind else ''),
        row('Detalle síntomas',    d.get('alergiasDetalle','')),

        seccion('Piel y estética'),
        row('Tipo de piel',        d.get('pielTipo','')),
        row('Problemas faciales',  d.get('pielProblemas','')),
        row('Áreas faciales',      d.get('areasFaciales','')),
        row('Áreas corporales',    d.get('areasCorporales','')),
        row('Rutina mañana',       d.get('rutinaManana','')),
        row('Rutina noche',        d.get('rutinaNoche','')),
        row('Productos frecuentes',d.get('productosFrecuentes','')),
        row('Protector solar',     d.get('usaProtectorSolar','')),
        row('Marca SPF',           d.get('protectorMarca','')),
        row('Factor SPF',          d.get('spf','')),
        row('Hora aplicación SPF', d.get('protectorHora','')),
        row('Reaplicación solar',  d.get('reaplicaSolar','')),
        row('Detalle láser',       d.get('laserActualDet','')),
        row('Complicaciones prev.',d.get('complicacionesDet','')),
        row('Historial estético',  hist_str or d.get('historialEstetico','')),

        seccion('Alimentación'),
        row('Síntomas digestivos', d.get('sintomasDigestivos','')),
        row('Notas alimentación',  d.get('notasAlimentacion','')),

        seccion('Hábitos y objetivos'),
        row('Act. física',         d.get('actFisica','')),
        row('Sueño',               d.get('sueno','')),
        row('Hora se despierta',   d.get('horaDespierta','')),
        row('Hora se duerme',      d.get('horaDuerme','')),
        row('Cansancio diurno',    d.get('cansancioDia','')),
        row('Nivel de estrés',     d.get('nivelEstres','')),
        row('Prioridad',           d.get('prioridad','')),
        row('Satisfacción actual', str(d.get('satisfaccion','')) + '/10' if d.get('satisfaccion') else ''),
    ])

    n_fotos = len([f for f in (fotos or []) if f])
    fotos_nota = (
        f'<div style="background:#f0f7e6;border:1px solid #c5d9a0;border-radius:4px;'
        f'padding:10px 16px;font-size:12px;color:#4a6020;margin:0">'
        f'{"📎 Fotos adjuntas: " + str(n_fotos) + " · " if n_fotos else ""}'
        f'El .docx con el cuestionario completo va adjunto a este correo.</div>'
    )

    return (
        '<!DOCTYPE html><html><head><meta charset="UTF-8"></head>'
        '<body style="background:#f0e8de;padding:20px;font-family:sans-serif">'
        '<div style="max-width:660px;margin:0 auto;background:#fff;border:1px solid #ddd">'
        '<div style="background:#1a1410;padding:20px 24px">'
        '<div style="color:#b8935a;font-size:11px;letter-spacing:3px;text-transform:uppercase">Centro Carvajal · Formulario Recibido</div>'
        f'<div style="color:#fff;font-size:18px;margin-top:4px">{nombre}</div>'
        f'<div style="color:rgba(255,255,255,0.4);font-size:11px;margin-top:2px">{fecha}</div>'
        '</div>'
        '<div style="background:#fffbf0;border-bottom:2px solid #e8d89a;padding:10px 16px;font-size:12px;color:#6a5a20">'
        '⚠ Respaldo completo del formulario. Si la IA falla, usa el .docx adjunto para generar el plan manualmente.'
        '</div>'
        f'{fotos_nota}'
        f'<table style="width:100%;border-collapse:collapse">{rows}</table>'
        '<div style="background:#1a1410;padding:12px 24px;text-align:center;font-size:10px;color:rgba(255,255,255,0.3)">'
        'Centro Carvajal · centrocarvajal.com</div>'
        '</div></body></html>'
    )

def email_plan_completo(d, borrador_url, html_url, faltantes=None):
    """Correo 2 — sale cuando el plan IA termino exitosamente.
    Contiene enlace al borrador editable y al plan final.
    El .docx con cuestionario + analisis clinico va adjunto."""
    nombre = d.get('nombre', '')
    fecha  = datetime.now().strftime('%d/%m/%Y a las %H:%M')

    borrador_btn = (
        f'<div style="text-align:center;margin:24px">'
        f'<a href="{borrador_url}" style="background:#8fa832;color:#fff;padding:13px 28px;'
        f'border-radius:4px;text-decoration:none;font-size:14px;font-weight:500;display:inline-block">'
        f'Revisar y editar borrador del plan</a>'
        f'<p style="font-size:11px;color:#999;margin-top:8px">Edita directamente en el navegador y genera el PDF final.</p>'
        f'</div>'
    ) if borrador_url else ''

    plan_btn = (
        f'<div style="text-align:center;margin:0 24px 24px">'
        f'<a href="{html_url}" style="background:#b8935a;color:#fff;padding:12px 28px;'
        f'border-radius:4px;text-decoration:none;font-size:13px;font-weight:500;display:inline-block">'
        f'Ver plan final</a>'
        f'</div>'
    ) if html_url else ''

    faltantes_html = ''
    if faltantes:
        items = ''.join(f'<li style="font-size:12px;color:#6a5a20;padding:2px 0">{f}</li>' for f in faltantes)
        faltantes_html = (
            f'<div style="background:#fffbf0;border:1px solid #e8d89a;border-radius:4px;'
            f'padding:14px 18px;margin:16px 24px">'
            f'<div style="font-size:11px;font-weight:700;color:#8a7030;margin-bottom:8px;'
            f'text-transform:uppercase;letter-spacing:1px">Campos sin datos</div>'
            f'<ul style="padding-left:16px;margin:0">{items}</ul></div>'
        )

    return (
        '<!DOCTYPE html><html><head><meta charset="UTF-8"></head>'
        '<body style="background:#f0e8de;padding:20px;font-family:sans-serif">'
        '<div style="max-width:620px;margin:0 auto;background:#fff;border:1px solid #ddd">'
        '<div style="background:#1a1410;padding:20px 24px">'
        '<div style="color:#b8935a;font-size:11px;letter-spacing:3px;text-transform:uppercase">Centro Carvajal - Plan IA Generado</div>'
        f'<div style="color:#fff;font-size:18px;margin-top:4px">{nombre}</div>'
        f'<div style="color:rgba(255,255,255,0.4);font-size:11px;margin-top:2px">{fecha}</div>'
        '</div>'
        '<div style="padding:20px 24px 8px;font-size:13px;color:#3d2e20;line-height:1.7">'
        f'El plan personalizado de <strong>{nombre}</strong> fue generado exitosamente. '
        'El .docx adjunto incluye el cuestionario completo y el analisis clinico.'
        '</div>'
        f'{borrador_btn}{plan_btn}{faltantes_html}'
        '<div style="background:#1a1410;padding:12px 24px;text-align:center;font-size:10px;color:rgba(255,255,255,0.3)">'
        'Centro Carvajal - centrocarvajal.com</div>'
        '</div></body></html>'
    )




# ════════════════════════════════════════════════════════════
# CORREO DE ERROR + RUTA DE REGENERACIÓN
# ════════════════════════════════════════════════════════════

def _enviar_correo_error_plan(nombre, session_id, error_msg=''):
    """Correo al staff cuando el worker falla — incluye botón para regenerar."""
    base_url = os.environ.get('BASE_URL', 'https://metodo.centrocarvajal.com')
    link = f'{base_url}/panel/regenerar/{session_id}'
    fecha = datetime.now().strftime('%d/%m/%Y a las %H:%M')
    cuerpo = (
        f'<!DOCTYPE html><html><head><meta charset="UTF-8"></head>'
        f'<body style="background:#f0e8de;padding:20px;font-family:sans-serif">'
        f'<div style="max-width:600px;margin:0 auto;background:#fff;border:1px solid #ddd">'
        f'<div style="background:#1a1410;padding:20px 24px">'
        f'<div style="color:#b8935a;font-size:11px;letter-spacing:3px;text-transform:uppercase">Centro Carvajal · Error generando plan</div>'
        f'<div style="color:#fff;font-size:18px;margin-top:4px">{nombre}</div>'
        f'<div style="color:rgba(255,255,255,0.4);font-size:11px;margin-top:2px">{fecha}</div>'
        f'</div>'
        f'<div style="padding:24px">'
        f'<div style="background:#fef2f2;border:1px solid #fecaca;border-radius:4px;padding:14px 18px;margin-bottom:20px">'
        f'<strong style="color:#c0392b;font-size:13px">Hubo un error al generar el plan de {nombre}.</strong><br>'
        f'<span style="font-size:12px;color:#666">El formulario fue recibido y guardado correctamente. '
        f'Puedes regenerar el plan con un solo clic.</span>'
        f'</div>'
        f'<div style="text-align:center;margin:24px 0">'
        f'<a href="{link}" style="background:#8fa832;color:#fff;padding:14px 32px;'
        f'border-radius:4px;text-decoration:none;font-size:15px;font-weight:500;display:inline-block">'
        f'🔄 Regenerar plan de {nombre}</a>'
        f'</div>'
        f'<p style="font-size:11px;color:#999;text-align:center">O copia este link: {link}</p>'
        f'<details style="margin-top:20px">'
        f'<summary style="font-size:11px;color:#999;cursor:pointer">Ver detalle del error</summary>'
        f'<pre style="font-size:10px;color:#666;background:#f9f9f9;padding:10px;border-radius:4px;margin-top:8px;overflow:auto">{error_msg[:500]}</pre>'
        f'</details>'
        f'</div>'
        f'<div style="background:#1a1410;padding:12px 24px;text-align:center;font-size:10px;color:rgba(255,255,255,0.3)">'
        f'Centro Carvajal · centrocarvajal.com</div>'
        f'</div></body></html>'
    )
    enviar_resend(
        f'⚠ Error generando plan — {nombre} ({fecha})',
        cuerpo,
        MAIL_TO,
        cc=MAIL_CC or None
    )
    print(f'[worker] Correo de error enviado: {nombre}')


@app.route('/panel/regenerar/<session_id>', methods=['GET'])
def panel_regenerar(session_id):
    """Página de regeneración — requiere login, redirige con next si no está autenticado."""
    from flask import make_response, redirect, url_for
    tok = request.cookies.get('cv_session', '')
    sesion = _verificar_sesion(tok)
    if not sesion:
        # Guardar destino en cookie temporal y redirigir a login
        resp = make_response(redirect(f'/panel?next=/panel/regenerar/{session_id}'))
        resp.set_cookie('cv_next', f'/panel/regenerar/{session_id}', max_age=300, httponly=True)
        return resp
    # Verificar que la sesión existe en Cloudinary
    datos_sesion = recuperar_sesion_cloudinary(session_id)
    nombre = datos_sesion.get('nombre', 'Paciente') if datos_sesion else 'Paciente'
    tiene_datos = datos_sesion is not None
    return REGENERAR_HTML.replace('{{SESSION_ID}}', session_id)                         .replace('{{NOMBRE}}', nombre)                         .replace('{{TIENE_DATOS}}', 'true' if tiene_datos else 'false')                         .replace('{{EMAIL}}', sesion['email'])


@app.route('/api/regenerar/<session_id>', methods=['POST'])
def api_regenerar(session_id):
    """Lanza la regeneración del plan desde el panel."""
    tok = request.cookies.get('cv_session', '')
    sesion = _verificar_sesion(tok)
    if not sesion:
        return jsonify({'ok': False, 'error': 'No autenticado'}), 401
    datos_sesion = recuperar_sesion_cloudinary(session_id)
    if not datos_sesion:
        return jsonify({'ok': False, 'error': 'Sesión no encontrada o expirada'}), 404
    # Limpiar clave interna del plan parcial si existe
    datos_sesion.pop('__plan_parcial__', None)
    job_id = uuid.uuid4().hex[:16]
    jobs[job_id] = {'status': 'working', 'msg': 'Iniciando regeneración del plan...'}
    modelo = request.get_json(silent=True, force=True).get('modelo', 'claude') if request.data else 'claude'
    t = threading.Thread(
        target=worker,
        args=(job_id, datos_sesion, [], [], modelo, session_id),
        daemon=True
    )
    t.start()
    print(f'[regenerar] Job {job_id} lanzado para sesión {session_id} por {sesion["email"]}')
    return jsonify({'ok': True, 'jobId': job_id, 'nombre': datos_sesion.get('nombre', '')})


REGENERAR_HTML = """<!DOCTYPE html>
<html lang="es">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>Regenerar Plan · Centro Carvajal</title>
<style>
:root{--dark:#1a1410;--olive:#8fa832;--gold:#b8935a;--cream:#f4f5ef;--muted:#6b7280;--red:#dc2626}
*{margin:0;padding:0;box-sizing:border-box}
body{font-family:'Segoe UI',system-ui,sans-serif;background:var(--cream);min-height:100vh;display:flex;flex-direction:column;align-items:center;justify-content:center;padding:20px}
.card{background:#fff;border-radius:12px;padding:40px;max-width:520px;width:100%;box-shadow:0 8px 40px rgba(0,0,0,.08);text-align:center}
.logo{font-size:11px;letter-spacing:3px;text-transform:uppercase;color:var(--gold);margin-bottom:8px}
h1{font-size:22px;color:var(--dark);margin-bottom:6px}
.nombre{font-size:26px;font-weight:600;color:var(--olive);margin-bottom:8px}
.desc{font-size:13px;color:var(--muted);margin-bottom:32px;line-height:1.6}
.btn{display:inline-block;padding:14px 32px;background:var(--olive);color:#fff;border:none;border-radius:8px;font-size:15px;font-weight:500;cursor:pointer;font-family:inherit;transition:background .2s;width:100%}
.btn:hover{background:#7a9428}
.btn:disabled{background:#9ca3af;cursor:not-allowed}
.btn-sec{background:transparent;border:1px solid #d4dcc0;color:var(--muted);margin-top:10px}
.btn-sec:hover{border-color:var(--olive);color:var(--olive);background:transparent}

/* Loading */
#screen-loading{display:none;margin-top:24px}
.ls-steps{display:flex;flex-direction:column;gap:10px;text-align:left;margin-top:20px}
.ls-item{display:flex;align-items:center;gap:10px;font-size:13px;color:var(--muted);padding:8px 12px;border-radius:6px;background:#f9f9f9}
.ls-item.active{color:var(--dark);background:rgba(143,168,50,.08);font-weight:500}
.ls-item.done{color:var(--olive)}
.ls-dot{width:8px;height:8px;border-radius:50%;background:#d4dcc0;flex-shrink:0}
.ls-item.active .ls-dot{background:var(--olive);animation:pulse 1s infinite}
.ls-item.done .ls-dot{background:var(--olive)}
@keyframes pulse{0%,100%{opacity:1}50%{opacity:.4}}
.timer{font-size:11px;color:var(--muted);margin-top:12px;text-align:center}

/* Success */
#screen-success{display:none;margin-top:8px}
.success-icon{width:60px;height:60px;border-radius:50%;background:rgba(143,168,50,.12);display:flex;align-items:center;justify-content:center;margin:0 auto 16px}
#screen-success h2{font-size:22px;color:var(--dark);margin-bottom:8px}
#screen-success p{font-size:13px;color:var(--muted);margin-bottom:24px}
.action-btns{display:flex;gap:10px;justify-content:center;flex-wrap:wrap}
.action-btn{padding:11px 22px;border-radius:6px;text-decoration:none;font-size:13px;font-weight:500;border:none;cursor:pointer;font-family:inherit}
.btn-ver{background:var(--olive);color:#fff}
.btn-editar{background:transparent;border:1px solid var(--olive);color:var(--olive)}

/* Error */
#screen-error{display:none;margin-top:8px}
.error-box{background:#fef2f2;border:1px solid #fecaca;border-radius:8px;padding:16px;margin-bottom:20px}
.error-box p{font-size:13px;color:var(--red)}

.no-datos{background:#fff7ed;border:1px solid #fed7aa;border-radius:8px;padding:16px;margin-bottom:24px}
.no-datos p{font-size:13px;color:#9a3412}
.user-tag{font-size:11px;color:var(--muted);margin-top:16px}
</style>
</head>
<body>
<div class="card">
  <div class="logo">Centro Carvajal</div>

  <!-- Estado inicial -->
  <div id="screen-init">
    <div class="nombre">{{NOMBRE}}</div>
    <h1>Regenerar plan personalizado</h1>
    <p class="desc" style="margin-top:8px">
      Los datos del formulario están guardados de forma segura.<br>
      Al dar clic el plan se generará automáticamente.
    </p>

    <div id="no-datos-msg" class="no-datos" style="display:none">
      <p>⚠ No se encontraron datos guardados para esta sesión. Puede que hayan expirado (72h). El paciente deberá llenar el formulario nuevamente.</p>
    </div>

    <button class="btn" id="btn-generar" onclick="iniciarGeneracion()" {{TIENE_DATOS == 'false' ? 'disabled' : ''}}>
      🔄 Generar plan ahora
    </button>
    <button class="btn btn-sec" onclick="window.location='/panel'">Volver al panel</button>
    <p class="user-tag">Sesión: {{EMAIL}}</p>
  </div>

  <!-- Loading -->
  <div id="screen-loading">
    <h1 style="margin-bottom:6px">Generando plan...</h1>
    <p class="desc">Esto tarda entre 2 y 4 minutos. Puedes dejar esta ventana abierta.</p>
    <div class="ls-steps">
      <div class="ls-item active" id="ls-1"><div class="ls-dot"></div> Iniciando generación del plan</div>
      <div class="ls-item" id="ls-2"><div class="ls-dot"></div> Sección 1/3 — Portada, diagnóstico y rutina</div>
      <div class="ls-item" id="ls-3"><div class="ls-dot"></div> Sección 2/3 — Nutrición, ejercicio y bienestar</div>
      <div class="ls-item" id="ls-4"><div class="ls-dot"></div> Sección 3/3 — Sueño, tratamientos y compromiso</div>
      <div class="ls-item" id="ls-5"><div class="ls-dot"></div> Enviando correo con plan completo</div>
    </div>
    <div class="timer" id="ls-timer">⏱ 0s transcurridos...</div>
  </div>

  <!-- Éxito -->
  <div id="screen-success">
    <div class="success-icon">
      <svg width="30" height="30" viewBox="0 0 30 30" fill="none"><path d="M6 15L12 21L24 9" stroke="#8fa832" stroke-width="2.5" stroke-linecap="round" stroke-linejoin="round"/></svg>
    </div>
    <h2>Plan generado correctamente</h2>
    <p id="success-desc">El plan de {{NOMBRE}} fue generado. El correo con el borrador y el plan completo ya llegó a la bandeja del equipo.</p>
    <div class="action-btns">
      <a id="btn-ver-plan" href="#" class="action-btn btn-ver" target="_blank">Ver plan</a>
      <a id="btn-editar-plan" href="#" class="action-btn btn-editar" target="_blank">Editar borrador</a>
    </div>
    <button class="btn btn-sec" onclick="window.location='/panel'" style="margin-top:20px;max-width:200px">Volver al panel</button>
  </div>

  <!-- Error -->
  <div id="screen-error">
    <div class="error-box">
      <p id="error-msg">Ocurrió un error al generar el plan.</p>
    </div>
    <button class="btn" onclick="iniciarGeneracion()">Reintentar</button>
    <button class="btn btn-sec" onclick="window.location='/panel'">Volver al panel</button>
  </div>
</div>

<script>
const SESSION_ID = '{{SESSION_ID}}';
const TIENE_DATOS = {{TIENE_DATOS}};
let jobId = null;
let pollInterval = null;
let tickerInterval = null;

document.addEventListener('DOMContentLoaded', () => {
  if (!TIENE_DATOS) {
    document.getElementById('no-datos-msg').style.display = 'block';
    document.getElementById('btn-generar').disabled = true;
  }
});

function show(id) {
  ['screen-init','screen-loading','screen-success','screen-error'].forEach(s => {
    document.getElementById(s).style.display = s === id ? 'block' : 'none';
  });
}

function iniciarGeneracion() {
  show('screen-loading');
  // Timer
  let secs = 0;
  tickerInterval = setInterval(() => {
    secs++;
    const m = Math.floor(secs/60), s = secs%60;
    document.getElementById('ls-timer').textContent =
      '⏱ ' + (m > 0 ? m + 'min ' + String(s).padStart(2,'0') + 's' : s + 's') + ' transcurridos...';
  }, 1000);

  fetch('/api/regenerar/' + SESSION_ID, {
    method: 'POST',
    headers: {'Content-Type': 'application/json'},
    body: JSON.stringify({modelo: 'claude'})
  })
  .then(r => r.json())
  .then(d => {
    if (!d.ok) { mostrarError(d.error || 'Error al iniciar'); return; }
    jobId = d.jobId;
    pollPlan();
  })
  .catch(e => mostrarError('Error de conexión: ' + e.message));
}

function pollPlan() {
  pollInterval = setInterval(async () => {
    try {
      const r = await fetch('/status?job=' + jobId);
      const d = await r.json();
      actualizarPasos(d.msg || '');
      if (d.status === 'done') {
        clearInterval(pollInterval); clearInterval(tickerInterval);
        mostrarExito(d);
      } else if (d.status === 'error') {
        clearInterval(pollInterval); clearInterval(tickerInterval);
        mostrarError(d.msg || 'Error desconocido');
      }
    } catch(e) { /* seguir intentando */ }
  }, 4000);
}

function actualizarPasos(msg) {
  const ls = s => document.getElementById(s);
  if (msg.includes('1/3')) {
    ls('ls-1').className='ls-item done'; ls('ls-2').className='ls-item active';
  } else if (msg.includes('2/3')) {
    ls('ls-2').className='ls-item done'; ls('ls-3').className='ls-item active';
  } else if (msg.includes('3/3')) {
    ls('ls-3').className='ls-item done'; ls('ls-4').className='ls-item active';
  } else if (msg.includes('correo') || msg.includes('Enviando')) {
    ls('ls-4').className='ls-item done'; ls('ls-5').className='ls-item active';
  }
}

function mostrarExito(d) {
  show('screen-success');
  if (d.html_url) document.getElementById('btn-ver-plan').href = d.html_url;
  if (d.borrador_url) document.getElementById('btn-editar-plan').href = d.borrador_url;
}

function mostrarError(msg) {
  show('screen-error');
  document.getElementById('error-msg').textContent = msg;
}
</script>
</body>
</html>"""


@app.route('/precios')
def precios():
    return render_template_string("""<!DOCTYPE html>
<html lang="es">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>Lista de Precios — Centro Carvajal</title>
<style>
:root{--dark:#1a1410;--olive:#8fa832;--olive-light:rgba(143,168,50,.1);--gold:#b8935a;--cream:#f4f5ef;--white:#fff;--gray:#6b7280;--border:rgba(143,168,50,.2);}
*{box-sizing:border-box;margin:0;padding:0}
body{font-family:'Segoe UI',system-ui,sans-serif;background:var(--cream);color:var(--dark);min-height:100vh;line-height:1.5}
.container{max-width:960px;margin:0 auto;padding:40px 20px}
header{text-align:center;margin-bottom:40px}
header h1{font-family:Georgia,serif;font-size:32px;font-weight:600;color:var(--dark);margin-bottom:6px}
header p{color:var(--gray);font-size:14px}
.section{margin-bottom:36px;background:var(--white);border:1px solid rgba(0,0,0,.06);border-radius:12px;overflow:hidden}
.section-header{background:var(--dark);padding:14px 20px;display:flex;align-items:center;gap:10px}
.section-header h2{font-size:15px;font-weight:600;color:var(--white);text-transform:uppercase;letter-spacing:.8px}
.badge{font-size:11px;font-weight:700;padding:3px 10px;border-radius:12px;background:var(--olive-light);color:var(--olive);text-transform:uppercase;letter-spacing:.5px}
table{width:100%;border-collapse:collapse}
th,td{padding:12px 16px;text-align:left;font-size:13px}
th{background:var(--olive-light);color:var(--olive);font-weight:600;font-size:11px;text-transform:uppercase;letter-spacing:.6px;border-bottom:1px solid var(--border)}
td{border-bottom:1px solid rgba(0,0,0,.04);color:var(--dark)}
tr:last-child td{border-bottom:none}
.price{font-weight:600;color:var(--dark);white-space:nowrap}
.sessions{color:var(--gray);font-size:12px}
.note{font-size:12px;color:var(--gray);padding:16px 20px;border-top:1px solid rgba(0,0,0,.04)}
.footer{text-align:center;padding:30px 0;font-size:12px;color:var(--gray)}
.footer a{color:var(--olive);text-decoration:none}
@media(max-width:640px){.container{padding:20px 12px}th,td{padding:10px 12px;font-size:12px}}
</style>
</head>
<body>
<div class="container">
<header>
  <h1>Lista de Precios</h1>
  <p>Centro Carvajal — Clínica de Medicina Estética — Panamá</p>
</header>

<div class="section">
  <div class="section-header"><h2>Faciales / Despigmentantes</h2><span class="badge">Rostro</span></div>
  <table>
    <tr><th>Tratamiento</th><th>Sesiones</th><th>Precio</th></tr>
    <tr><td>Cosmelan Kit</td><td class="sessions">Completo</td><td class="price">$600.00</td></tr>
    <tr><td>Cosmelan Mantenimiento</td><td class="sessions">Kit</td><td class="price">$300.00</td></tr>
    <tr><td>Melas Peel</td><td class="sessions">3 ses</td><td class="price">$200.00</td></tr>
    <tr><td>Regenerador Facial</td><td class="sessions">3 ses</td><td class="price">$613.00</td></tr>
    <tr><td>Regenerador Facial</td><td class="sessions">1 ses</td><td class="price">$313.00</td></tr>
    <tr><td>Fine Lift</td><td class="sessions">Paquete</td><td class="price">$999.00</td></tr>
    <tr><td>Skin Lift Pro</td><td class="sessions">1 ses</td><td class="price">$600.00</td></tr>
    <tr><td>Total Lift</td><td class="sessions">Paquete</td><td class="price">$1,300.00</td></tr>
    <tr><td>De Age Treatment</td><td class="sessions">Paquete</td><td class="price">$975.00</td></tr>
    <tr><td>Blanqueamiento Facial</td><td class="sessions">6 ses</td><td class="price">$266.00</td></tr>
    <tr><td>Plasma Facial</td><td class="sessions">1 ses</td><td class="price">$200.00</td></tr>
    <tr><td>Plasma Gel</td><td class="sessions">1 ses</td><td class="price">$250.00</td></tr>
    <tr><td>Peeling Periocular</td><td class="sessions">3 ses</td><td class="price">$151.00</td></tr>
    <tr><td>Acthyderm Rostro</td><td class="sessions">3 ses</td><td class="price">$334.00</td></tr>
    <tr><td>Péptidos Rejuvenecedores Rostro</td><td class="sessions">3 ses</td><td class="price">$544.00</td></tr>
    <tr><td>Péptidos Párpados</td><td class="sessions">3 ses</td><td class="price">$187.00</td></tr>
    <tr><td>Foto Facial</td><td class="sessions">3 ses</td><td class="price">$367.00</td></tr>
    <tr><td>Gleaming Skin</td><td class="sessions">6 ses</td><td class="price">$616.00</td></tr>
    <tr><td>Beauty Light</td><td class="sessions">2 ses</td><td class="price">$300.00</td></tr>
    <tr><td>Bright Eyes</td><td class="sessions">6 ses</td><td class="price">$241.00</td></tr>
    <tr><td>Hidratación Piel</td><td class="sessions">3 ses</td><td class="price">$236.00</td></tr>
    <tr><td>Vita C Peel</td><td class="sessions">3 ses</td><td class="price">$286.00</td></tr>
    <tr><td>Hidrofacial</td><td class="sessions">Por sesión</td><td class="price">$90.00</td></tr>
    <tr><td>Microdermoabrasión</td><td class="sessions">Por sesión</td><td class="price">$45.00</td></tr>
    <tr><td>Luz Anti-Acné</td><td class="sessions">3 ses</td><td class="price">$290.00</td></tr>
    <tr><td>Toxina Botulínica</td><td class="sessions">30 u</td><td class="price">$450.00</td></tr>
    <tr><td>Toxina Botulínica</td><td class="sessions">50 u</td><td class="price">$750.00</td></tr>
    <tr><td>Hilos PDO</td><td class="sessions">1 ses</td><td class="price">$800.00</td></tr>
    <tr><td>Rellenos / Ácido Hialurónico</td><td class="sessions">Criterio médico</td><td class="price">—</td></tr>
  </table>
</div>

<div class="section">
  <div class="section-header"><h2>Corporales</h2><span class="badge">Cuerpo</span></div>
  <table>
    <tr><th>Tratamiento</th><th>Sesiones</th><th>Precio</th></tr>
    <tr><td>EXILIS Abdomen</td><td class="sessions">8 ses</td><td class="price">$1,000.00</td></tr>
    <tr><td>Lipoláser</td><td class="sessions">10 ses</td><td class="price">$558.00</td></tr>
    <tr><td>Sculped Body</td><td class="sessions">12 ses</td><td class="price">$458.00</td></tr>
    <tr><td>Cellulite Shock / BTL X-Wave</td><td class="sessions">10 ses</td><td class="price">$790.00</td></tr>
    <tr><td>Electro Fit / Gimnasia Pasiva</td><td class="sessions">12 ses</td><td class="price">$408.00</td></tr>
    <tr><td>Tensor Cuerpo RF</td><td class="sessions">8 ses</td><td class="price">$808.00</td></tr>
    <tr><td>Acthyderm Cuerpo</td><td class="sessions">12 ses</td><td class="price">$783.00</td></tr>
    <tr><td>Post Parto</td><td class="sessions">10 ses</td><td class="price">$218.00</td></tr>
    <tr><td>Blanqueamiento Corporal</td><td class="sessions">6 ses</td><td class="price">$266.00</td></tr>
  </table>
</div>

<div class="section">
  <div class="section-header"><h2>Capilares</h2><span class="badge">Cabello</span></div>
  <table>
    <tr><th>Tratamiento</th><th>Sesiones</th><th>Precio</th></tr>
    <tr><td>Plasma Capilar</td><td class="sessions">2 ses</td><td class="price">$400.00</td></tr>
    <tr><td>Capilar Plus</td><td class="sessions">2 ses</td><td class="price">$499.00</td></tr>
  </table>
</div>

<div class="section">
  <div class="section-header"><h2>Depilación IPL</h2><span class="badge">Luz pulsada</span></div>
  <table>
    <tr><th>Zona</th><th>Sesiones</th><th>Precio</th></tr>
    <tr><td>IPL Facial</td><td class="sessions">6 ses</td><td class="price">$350.00</td></tr>
    <tr><td>IPL Axilas</td><td class="sessions">6 ses</td><td class="price">$350.00</td></tr>
    <tr><td>IPL Piernas</td><td class="sessions">8 ses</td><td class="price">$650.00</td></tr>
    <tr><td>IPL Brasileño</td><td class="sessions">8 ses</td><td class="price">$600.00</td></tr>
  </table>
  <div class="note">Garantía: 6 sesiones garantizadas. Si al finalizar no hay resultados, 6 sesiones adicionales gratuitas.</div>
</div>

<div class="footer">
  <p>Centro Carvajal · <a href="https://centrocarvajal.com">centrocarvajal.com</a> · Tel: 263-8134 / 209-4284</p>
  <p style="margin-top:4px;font-size:11px">Precios en USD/Balboas · Sujetos a cambio sin previo aviso · Consulte con nuestro equipo médico</p>
</div>
</div>
</body>
</html>""")

# Sincronizar catálogo de tratamientos al arrancar (local y producción)
try:
    _sincronizar_catalogo_inicial()
except Exception as e:
    print(f'[startup] Error sincronizando catálogo: {e}')

if __name__ == '__main__':
    # Inicializar usuarios al arrancar si no existen
    try:
        _inicializar_usuarios()
    except Exception as e:
        print(f'[startup] Error inicializando usuarios: {e}')
    app.run(debug=True, port=5000)
